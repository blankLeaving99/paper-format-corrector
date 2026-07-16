"""文档分析器测试"""

import pytest
from docx import Document
from docx.shared import Pt

from paper_format_corrector.parsers.document_analyzer import (
    DocumentAnalyzer,
    ParagraphType,
    analyze_document,
)


@pytest.fixture
def sample_doc(tmp_path):
    """创建测试文档"""
    doc = Document()

    # 添加标题
    title = doc.add_paragraph("论文题目：测试论文")
    title.runs[0].font.size = Pt(22)
    title.runs[0].font.bold = True
    title.runs[0].font.name = "黑体"  # 显式设置字体

    # 添加摘要
    doc.add_paragraph("摘  要")
    doc.add_paragraph("这是摘要内容。")

    # 添加正文
    doc.add_paragraph("第一章 绪论")
    body = doc.add_paragraph("这是正文第一段。")
    body.runs[0].font.name = "宋体"  # 显式设置字体
    doc.add_paragraph("这是正文第二段。")

    # 添加参考文献
    doc.add_paragraph("参考文献")
    doc.add_paragraph("[1] 张三. 论文标题[J]. 期刊名, 2020, 1(1): 1-10.")

    # 保存
    doc_path = tmp_path / "test.docx"
    doc.save(str(doc_path))
    return doc_path


@pytest.fixture
def empty_doc(tmp_path):
    """创建空文档"""
    doc = Document()
    doc_path = tmp_path / "empty.docx"
    doc.save(str(doc_path))
    return doc_path


class TestDocumentAnalyzer:
    """DocumentAnalyzer 测试"""

    def test_analyze_sample_doc(self, sample_doc):
        """测试分析样本文档"""
        analyzer = DocumentAnalyzer()
        result = analyzer.analyze(sample_doc)

        assert result.total_paragraphs > 0
        assert len(result.paragraphs) == result.total_paragraphs

    def test_paragraph_type_detection(self, sample_doc):
        """测试段落类型检测"""
        analyzer = DocumentAnalyzer()
        result = analyzer.analyze(sample_doc)

        # 检查是否检测到各类段落
        types_found = {p.paragraph_type for p in result.paragraphs}
        assert ParagraphType.BODY in types_found

    def test_font_extraction(self, sample_doc):
        """测试字体提取"""
        analyzer = DocumentAnalyzer()
        result = analyzer.analyze(sample_doc)

        # 应该检测到字体
        assert len(result.fonts_used) > 0

    def test_structure统计(self, sample_doc):
        """测试结构统计"""
        analyzer = DocumentAnalyzer()
        result = analyzer.analyze(sample_doc)

        # 应该有正文
        assert result.structure.get("body", 0) > 0

    def test_empty_doc(self, empty_doc):
        """测试空文档"""
        analyzer = DocumentAnalyzer()
        result = analyzer.analyze(empty_doc)

        assert result.total_paragraphs == 0
        assert len(result.paragraphs) == 0

    def test_custom_patterns(self, sample_doc):
        """测试自定义检测模式"""
        custom_patterns = {
            "title": r"^论文题目",
            "chapter": r"^第.+章",
            "body": r".*",  # 匹配所有
        }
        analyzer = DocumentAnalyzer(patterns=custom_patterns)
        result = analyzer.analyze(sample_doc)

        assert result.total_paragraphs > 0

    def test_compare_with_config(self, sample_doc):
        """测试与配置比较"""
        analyzer = DocumentAnalyzer()
        result = analyzer.analyze(sample_doc)

        config = {
            "format_rules": {
                "font": {"chinese": "宋体"},
                "body_text": {"font_size": 12},
                "margins": {"top": 2.54, "bottom": 2.54, "left": 3.17, "right": 3.17},
            }
        }

        diffs = analyzer.compare_with_config(result, config)
        # 应该返回差异列表（可能为空或有差异）
        assert isinstance(diffs, list)


class TestParagraphType:
    """ParagraphType 枚举测试"""

    def test_all_types_exist(self):
        """测试所有类型都存在"""
        assert ParagraphType.TITLE.value == "title"
        assert ParagraphType.BODY.value == "body"
        assert ParagraphType.HEADING1.value == "heading1"
        assert ParagraphType.REFERENCE.value == "reference"


class TestConvenienceFunction:
    """便捷函数测试"""

    def test_analyze_document(self, sample_doc):
        """测试 analyze_document"""
        result = analyze_document(sample_doc)

        assert result.total_paragraphs > 0
        assert len(result.paragraphs) > 0
