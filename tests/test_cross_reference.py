"""Tests for cross-reference updater and citation consistency checker."""

from docx import Document

from paper_format_corrector.core.document.cross_reference import (
    CitationConsistencyReport,
    CrossReferenceUpdater,
)


def test_cross_reference_updater_updates_fig_citations():
    doc = Document()
    para = doc.add_paragraph("如图1所示，系统架构如下。")
    updater = CrossReferenceUpdater()
    count = updater.update(doc, fig_map={"1": "2"})
    assert count == 1
    assert "图2" in para.text


def test_cross_reference_updater_updates_table_citations():
    doc = Document()
    para = doc.add_paragraph("见表3，数据对比。")
    updater = CrossReferenceUpdater()
    count = updater.update(doc, tab_map={"3": "5"})
    assert count == 1
    assert "表5" in para.text


def test_cross_reference_updater_updates_formula_citations():
    doc = Document()
    para = doc.add_paragraph("式(2-1)给出了公式推导。")
    updater = CrossReferenceUpdater()
    count = updater.update(doc, eq_map={"2-1": "3-1"})
    assert count == 1
    assert "式(3-1)" in para.text


def test_citation_consistency_detects_missing_in_references():
    doc = Document()
    # 正文引用
    doc.add_paragraph("研究表明[1]是正确的。")
    doc.add_paragraph("但[2]可能有误。")
    # 空段落
    doc.add_paragraph("")
    # 参考文献标题
    doc.add_paragraph("参考文献")
    # 只有 [1] 的条目
    doc.add_paragraph("[1] 张三. 论文标题[J]. 期刊, 2020.")

    updater = CrossReferenceUpdater()
    ref_start = 3  # "参考文献" 所在段落
    issues = updater.check_citation_consistency(doc, ref_start)

    # [2] 在参考文献中未找到
    missing_issues = [i for i in issues if i["type"] == "missing_in_references"]
    assert len(missing_issues) >= 1
    assert any("[2]" in i["message"] for i in missing_issues)


def test_citation_consistency_detects_missing_in_text():
    doc = Document()
    # 正文只引用了 [1]
    doc.add_paragraph("如文献[1]所述。")
    doc.add_paragraph("")
    doc.add_paragraph("参考文献")
    # 但参考文献有 [1] 和 [2]
    doc.add_paragraph("[1] 张三. 论文A[J]. 期刊, 2020.")
    doc.add_paragraph("[2] 李四. 论文B[J]. 期刊, 2021.")

    updater = CrossReferenceUpdater()
    issues = updater.check_citation_consistency(doc, 2)

    missing_in_text = [i for i in issues if i["type"] == "missing_in_text"]
    assert len(missing_in_text) >= 1
    assert any("[2]" in i["message"] for i in missing_in_text)


def test_citation_consistency_full_report():
    doc = Document()
    doc.add_paragraph("引用[1]和[2]。")
    doc.add_paragraph("")
    doc.add_paragraph("参考文献")
    doc.add_paragraph("[1] 张三. 论文A[J]. 2020.")
    doc.add_paragraph("[2] 李四. 论文B[J]. 2021.")
    doc.add_paragraph("[3] 王五. 论文C[J]. 2022.")

    updater = CrossReferenceUpdater()
    report = updater.get_full_consistency_report(doc, 2)

    assert isinstance(report, CitationConsistencyReport)
    assert report.total_text_citations == 2
    assert report.total_reference_items == 3
    assert len(report.matched_citations) == 2
    # [3] 未被正文引用
    assert any(i.issue_type == "missing_in_text" for i in report.issues)


def test_cross_reference_no_update_when_no_maps():
    doc = Document()
    doc.add_paragraph("如图1所示。")
    updater = CrossReferenceUpdater()
    count = updater.update(doc)
    assert count == 0
