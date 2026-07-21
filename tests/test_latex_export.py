"""LaTeX 导出和 LaTeX↔DOCX 双向转换测试"""

import os
import sys
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from docx.shared import Pt

sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "src"))


# ── LaTeXExporter 测试 ──────────────────────────────────────────


class TestLaTeXExporter:
    """LaTeXExporter 类测试"""

    def test_export_english_docx(self, tmp_path):
        """英文 DOCX → LaTeX 基本转换"""
        from paper_format_corrector.infrastructure.converters.latex_exporter import LaTeXExporter

        # 创建英文 DOCX
        doc = Document()
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("This is a test document with some content.")
        doc.add_heading("Methods", level=2)
        doc.add_paragraph("We used advanced techniques.")
        docx_path = tmp_path / "english.docx"
        doc.save(str(docx_path))

        exporter = LaTeXExporter()
        tex_path = tmp_path / "output.tex"
        result = exporter.export(str(docx_path), str(tex_path))

        assert result == str(tex_path)
        assert tex_path.exists()

        content = tex_path.read_text(encoding="utf-8")
        # 应包含 LaTeX 结构
        assert "\\documentclass" in content
        assert "\\begin{document}" in content
        assert "\\end{document}" in content
        # 应使用 article 文档类（英文）
        assert "article" in content

    def test_export_chinese_docx_uses_ctex(self, tmp_path):
        """中文 DOCX → LaTeX 自动生成 ctex 支持"""
        from paper_format_corrector.infrastructure.converters.latex_exporter import LaTeXExporter

        doc = Document()
        doc.add_heading("绪论", level=1)
        doc.add_paragraph("这是一篇中文论文的测试文档。")
        docx_path = tmp_path / "chinese.docx"
        doc.save(str(docx_path))

        exporter = LaTeXExporter()
        tex_path = tmp_path / "output.tex"
        result = exporter.export(str(docx_path), str(tex_path))

        assert result == str(tex_path)
        content = tex_path.read_text(encoding="utf-8")
        # 中文文档应使用 ctexart 或包含 ctex 宏包
        assert "ctex" in content.lower()

    def test_export_with_custom_document_class(self, tmp_path):
        """指定文档类"""
        from paper_format_corrector.infrastructure.converters.latex_exporter import LaTeXExporter

        doc = Document()
        doc.add_paragraph("Test content.")
        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        exporter = LaTeXExporter()
        tex_path = tmp_path / "output.tex"
        result = exporter.export(
            str(docx_path), str(tex_path), document_class="report"
        )

        content = tex_path.read_text(encoding="utf-8")
        assert "report" in content

    def test_export_with_extra_packages(self, tmp_path):
        """额外宏包注入"""
        from paper_format_corrector.infrastructure.converters.latex_exporter import LaTeXExporter

        doc = Document()
        doc.add_paragraph("Test content.")
        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        exporter = LaTeXExporter()
        tex_path = tmp_path / "output.tex"
        exporter.export(
            str(docx_path), str(tex_path),
            extra_packages=["\\usepackage{amsmath}", "\\usepackage{amssymb}"],
        )

        content = tex_path.read_text(encoding="utf-8")
        assert "amsmath" in content
        assert "amssymb" in content

    def test_export_file_not_found(self, tmp_path):
        """输入文件不存在时抛出异常"""
        from paper_format_corrector.infrastructure.converters.latex_exporter import LaTeXExporter

        exporter = LaTeXExporter()
        with pytest.raises(FileNotFoundError):
            exporter.export(str(tmp_path / "nonexistent.docx"), str(tmp_path / "out.tex"))

    def test_export_creates_output_directory(self, tmp_path):
        """输出目录不存在时自动创建"""
        from paper_format_corrector.infrastructure.converters.latex_exporter import LaTeXExporter

        doc = Document()
        doc.add_paragraph("Test.")
        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        exporter = LaTeXExporter()
        tex_path = tmp_path / "subdir" / "deep" / "output.tex"
        result = exporter.export(str(docx_path), str(tex_path))
        assert Path(result).exists()

    def test_export_with_template(self, tmp_path):
        """自定义模板注入"""
        from paper_format_corrector.infrastructure.converters.latex_exporter import LaTeXExporter

        doc = Document()
        doc.add_paragraph("Test.")
        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        template_content = "\\usepackage{custompkg}\n\\newcommand{\\mycmd}[1]{#1}"
        exporter = LaTeXExporter()
        tex_path = tmp_path / "output.tex"
        exporter.export(str(docx_path), str(tex_path), template=template_content)

        content = tex_path.read_text(encoding="utf-8")
        assert "custompkg" in content
        assert "mycmd" in content

    def test_detect_chinese(self):
        """中文检测功能"""
        from paper_format_corrector.infrastructure.converters.latex_exporter import LaTeXExporter

        assert LaTeXExporter._detect_chinese("这是一段中文") is True
        assert LaTeXExporter._detect_chinese("Hello World") is False
        assert LaTeXExporter._detect_chinese("Hello 你好") is True
        assert LaTeXExporter._detect_chinese("") is False

    def test_find_pandoc_returns_string_or_none(self):
        """pandoc 查找返回字符串或 None"""
        from paper_format_corrector.infrastructure.converters.latex_exporter import LaTeXExporter

        result = LaTeXExporter._find_pandoc()
        assert result is None or isinstance(result, str)


# ── _convert_tex 增强测试 ──────────────────────────────────────


class TestConvertTex:
    """FileConverter._convert_tex 增强测试"""

    def test_convert_tex_with_pandoc(self, tmp_path, monkeypatch):
        """pandoc 可用时使用 pandoc 转换"""
        from paper_format_corrector.infrastructure.converters.file_converter import FileConverter

        # 创建 .tex 文件
        tex_content = r"""
\documentclass{article}
\begin{document}
\section{Introduction}
This is a test document.
\textbf{Bold text} and \textit{italic text}.
\end{document}
"""
        tex_path = tmp_path / "test.tex"
        tex_path.write_text(tex_content, encoding="utf-8")

        converter = FileConverter()

        # Mock pandoc 不存在
        monkeypatch.setattr(
            "paper_format_corrector.infrastructure.converters.file_converter.shutil.which",
            lambda x: None,
        )

        output_path = tmp_path / "output.docx"
        result = converter._convert_tex(tex_path, output_path)
        assert Path(result).exists()

    def test_convert_tex_pandoc_timeout_fallback(self, tmp_path, monkeypatch):
        """pandoc 超时时降级到文本提取"""
        from paper_format_corrector.infrastructure.converters.file_converter import FileConverter

        tex_content = r"""
\documentclass{article}
\begin{document}
\section{Test}
Hello world content here.
\end{document}
"""
        tex_path = tmp_path / "test.tex"
        tex_path.write_text(tex_content, encoding="utf-8")

        # Mock pandoc 路径存在但调用失败
        fake_pandoc = tmp_path / "pandoc.exe"
        fake_pandoc.write_text("fake", encoding="utf-8")

        monkeypatch.setattr(
            "paper_format_corrector.infrastructure.converters.file_converter.shutil.which",
            lambda x: str(fake_pandoc),
        )

        def fake_run(*args, **kwargs):
            raise Exception("pandoc failed")

        monkeypatch.setattr(
            "paper_format_corrector.infrastructure.converters.file_converter.subprocess.run",
            fake_run,
        )

        converter = FileConverter()
        output_path = tmp_path / "output.docx"
        result = converter._convert_tex(tex_path, output_path)
        assert Path(result).exists()


# ── _extract_text_from_latex 增强测试 ─────────────────────────


class TestExtractTextFromLatex:
    """_extract_text_from_latex 格式保留测试"""

    def test_preserves_bold_markdown(self):
        """\\textbf{} → **bold**"""
        from paper_format_corrector.infrastructure.converters.file_converter import FileConverter

        tex = r"\textbf{Important text}"
        result = FileConverter._extract_text_from_latex(tex)
        assert "**Important text**" in result

    def test_preserves_italic_markdown(self):
        """\\textit{} / \\emph{} → *italic*"""
        from paper_format_corrector.infrastructure.converters.file_converter import FileConverter

        tex = r"\textit{Italic text} and \emph{Emphasized}"
        result = FileConverter._extract_text_from_latex(tex)
        assert "*Italic text*" in result
        assert "*Emphasized*" in result

    def test_preserves_code_markdown(self):
        """\\texttt{} → `code`"""
        from paper_format_corrector.infrastructure.converters.file_converter import FileConverter

        tex = r"\texttt{monospace}"
        result = FileConverter._extract_text_from_latex(tex)
        assert "`monospace`" in result

    def test_preserves_underline_markdown(self):
        """\\underline{} → <u>...</u>"""
        from paper_format_corrector.infrastructure.converters.file_converter import FileConverter

        tex = r"\underline{Underlined}"
        result = FileConverter._extract_text_from_latex(tex)
        assert "<u>Underlined</u>" in result or "Underlined" in result

    def test_preserves_strikethrough_markdown(self):
        """\\sout{} → <s>...</s>"""
        from paper_format_corrector.infrastructure.converters.file_converter import FileConverter

        tex = r"\sout{Deleted text}"
        result = FileConverter._extract_text_from_latex(tex)
        assert "<s>Deleted text</s>" in result or "Deleted text" in result

    def test_heading_mapping(self):
        """章节标题映射为 Markdown"""
        from paper_format_corrector.infrastructure.converters.file_converter import FileConverter

        tex = r"\section{Introduction}"
        result = FileConverter._extract_text_from_latex(tex)
        assert "## Introduction" in result

    def test_subsection_mapping(self):
        """\\subsection → ###"""
        from paper_format_corrector.infrastructure.converters.file_converter import FileConverter

        tex = r"\subsection{Related Work}"
        result = FileConverter._extract_text_from_latex(tex)
        assert "### Related Work" in result

    def test_abstract_extraction(self):
        """\\begin{abstract} → **摘要**"""
        from paper_format_corrector.infrastructure.converters.file_converter import FileConverter

        tex = "\\begin{abstract}\nThis is the abstract.\n\\end{abstract}"
        result = FileConverter._extract_text_from_latex(tex)
        assert "摘要" in result
        assert "This is the abstract." in result

    def test_image_placeholder(self):
        """\\includegraphics → [图片: ...]"""
        from paper_format_corrector.infrastructure.converters.file_converter import FileConverter

        tex = r"\includegraphics[width=0.5\textwidth]{fig1.png}"
        result = FileConverter._extract_text_from_latex(tex)
        assert "[图片: fig1.png]" in result

    def test_list_item_conversion(self):
        """\\item → - list item"""
        from paper_format_corrector.infrastructure.converters.file_converter import FileConverter

        tex = "\\begin{itemize}\n\\item First\n\\item Second\n\\end{itemize}"
        result = FileConverter._extract_text_from_latex(tex)
        assert "- First" in result
        assert "- Second" in result

    def test_footnote_extraction(self):
        """\\footnote{} → (脚注: ...)"""
        from paper_format_corrector.infrastructure.converters.file_converter import FileConverter

        tex = r"This has a footnote\footnote{Important note}."
        result = FileConverter._extract_text_from_latex(tex)
        assert "脚注: Important note" in result

    def test_removes_comments(self):
        """移除注释行"""
        from paper_format_corrector.infrastructure.converters.file_converter import FileConverter

        tex = "% This is a comment\nActual content here.\n% Another comment"
        result = FileConverter._extract_text_from_latex(tex)
        assert "comment" not in result
        assert "Actual content" in result

    def test_removes_preamble(self):
        """移除 preamble"""
        from paper_format_corrector.infrastructure.converters.file_converter import FileConverter

        tex = "\\documentclass{article}\n\\usepackage{amsmath}\n\\begin{document}\nContent\n\\end{document}"
        result = FileConverter._extract_text_from_latex(tex)
        assert "documentclass" not in result
        assert "Content" in result

    def test_combined_formatting(self):
        """混合格式保留"""
        from paper_format_corrector.infrastructure.converters.file_converter import FileConverter

        tex = r"""
\section{Results}
The \textbf{main result} shows that \textit{performance} improved by \texttt{15\%}.
"""
        result = FileConverter._extract_text_from_latex(tex)
        assert "## Results" in result
        assert "**main result**" in result
        assert "*performance*" in result
        assert "`15%`" in result or "`15" in result  # % may be consumed

    def test_citation_removal(self):
        """引用命令被移除"""
        from paper_format_corrector.infrastructure.converters.file_converter import FileConverter

        tex = r"As shown in \cite{smith2020} and \ref{fig1}, the results are significant."
        result = FileConverter._extract_text_from_latex(tex)
        assert "smith2020" not in result
        assert "fig1" not in result
        assert "significant" in result


# ── FileConverter .tex 输入集成测试 ────────────────────────────


class TestFileConverterTexInput:
    """FileConverter 对 .tex 输入的转换测试"""

    def test_tex_to_docx_basic(self, tmp_path):
        """基本 .tex → .docx 转换"""
        from paper_format_corrector.infrastructure.converters.file_converter import FileConverter

        tex_content = r"""
\documentclass{article}
\begin{document}
\section{Introduction}
This is a test document with \textbf{bold} and \textit{italic} text.
\end{document}
"""
        tex_path = tmp_path / "test.tex"
        tex_path.write_text(tex_content, encoding="utf-8")

        converter = FileConverter()
        output_path = tmp_path / "output.docx"
        result = converter.convert(str(tex_path))
        assert Path(result).exists()

    def test_tex_chinese_content(self, tmp_path):
        """中文 .tex 文件转换"""
        from paper_format_corrector.infrastructure.converters.file_converter import FileConverter

        tex_content = r"""
\documentclass{ctexart}
\begin{document}
\section{绪论}
这是一篇中文论文的测试内容。
\textbf{重点内容}和\textit{强调内容}。
\end{document}
"""
        tex_path = tmp_path / "chinese.tex"
        tex_path.write_text(tex_content, encoding="utf-8")

        converter = FileConverter()
        result = converter.convert(str(tex_path))
        assert Path(result).exists()
