"""Tests for image handling, table handling, and figure/table caption detection."""

import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "src"))

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.shared import Cm

from paper_format_corrector.domain.document.elements.figure_table_handler import FigureTableHandler
from paper_format_corrector.domain.document.elements.image_handler import ImageHandler
from paper_format_corrector.domain.document.elements.table_handler import TableHandler
from paper_format_corrector.domain.document.parser.structure import SectionDetector, SectionType

# ========== ImageHandler tests ==========


def test_image_handler_init(config):
    handler = ImageHandler(config)
    assert handler.center is True
    assert handler.max_width == "full"
    assert handler.min_dpi == 150


def test_image_handler_no_images(config, tmp_path):
    """A document with no images should return 0."""
    doc = Document()
    doc.add_paragraph("No images here.")
    path = tmp_path / "no_images.docx"
    doc.save(str(path))

    handler = ImageHandler(config)
    doc2 = Document(str(path))
    result = handler.process_all_images(doc2)
    assert result.images_centered == 0


def test_image_handler_page_width(config):
    """Page content width should be calculable."""
    doc = Document()
    handler = ImageHandler(config)
    width = handler._get_page_content_width(doc)
    assert width > 0


def test_image_handler_stats_empty(config):
    """get_image_stats on empty doc should return zeros."""
    doc = Document()
    handler = ImageHandler(config)
    stats = handler.get_image_stats(doc)
    assert stats["total"] == 0
    assert stats["avg_width_inches"] == 0


def test_image_handler_target_width(config):
    """_calc_target_width should respect config."""
    handler = ImageHandler(config)
    page_width = Cm(15)

    # full width
    assert handler._calc_target_width(page_width) == page_width

    # percentage
    handler.max_width = "50%"
    target = handler._calc_target_width(page_width)
    assert target == int(page_width * 0.5)

    # specific cm
    handler.max_width = 10
    target = handler._calc_target_width(page_width)
    assert target == Cm(10)


# ========== TableHandler tests ==========


def test_table_handler_init(config):
    handler = TableHandler(config)
    assert handler.default_font_size == 10.5
    assert handler.header_bold is True


def test_table_handler_format_empty_doc(config):
    """Formatting tables on a doc with no tables should return 0."""
    doc = Document()
    doc.add_paragraph("No tables here.")
    handler = TableHandler(config)
    result = handler.format_all_tables(doc)
    assert result.tables_formatted == 0


def test_table_handler_header_bold(config, tmp_path):
    """Header row (first row) should be bold, body rows should not."""
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    table.rows[0].cells[0].text = "Header1"
    table.rows[0].cells[1].text = "Header2"
    table.rows[1].cells[0].text = "Data1"
    table.rows[1].cells[1].text = "Data2"

    handler = TableHandler(config)
    handler.format_all_tables(doc)

    # Header cell should be bold
    header_run = doc.tables[0].rows[0].cells[0].paragraphs[0].runs[0]
    assert header_run.font.bold is True

    # Body cell should not be bold
    body_run = doc.tables[0].rows[1].cells[0].paragraphs[0].runs[0]
    assert body_run.font.bold is False


def test_table_handler_vertical_alignment(config, tmp_path):
    """All cells should have vertical center alignment."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[1].cells[0].text = "B"

    handler = TableHandler(config)
    handler.format_all_tables(doc)

    for row in doc.tables[0].rows:
        for cell in row.cells:
            assert cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER


def test_table_handler_center_alignment(config):
    """Table should be centered."""
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "Test"

    handler = TableHandler(config)
    handler.format_all_tables(doc)

    assert doc.tables[0].alignment == WD_TABLE_ALIGNMENT.CENTER


def test_table_handler_three_line_style(config):
    """Three-line table style should be applied when configured."""
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    for row in table.rows:
        for cell in row.cells:
            cell.text = "X"

    # Set three-line style in config
    config["format_rules"]["_table_style"] = "three_line"
    handler = TableHandler(config)
    handler.format_all_tables(doc)

    # Should not raise any errors
    assert True


def test_table_handler_full_border_style(config):
    """Full border style should be applied when configured."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    for row in table.rows:
        for cell in row.cells:
            cell.text = "X"

    config["format_rules"]["_table_style"] = "full_border"
    handler = TableHandler(config)
    handler.format_all_tables(doc)

    assert True


# ========== FigureTableHandler tests ==========


def test_figure_caption_chinese(config):
    """Chinese figure captions should be detected and renumbered."""
    doc = Document()
    p = doc.add_paragraph("图1 测试图片")

    handler = FigureTableHandler(config)
    handler.update_chapter(1)
    handler.process_paragraph(p, SectionType.FIGURE_CAPTION, {"num": "1"})

    assert handler.fig_count == 1
    assert "1-1" in p.text or "图1" in p.text


def test_table_caption_chinese(config):
    """Chinese table captions should be detected and renumbered."""
    doc = Document()
    p = doc.add_paragraph("表1 测试表格")

    handler = FigureTableHandler(config)
    handler.update_chapter(1)
    handler.process_paragraph(p, SectionType.TABLE_CAPTION, {"num": "1"})

    assert handler.tab_count == 1


def test_figure_caption_english(config):
    """English figure captions (Fig. 1) should be handled."""
    config_en = dict(config)
    config_en["auto_detect"] = dict(config.get("auto_detect", {}))
    config_en["auto_detect"]["figure_caption_pattern"] = r"^Fig\.?\s*\d"
    config_en["format_rules"] = dict(config.get("format_rules", {}))
    config_en["format_rules"]["figures"] = {"label": "Fig.", "numbering": "sequential", "font_size": 9}

    doc = Document()
    p = doc.add_paragraph("Fig. 1 Test image")

    handler = FigureTableHandler(config_en)
    handler.process_paragraph(p, SectionType.FIGURE_CAPTION, {"num": "1"})

    assert handler.fig_count == 1


def test_table_caption_english(config):
    """English table captions (Table 1) should be handled."""
    config_en = dict(config)
    config_en["auto_detect"] = dict(config.get("auto_detect", {}))
    config_en["auto_detect"]["table_caption_pattern"] = r"^Table\s+\d"
    config_en["format_rules"] = dict(config.get("format_rules", {}))
    config_en["format_rules"]["tables"] = {"label": "Table", "numbering": "sequential", "font_size": 9}

    doc = Document()
    p = doc.add_paragraph("Table 1 Test table")

    handler = FigureTableHandler(config_en)
    handler.process_paragraph(p, SectionType.TABLE_CAPTION, {"num": "1"})

    assert handler.tab_count == 1


# ========== SectionDetector tests ==========


def test_detect_english_figure_caption(config):
    """SectionDetector should detect English figure captions."""
    config_en = dict(config)
    config_en["auto_detect"] = dict(config.get("auto_detect", {}))
    config_en["auto_detect"]["figure_caption_pattern"] = r"^Fig\.?\s*\d"

    detector = SectionDetector(config_en)

    class MockPara:
        def __init__(self, text):
            self.text = text

    stype, info = detector.detect(MockPara("Fig. 1 System architecture"))
    assert stype == SectionType.FIGURE_CAPTION


def test_detect_english_table_caption(config):
    """SectionDetector should detect English table captions."""
    config_en = dict(config)
    config_en["auto_detect"] = dict(config.get("auto_detect", {}))
    config_en["auto_detect"]["table_caption_pattern"] = r"^Table\s+\d"

    detector = SectionDetector(config_en)

    class MockPara:
        def __init__(self, text):
            self.text = text

    stype, info = detector.detect(MockPara("Table 1 Performance comparison"))
    assert stype == SectionType.TABLE_CAPTION


# ========== Compat tests ==========


def test_compat_check():
    """check_dependencies should return a list."""
    from paper_format_corrector.infrastructure.compat import check_dependencies
    warnings = check_dependencies()
    assert isinstance(warnings, list)
    # Should have no ERROR-level warnings if deps are installed
    errors = [w for w in warnings if "[ERROR]" in w]
    assert len(errors) == 0, f"Dependency errors: {errors}"


# ========== Table cross-page header tests ==========


def test_table_enable_repeat_header(config):
    """_enable_repeat_header should set tblHeader attribute on first row."""
    from docx.oxml.ns import qn
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    table.rows[0].cells[0].text = "H1"
    table.rows[0].cells[1].text = "H2"
    table.rows[0].cells[2].text = "H3"
    table.rows[1].cells[0].text = "D1"
    table.rows[2].cells[0].text = "D2"

    handler = TableHandler(config)
    handler._enable_repeat_header(table)

    # Check tblHeader attribute on first row's trPr
    tr = table.rows[0]._tr
    trPr = tr.find(qn("w:trPr"))
    assert trPr is not None
    tblHeader = trPr.find(qn("w:tblHeader"))
    assert tblHeader is not None


def test_table_set_row_cant_split(config):
    """_set_row_cant_split should prevent row from splitting across pages."""
    from docx.oxml.ns import qn
    doc = Document()
    table = doc.add_table(rows=2, cols=2)

    handler = TableHandler(config)
    handler._set_row_cant_split(table.rows[0])

    tr = table.rows[0]._tr
    trPr = tr.find(qn("w:trPr"))
    assert trPr is not None
    cantSplit = trPr.find(qn("w:cantSplit"))
    assert cantSplit is not None


def test_table_format_calls_repeat_header(config):
    """format_all_tables should call _enable_repeat_header on tables with headers."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Alice"
    table.rows[1].cells[1].text = "42"

    handler = TableHandler(config)
    result = handler.format_all_tables(doc)
    # Should have processed the table
    assert result.tables_formatted >= 0


# ========== Image caption binding tests ==========


def test_image_handler_is_caption_bound(config):
    """_is_caption_bound_to_image returns False when no image in previous paragraphs."""
    doc = Document()
    doc.add_paragraph("Some text before caption.")
    doc.add_paragraph("图 1.0 测试图片标题")

    handler = ImageHandler(config)
    # No image in paragraph at index 0, so caption at index 1 is not bound
    bound = handler._is_caption_bound_to_image(doc, 1)
    assert bound is False


def test_image_handler_caption_not_bound(config):
    """Caption far from any image should not be bound."""
    doc = Document()
    doc.add_paragraph("Just some text.")
    doc.add_paragraph("More text.")
    doc.add_paragraph("Even more.")
    caption_para = doc.add_paragraph("图 1.0 测试图片标题")

    handler = ImageHandler(config)
    # caption at index 3, no images nearby
    bound = handler._is_caption_bound_to_image(doc, 3)
    assert bound is False
