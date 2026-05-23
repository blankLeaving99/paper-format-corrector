"""Tests for code block and formula content preservation."""

import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "src"))


import pytest
from docx import Document
from docx.shared import Pt

# ── Section detection ──────────────────────────────────────

class TestCodeDetection:
    def test_monospace_font_detected_as_code(self, config):
        """Paragraphs with Consolas font should be detected as CODE."""
        from paper_format_corrector.parsers.section_detector import SectionDetector, SectionType

        detector = SectionDetector(config)
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("def hello():\n    print('world')")
        run.font.name = "Consolas"
        run.font.size = Pt(10)

        section_type, _ = detector.detect(p)
        assert section_type == SectionType.CODE

    def test_courier_new_detected_as_code(self, config):
        """Paragraphs with Courier New font should be detected as CODE."""
        from paper_format_corrector.parsers.section_detector import SectionDetector, SectionType

        detector = SectionDetector(config)
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("int main() { return 0; }")
        run.font.name = "Courier New"

        section_type, _ = detector.detect(p)
        assert section_type == SectionType.CODE

    def test_indented_code_detected(self, config):
        """Indented text with code chars should be detected as CODE."""
        from paper_format_corrector.parsers.section_detector import SectionDetector, SectionType

        detector = SectionDetector(config)
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("    for i in range(10):")

        section_type, _ = detector.detect(p)
        assert section_type == SectionType.CODE

    def test_normal_text_not_detected_as_code(self, config):
        """Normal body text should not be detected as CODE."""
        from paper_format_corrector.parsers.section_detector import SectionDetector, SectionType

        detector = SectionDetector(config)
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("这是一段普通的正文内容，用于测试。")

        section_type, _ = detector.detect(p)
        assert section_type == SectionType.BODY

    def test_english_body_not_detected_as_code(self, config):
        """Normal English body text should not be detected as CODE."""
        from paper_format_corrector.parsers.section_detector import SectionDetector, SectionType

        detector = SectionDetector(config)
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("This is a normal paragraph in English that should not be code.")

        section_type, _ = detector.detect(p)
        assert section_type == SectionType.BODY


class TestFormulaContentDetection:
    def test_cambria_math_detected_as_formula(self, config):
        """Paragraphs with Cambria Math font should be detected as FORMULA_CONTENT."""
        from paper_format_corrector.parsers.section_detector import SectionDetector, SectionType

        detector = SectionDetector(config)
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("E = mc²")
        run.font.name = "Cambria Math"

        section_type, _ = detector.detect(p)
        assert section_type == SectionType.FORMULA_CONTENT

    def test_math_unicode_detected_as_formula(self, config):
        """Paragraphs with math Unicode should be detected as FORMULA_CONTENT."""
        from paper_format_corrector.parsers.section_detector import SectionDetector, SectionType

        detector = SectionDetector(config)
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("∑(xi − x̄)² / n")

        section_type, _ = detector.detect(p)
        assert section_type == SectionType.FORMULA_CONTENT

    def test_formula_numbering_still_detected(self, config):
        """Formula numbering lines like (1-1) should still be FORMULA."""
        from paper_format_corrector.parsers.section_detector import SectionDetector, SectionType

        detector = SectionDetector(config)
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("(1-1)")

        section_type, _ = detector.detect(p)
        assert section_type == SectionType.FORMULA

    def test_normal_text_not_detected_as_formula(self, config):
        """Normal text should not be detected as FORMULA_CONTENT."""
        from paper_format_corrector.parsers.section_detector import SectionDetector, SectionType

        detector = SectionDetector(config)
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("本文提出了一种新的算法。")

        section_type, _ = detector.detect(p)
        assert section_type == SectionType.BODY


# ── Format preservation ───────────────────────────────────

class TestCodeFormatPreservation:
    def test_code_font_not_changed(self, config, tmp_path):
        """Code paragraphs should preserve their original font."""
        from paper_format_corrector.core.format_corrector import FormatCorrector

        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("def hello():\n    print('world')")
        run.font.name = "Consolas"
        run.font.size = Pt(10)

        input_path = str(tmp_path / "input.docx")
        doc.save(input_path)

        corrector = FormatCorrector(str(tmp_path / "missing.docx"), config)
        corrector.load_template_styles()
        output_path = str(tmp_path / "output.docx")
        corrector.correct_document(input_path, output_path)

        result = Document(output_path)
        # Find the code paragraph
        for para in result.paragraphs:
            if "def hello" in para.text:
                for run in para.runs:
                    if run.text.strip():
                        assert run.font.name == "Consolas" or run.font.name is None
                return
        pytest.fail("Code paragraph not found in output")

    def test_code_line_spacing_not_changed(self, config, tmp_path):
        """Code paragraphs should preserve their original line spacing."""

        from paper_format_corrector.core.format_corrector import FormatCorrector

        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("line1\nline2\nline3")
        run.font.name = "Consolas"
        p.paragraph_format.line_spacing = 1.0

        input_path = str(tmp_path / "input.docx")
        doc.save(input_path)

        corrector = FormatCorrector(str(tmp_path / "missing.docx"), config)
        corrector.load_template_styles()
        output_path = str(tmp_path / "output.docx")
        corrector.correct_document(input_path, output_path)

        result = Document(output_path)
        for para in result.paragraphs:
            if "line1" in para.text:
                # Line spacing should be preserved (1.0), not changed to 1.5
                ls = para.paragraph_format.line_spacing
                assert ls is None or ls == 1.0 or ls == 1
                return
        pytest.fail("Code paragraph not found in output")

    def test_body_text_still_corrected(self, config, tmp_path):
        """Normal body text should still be corrected."""
        from paper_format_corrector.core.format_corrector import FormatCorrector

        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("这是一段普通的正文内容，应该被正常矫正。")
        run.font.name = "Arial"
        run.font.size = Pt(10)

        input_path = str(tmp_path / "input.docx")
        doc.save(input_path)

        corrector = FormatCorrector(str(tmp_path / "missing.docx"), config)
        corrector.load_template_styles()
        output_path = str(tmp_path / "output.docx")
        report = corrector.correct_document(input_path, output_path)

        assert report["body_fixed"] > 0


class TestFormulaContentPreservation:
    def test_formula_content_font_not_changed(self, config, tmp_path):
        """Formula content paragraphs should preserve their original font."""
        from paper_format_corrector.core.format_corrector import FormatCorrector

        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("E = mc²")
        run.font.name = "Cambria Math"
        run.font.size = Pt(12)

        input_path = str(tmp_path / "input.docx")
        doc.save(input_path)

        corrector = FormatCorrector(str(tmp_path / "missing.docx"), config)
        corrector.load_template_styles()
        output_path = str(tmp_path / "output.docx")
        corrector.correct_document(input_path, output_path)

        result = Document(output_path)
        for para in result.paragraphs:
            if "mc²" in para.text:
                for run in para.runs:
                    if run.text.strip():
                        assert run.font.name == "Cambria Math" or run.font.name is None
                return
        pytest.fail("Formula paragraph not found in output")
