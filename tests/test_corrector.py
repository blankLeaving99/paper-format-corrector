"""测试脚本：生成一个格式错误的样例论文，然后运行矫正器验证。"""

import sys
import os

# 添加 src 目录到 Python 路径
sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "src"))

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def create_sample_document(output_path="input/sample_paper.docx"):
    """生成一份格式不规范的样例论文"""
    doc = Document()

    # 页面边距故意设置不对
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # 题目（字号不对，字体不对）
    p = doc.add_paragraph()
    run = p.add_run("基于深度学习的论文格式自动矫正系统研究")
    run.font.size = Pt(14)
    run.font.name = "Arial"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 应该居中

    # 作者
    p = doc.add_paragraph()
    run = p.add_run("张三 李四")
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 摘要（格式不对）
    p = doc.add_paragraph()
    run = p.add_run("摘要")
    run.font.size = Pt(12)
    run.font.bold = False  # 应该加粗

    p = doc.add_paragraph()
    run = p.add_run(
        "本文提出了一种基于深度学习的论文格式自动矫正系统。该系统能够自动检测论文中的格式问题，"
        "包括字体、字号、行距、标题层级等，并进行自动修正。实验结果表明，该系统在多种论文格式上"
        "均能取得较好的矫正效果。"
    )
    run.font.size = Pt(10.5)  # 应该12pt
    run.font.name = "微软雅黑"  # 应该宋体

    # 关键词
    p = doc.add_paragraph()
    run = p.add_run("关键词：深度学习；论文格式；自动矫正；自然语言处理")
    run.font.size = Pt(10.5)

    # 第一章（格式不对）
    p = doc.add_paragraph()
    run = p.add_run("第一章 绪论")
    run.font.size = Pt(14)
    run.font.bold = False  # 应该加粗
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 应该居中

    # 1.1 节
    p = doc.add_paragraph()
    run = p.add_run("1.1 研究背景")
    run.font.size = Pt(12)

    # 正文（没有首行缩进，行距不对）
    p = doc.add_paragraph()
    run = p.add_run(
        "随着学术研究的不断发展，论文写作的规范性越来越受到重视。然而，许多研究者在撰写论文时，"
        "常常忽视格式规范，导致论文格式参差不齐。因此，开发一种自动化的论文格式矫正工具具有重要"
        "的现实意义。"
    )
    run.font.size = Pt(12)
    run.font.name = "宋体"

    # 1.2 节
    p = doc.add_paragraph()
    run = p.add_run("1.2 研究现状")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(
        "目前，关于论文格式矫正的研究主要集中在规则匹配和模板比对两个方面。"
        "规则匹配方法通过预定义的格式规则对论文进行检查和修正；模板比对方法则通过与标准模板进行"
        "对比，找出差异并进行修正。"
    )
    run.font.size = Pt(12)

    # 1.2.1 小节
    p = doc.add_paragraph()
    run = p.add_run("1.2.1 国内研究现状")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(
        "国内学者在论文格式自动化方面做了大量工作。王五等人提出了基于Word模板的格式检测方法，"
        "该方法能够检测字体、字号、行距等基本格式问题。"
    )
    run.font.size = Pt(12)

    # 第二章
    p = doc.add_paragraph()
    run = p.add_run("第二章 系统设计")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    run = p.add_run("2.1 系统架构")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(
        "本系统采用模块化设计，主要包括样式提取模块、格式检测模块、格式矫正模块三个部分。"
        "样式提取模块负责从模板文档中提取格式信息；格式检测模块负责检测待处理文档中的格式问题；"
        "格式矫正模块负责根据检测结果进行自动修正。"
    )
    run.font.size = Pt(12)

    # 图标题（编号不对）
    p = doc.add_paragraph()
    run = p.add_run("图2 系统架构图")
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 应该居中

    p = doc.add_paragraph()
    run = p.add_run("2.2 关键技术")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(
        "本系统的关键技术包括：（1）文档解析技术；（2）模式识别技术；（3）自动排版技术。"
    )
    run.font.size = Pt(12)

    # 表标题
    p = doc.add_paragraph()
    run = p.add_run("表1 系统性能对比")
    run.font.size = Pt(10)

    # 第三章
    p = doc.add_paragraph()
    run = p.add_run("第三章 实验结果")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    run = p.add_run("3.1 实验设置")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(
        "实验使用了100篇不同格式的论文作为测试集，包括期刊论文、会议论文和学位论文三种类型。"
        "评估指标包括格式正确率、矫正成功率和处理时间。"
    )
    run.font.size = Pt(12)

    # 参考文献
    p = doc.add_paragraph()
    run = p.add_run("参考文献")
    run.font.size = Pt(14)
    run.font.bold = True

    refs = [
        "[1] 王五, 赵六. 基于模板的论文格式检测方法[J]. 计算机应用, 2020, 40(5): 123-128.",
        "[2] Smith J, Brown K. Automatic Paper Formatting System[J]. IEEE Trans, 2019, 15(3): 45-52.",
        "[3] 陈七. 学术论文自动排版系统设计与实现[D]. 北京: 清华大学, 2021.",
        "[4] Wilson M. AI-based Document Formatting[C]//Proc of ACL. 2022: 100-108.",
        "[5] 李八. 自然语言处理在论文审查中的应用[J]. 软件学报, 2023, 34(2): 56-67.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(10)
        run.font.name = "Arial"

    doc.save(output_path)
    print(f"样例论文已生成: {output_path}")
    return output_path


def main():
    print("=" * 60)
    print("  论文格式矫正工具 - 测试")
    print("=" * 60)

    # 确保目录存在
    Path("input").mkdir(exist_ok=True)
    Path("output").mkdir(exist_ok=True)

    # 1. 生成样例论文
    print("\n[1/3] 生成格式不规范的样例论文...")
    sample_path = create_sample_document()

    # 2. 检查是否有模板
    template_path = Path("template/template.docx")
    if not template_path.exists():
        print("\n[2/3] 未找到模板文件，创建默认模板...")
        create_default_template(str(template_path))

    # 3. 运行矫正器
    print("\n[3/3] 运行格式矫正器...")
    from paper_format_corrector.app import PaperFormatCorrector

    corrector = PaperFormatCorrector()
    report = corrector.process_single(str(sample_path), "output/test_formatted.docx")

    if report:
        print("\n测试通过！")
    else:
        print("\n测试失败！")


def create_default_template(output_path):
    """创建一个默认模板"""
    from docx import Document
    from docx.shared import Pt, Cm

    doc = Document()

    # 正确的页面边距
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # Normal 样式
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Heading 1
    style = doc.styles["Heading 1"]
    style.font.size = Pt(16)
    style.font.bold = True

    # Heading 2
    style = doc.styles["Heading 2"]
    style.font.size = Pt(14)
    style.font.bold = True

    # Heading 3
    style = doc.styles["Heading 3"]
    style.font.size = Pt(12)
    style.font.bold = True

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"默认模板已创建: {output_path}")


if __name__ == "__main__":
    main()
