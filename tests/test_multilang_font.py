"""多语言字体测试

测试 _get_script_type()、_apply_mixed_font()、_set_run_font() 对中文/日文/韩文/英文的正确处理。
测试 get_east_asian_font() 工具方法的语言感知字体选择。
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "src"))

import pytest
from docx import Document

from paper_format_corrector.adapters.word.file_formatter import FormatCorrector


class TestGetScriptType:
    """测试 _get_script_type() 静态方法"""

    def test_chinese_characters(self):
        assert FormatCorrector._get_script_type("你") == "zh"
        assert FormatCorrector._get_script_type("世") == "zh"
        assert FormatCorrector._get_script_type("界") == "zh"

    def test_japanese_hiragana(self):
        assert FormatCorrector._get_script_type("あ") == "ja"
        assert FormatCorrector._get_script_type("ん") == "ja"
        assert FormatCorrector._get_script_type("は") == "ja"

    def test_japanese_katakana(self):
        assert FormatCorrector._get_script_type("ア") == "ja"
        assert FormatCorrector._get_script_type("ン") == "ja"
        assert FormatCorrector._get_script_type("カ") == "ja"

    def test_korean_hangul(self):
        assert FormatCorrector._get_script_type("가") == "ko"
        assert FormatCorrector._get_script_type("힣") == "ko"
        assert FormatCorrector._get_script_type("한") == "ko"

    def test_english_characters(self):
        assert FormatCorrector._get_script_type("A") == "en"
        assert FormatCorrector._get_script_type("z") == "en"
        assert FormatCorrector._get_script_type("H") == "en"

    def test_other_characters(self):
        assert FormatCorrector._get_script_type("1") == "other"
        assert FormatCorrector._get_script_type(" ") == "other"
        assert FormatCorrector._get_script_type(".") == "other"


class TestApplyMixedFont:
    """测试 _apply_mixed_font() 多语言字体应用"""

    @pytest.fixture
    def corrector(self):
        config = {
            "format_rules": {
                "font": {
                    "chinese": "宋体",
                    "japanese": "MS Mincho",
                    "korean": "Batang",
                    "english": "Times New Roman",
                }
            }
        }
        return FormatCorrector(template_path=None, config=config)

    def test_japanese_run_uses_japanese_font(self, corrector):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("こんにちは")
        font_rules = {
            "chinese": "宋体",
            "japanese": "MS Mincho",
            "korean": "Batang",
            "english": "Times New Roman",
        }
        style_rules = {"font_size": 12}
        corrector._apply_mixed_font(para, font_rules, style_rules)
        assert run.font.name == "MS Mincho"

    def test_korean_run_uses_korean_font(self, corrector):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("안녕하세요")
        font_rules = {
            "chinese": "宋体",
            "japanese": "MS Mincho",
            "korean": "Batang",
            "english": "Times New Roman",
        }
        style_rules = {"font_size": 12}
        corrector._apply_mixed_font(para, font_rules, style_rules)
        assert run.font.name == "Batang"

    def test_chinese_run_uses_chinese_font(self, corrector):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("你好世界")
        font_rules = {
            "chinese": "宋体",
            "japanese": "MS Mincho",
            "korean": "Batang",
            "english": "Times New Roman",
        }
        style_rules = {"font_size": 12}
        corrector._apply_mixed_font(para, font_rules, style_rules)
        assert run.font.name == "宋体"

    def test_english_run_uses_english_font(self, corrector):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Hello World")
        font_rules = {
            "chinese": "宋体",
            "japanese": "MS Mincho",
            "korean": "Batang",
            "english": "Times New Roman",
        }
        style_rules = {"font_size": 12}
        corrector._apply_mixed_font(para, font_rules, style_rules)
        assert run.font.name == "Times New Roman"

    def test_mixed_chinese_english_uses_dominant(self, corrector):
        doc = Document()
        para = doc.add_paragraph()
        # 中文为主，英文为辅
        run = para.add_run("这是一个测试test")
        font_rules = {
            "chinese": "宋体",
            "japanese": "MS Mincho",
            "korean": "Batang",
            "english": "Times New Roman",
        }
        style_rules = {"font_size": 12}
        corrector._apply_mixed_font(para, font_rules, style_rules)
        # 中文字符多于英文，应使用中文字体
        assert run.font.name == "宋体"


class TestSetRunFont:
    """测试 _set_run_font() 语言感知字体选择"""

    def test_chinese_language_uses_chinese_font(self):
        config = {"format_rules": {"font": {}}}
        corrector = FormatCorrector(template_path=None, config=config)
        corrector._detected_language = "chinese"

        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("测试")
        font_rules = {"chinese": "宋体", "english": "Times New Roman"}
        style_rules = {}

        corrector._set_run_font(run, font_rules, style_rules)
        assert run.font.name == "Times New Roman"

    def test_japanese_language_uses_japanese_font(self):
        config = {"format_rules": {"font": {}}}
        corrector = FormatCorrector(template_path=None, config=config)
        corrector._detected_language = "japanese"

        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("テスト")
        font_rules = {
            "chinese": "宋体",
            "japanese": "MS Mincho",
            "english": "Times New Roman",
        }
        style_rules = {}

        corrector._set_run_font(run, font_rules, style_rules)
        assert run.font.name == "Times New Roman"

    def test_korean_language_uses_korean_font(self):
        config = {"format_rules": {"font": {}}}
        corrector = FormatCorrector(template_path=None, config=config)
        corrector._detected_language = "korean"

        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("테스트")
        font_rules = {
            "chinese": "宋体",
            "korean": "Batang",
            "english": "Times New Roman",
        }
        style_rules = {}

        corrector._set_run_font(run, font_rules, style_rules)
        assert run.font.name == "Times New Roman"


class TestGetEastAsianFont:
    """测试 get_east_asian_font() 工具方法"""

    def test_chinese_body_default(self):
        from paper_format_corrector.shared.docx_utils import get_east_asian_font
        assert get_east_asian_font({}, "chinese", False) == "宋体"

    def test_chinese_heading_default(self):
        from paper_format_corrector.shared.docx_utils import get_east_asian_font
        assert get_east_asian_font({}, "chinese", True) == "黑体"

    def test_japanese_body_default(self):
        from paper_format_corrector.shared.docx_utils import get_east_asian_font
        assert get_east_asian_font({}, "japanese", False) == "MS Mincho"

    def test_japanese_heading_default(self):
        from paper_format_corrector.shared.docx_utils import get_east_asian_font
        assert get_east_asian_font({}, "japanese", True) == "MS Gothic"

    def test_korean_body_default(self):
        from paper_format_corrector.shared.docx_utils import get_east_asian_font
        assert get_east_asian_font({}, "korean", False) == "Batang"

    def test_korean_heading_default(self):
        from paper_format_corrector.shared.docx_utils import get_east_asian_font
        assert get_east_asian_font({}, "korean", True) == "Dotum"

    def test_chinese_with_custom_font_rules(self):
        from paper_format_corrector.shared.docx_utils import get_east_asian_font
        rules = {"chinese": "楷体", "heading_chinese": "仿宋"}
        assert get_east_asian_font(rules, "chinese", False) == "楷体"
        assert get_east_asian_font(rules, "chinese", True) == "仿宋"

    def test_japanese_with_custom_font_rules(self):
        from paper_format_corrector.shared.docx_utils import get_east_asian_font
        rules = {"japanese": "Yu Mincho", "japanese_heading": "Yu Gothic"}
        assert get_east_asian_font(rules, "japanese", False) == "Yu Mincho"
        assert get_east_asian_font(rules, "japanese", True) == "Yu Gothic"

    def test_korean_with_custom_font_rules(self):
        from paper_format_corrector.shared.docx_utils import get_east_asian_font
        rules = {"korean": "Gulim", "korean_heading": "DotumChe"}
        assert get_east_asian_font(rules, "korean", False) == "Gulim"
        assert get_east_asian_font(rules, "korean", True) == "DotumChe"

    def test_unknown_language_falls_back_to_chinese(self):
        from paper_format_corrector.shared.docx_utils import get_east_asian_font
        assert get_east_asian_font({}, "english", False) == "宋体"
        assert get_east_asian_font({}, "", False) == "宋体"

    def test_empty_font_rules_uses_defaults(self):
        from paper_format_corrector.shared.docx_utils import get_east_asian_font
        assert get_east_asian_font({}, "chinese") == "宋体"
        assert get_east_asian_font({}, "japanese") == "MS Mincho"
        assert get_east_asian_font({}, "korean") == "Batang"
