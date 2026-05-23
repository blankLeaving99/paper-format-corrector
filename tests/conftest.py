"""Shared pytest fixtures for paper-format-corrector tests."""

import os
import shutil
import sys

sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "src"))

from pathlib import Path

import pytest
import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


@pytest.fixture
def project_root():
    """Return the project root directory."""
    return Path(__file__).parent.parent


@pytest.fixture
def config(project_root):
    """Load the default config.yaml."""
    config_path = project_root / "config" / "config.yaml"
    with open(config_path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


@pytest.fixture
def template_path(project_root, tmp_path):
    """Create a default template docx and return its path."""
    src = project_root / "template" / "template.docx"
    if src.exists():
        dst = tmp_path / "template.docx"
        shutil.copy2(str(src), str(dst))
        return str(dst)
    # Create a minimal template if none exists
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for level, size in [(1, 16), (2, 14), (3, 12)]:
        style = doc.styles[f"Heading {level}"]
        style.font.size = Pt(size)
        style.font.bold = True

    dst = tmp_path / "template.docx"
    doc.save(str(dst))
    return str(dst)


@pytest.fixture
def sample_paper_path(tmp_path):
    """Create a format-incorrect sample paper and return its path."""
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Title (wrong font, wrong size, wrong alignment)
    p = doc.add_paragraph()
    run = p.add_run("基于深度学习的论文格式自动矫正系统研究")
    run.font.size = Pt(14)
    run.font.name = "Arial"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Author
    p = doc.add_paragraph()
    run = p.add_run("张三 李四")
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Abstract (wrong format)
    p = doc.add_paragraph()
    run = p.add_run("摘要")
    run.font.size = Pt(12)
    run.font.bold = False

    p = doc.add_paragraph()
    run = p.add_run(
        "本文提出了一种基于深度学习的论文格式自动矫正系统。该系统能够自动检测论文中的格式问题，"
        "包括字体、字号、行距、标题层级等，并进行自动修正。"
    )
    run.font.size = Pt(10.5)
    run.font.name = "微软雅黑"

    # Keywords
    p = doc.add_paragraph()
    run = p.add_run("关键词：深度学习；论文格式；自动矫正；自然语言处理")
    run.font.size = Pt(10.5)

    # Chapter 1 (wrong format)
    p = doc.add_paragraph()
    run = p.add_run("第一章 绪论")
    run.font.size = Pt(14)
    run.font.bold = False
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Section 1.1
    p = doc.add_paragraph()
    run = p.add_run("1.1 研究背景")
    run.font.size = Pt(12)

    # Body text (no indent, wrong spacing)
    p = doc.add_paragraph()
    run = p.add_run(
        "随着学术研究的不断发展，论文写作的规范性越来越受到重视。然而，许多研究者在撰写论文时，"
        "常常忽视格式规范，导致论文格式参差不齐。"
    )
    run.font.size = Pt(12)
    run.font.name = "宋体"

    # Section 1.2
    p = doc.add_paragraph()
    run = p.add_run("1.2 研究现状")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(
        "目前，关于论文格式矫正的研究主要集中在规则匹配和模板比对两个方面。"
    )
    run.font.size = Pt(12)

    # Subsection 1.2.1
    p = doc.add_paragraph()
    run = p.add_run("1.2.1 国内研究现状")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("国内学者在论文格式自动化方面做了大量工作。")
    run.font.size = Pt(12)

    # Chapter 2
    p = doc.add_paragraph()
    run = p.add_run("第二章 系统设计")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    run = p.add_run("2.1 系统架构")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("本系统采用模块化设计，主要包括样式提取模块、格式检测模块、格式矫正模块三个部分。")
    run.font.size = Pt(12)

    # Figure caption (wrong numbering)
    p = doc.add_paragraph()
    run = p.add_run("图2 系统架构图")
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Table caption
    p = doc.add_paragraph()
    run = p.add_run("表1 系统性能对比")
    run.font.size = Pt(10)

    # Chapter 3
    p = doc.add_paragraph()
    run = p.add_run("第三章 实验结果")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    run = p.add_run("3.1 实验设置")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("实验使用了100篇不同格式的论文作为测试集。")
    run.font.size = Pt(12)

    # References
    p = doc.add_paragraph()
    run = p.add_run("参考文献")
    run.font.size = Pt(14)
    run.font.bold = True

    refs = [
        "[1] 王五, 赵六. 基于模板的论文格式检测方法[J]. 计算机应用, 2020, 40(5): 123-128.",
        "[2] Smith J, Brown K. Automatic Paper Formatting System[J]. IEEE Trans, 2019, 15(3): 45-52.",
        "[3] 陈七. 学术论文自动排版系统设计与实现[D]. 北京: 清华大学, 2021.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(10)
        run.font.name = "Arial"

    path = tmp_path / "sample_paper.docx"
    doc.save(str(path))
    return str(path)


@pytest.fixture
def thesis_paper_path(tmp_path):
    """Create a thesis-style sample paper with abstracts, figures, tables, acknowledgment."""
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Cover
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("基于深度学习的中文文本情感分析研究")
    run.font.size = Pt(18)
    run.font.name = "Arial"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("张  三")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    run = p.add_run("计算机科学与技术学院")
    run.font.size = Pt(12)

    doc.add_page_break()

    # Chinese abstract
    p = doc.add_paragraph()
    run = p.add_run("摘  要")
    run.font.size = Pt(14)
    run.font.bold = False

    p = doc.add_paragraph()
    run = p.add_run(
        "随着互联网的快速发展，网络上的文本信息呈爆炸式增长。情感分析作为自然语言处理的重要分支，"
        "在舆情监控、产品评价分析等领域有着广泛的应用。"
    )
    run.font.size = Pt(10.5)
    run.font.name = "微软雅黑"

    p = doc.add_paragraph()
    run = p.add_run("关键词：深度学习；情感分析；BERT；注意力机制；自然语言处理")
    run.font.size = Pt(10.5)

    doc.add_page_break()

    # English abstract
    p = doc.add_paragraph()
    run = p.add_run("ABSTRACT")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    run = p.add_run(
        "With the rapid development of the Internet, text information on the web has "
        "experienced explosive growth. Sentiment analysis, as an important branch of "
        "natural language processing, has been widely applied in public opinion monitoring."
    )
    run.font.size = Pt(10.5)

    p = doc.add_paragraph()
    run = p.add_run("Key Words: Deep Learning; Sentiment Analysis; BERT; Attention Mechanism")
    run.font.size = Pt(10.5)

    doc.add_page_break()

    # Chapter 1
    p = doc.add_paragraph()
    run = p.add_run("第一章 绪论")
    run.font.size = Pt(14)
    run.font.bold = False
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    run = p.add_run("1.1 研究背景及意义")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(
        "近年来，随着社交媒体和电子商务的蓬勃发展，用户生成内容呈指数级增长。"
        "这些海量文本数据中蕴含着丰富的情感信息。"
    )
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("1.2 国内外研究现状")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("情感分析的研究可以追溯到早期基于规则和词典的方法。")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("1.2.1 基于词典的方法")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("基于情感词典的方法通过构建情感词典，利用词典中的情感极性信息对文本进行情感分类。")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("1.2.2 基于深度学习的方法")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("深度学习方法能够自动学习文本的特征表示，避免了人工特征工程的繁琐过程。")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("1.3 本文主要工作")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("本文的主要贡献包括：（1）提出了一种基于BERT的中文情感分析模型。")
    run.font.size = Pt(12)

    # Chapter 2
    p = doc.add_paragraph()
    run = p.add_run("第二章 相关理论与技术")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    run = p.add_run("2.1 预训练语言模型")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("预训练语言模型是近年来自然语言处理领域的重大突破。")
    run.font.size = Pt(12)

    # Table caption
    p = doc.add_paragraph()
    run = p.add_run("表2 BERT模型参数对比")
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Table
    table = doc.add_table(rows=4, cols=4)
    headers = ["模型", "层数", "隐藏维度", "参数量"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    # Figure caption
    p = doc.add_paragraph()
    run = p.add_run("图2 BERT模型架构")
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run("2.3 本章小结")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("本章介绍了预训练语言模型和注意力机制的基本原理。")
    run.font.size = Pt(12)

    # Chapter 3
    p = doc.add_paragraph()
    run = p.add_run("第三章 基于BERT的情感分析模型")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    run = p.add_run("3.1 模型架构")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("本文提出的模型基于BERT-Base架构，在其基础上添加了多头注意力层和全连接层。")
    run.font.size = Pt(12)

    # Table caption
    p = doc.add_paragraph()
    run = p.add_run("表1 模型超参数设置")
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run("3.2 训练策略")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("模型采用Adam优化器，学习率设置为2e-5。")
    run.font.size = Pt(12)

    # References
    p = doc.add_paragraph()
    run = p.add_run("参考文献")
    run.font.size = Pt(14)
    run.font.bold = True

    refs = [
        "[1] Devlin J, Chang M W, Lee K, et al. BERT: Pre-training of Deep Bidirectional Transformers[J]. NAACL, 2019: 4171-4186.",
        "[2] 王小明, 李红. 基于深度学习的中文情感分析综述[J]. 计算机学报, 2021, 44(6): 1125-1148.",
        "[3] Liu Y, Ott M, Goyal N, et al. RoBERTa: A Robustly Optimized BERT Pretraining Approach[J]. arXiv, 2019.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(10)

    # Acknowledgment
    p = doc.add_paragraph()
    run = p.add_run("致  谢")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    run = p.add_run("感谢我的导师在论文写作过程中给予的悉心指导和帮助。")
    run.font.size = Pt(12)

    path = tmp_path / "thesis_sample.docx"
    doc.save(str(path))
    return str(path)


@pytest.fixture
def requirement_natural_path(tmp_path):
    """Create a natural language requirement document."""
    content = """毕业论文格式要求

一、论文题目
论文题目用黑体，二号字，居中，加粗。

二、作者与单位
作者姓名用宋体，小四号，居中。
作者单位用宋体，五号，居中。

三、摘要
摘要标题用黑体，三号字，居中，加粗。
摘要正文用宋体，小四号，1.5倍行距，首行缩进2字符。

四、关键词
关键词用宋体，小四号，关键词三个字加粗。

五、一级标题
一级标题用黑体，二号字，居中，加粗。

六、二级标题
二级标题用黑体，三号字，左对齐，加粗。

七、三级标题
三级标题用黑体，四号字，左对齐，加粗。

八、正文
正文用宋体，小四号，1.5倍行距，首行缩进2字符，两端对齐。

九、参考文献
参考文献标题用黑体，三号字，居中，加粗。
参考文献正文用宋体，五号，1.25倍行距，悬挂缩进。

十、页边距
上边距2.54cm，下边距2.54cm，左边距3.17cm，右边距3.17cm。

十一、图标题
图标题用宋体，五号，居中。

十二、表标题
表标题用宋体，五号，居中。
"""
    path = tmp_path / "requirement_natural.txt"
    path.write_text(content, encoding="utf-8")
    return str(path)


@pytest.fixture
def requirement_table_path(tmp_path):
    """Create a table-format requirement document."""
    content = """毕业论文格式要求（表格版）

项目 | 字体 | 字号 | 对齐 | 加粗
论文题目 | 黑体 | 二号 | 居中 | 加粗
作者姓名 | 宋体 | 小四 | 居中 |
作者单位 | 宋体 | 五号 | 居中 |
摘要标题 | 黑体 | 三号 | 居中 | 加粗
摘要正文 | 宋体 | 小四 | |
一级标题 | 黑体 | 二号 | 居中 | 加粗
二级标题 | 黑体 | 三号 | 左对齐 | 加粗
三级标题 | 黑体 | 四号 | 左对齐 | 加粗
正文 | 宋体 | 小四 | 两端对齐 |
参考文献标题 | 黑体 | 三号 | 居中 | 加粗
参考文献正文 | 宋体 | 五号 | |
图标题 | 宋体 | 五号 | 居中 |
表标题 | 宋体 | 五号 | 居中 |

行距要求：
- 正文：1.5倍行距
- 摘要正文：1.5倍行距
- 参考文献正文：1.25倍行距

缩进要求：
- 正文：首行缩进2字符
- 摘要正文：首行缩进2字符

页边距：上下2.54cm，左右3.17cm
"""
    path = tmp_path / "requirement_table.txt"
    path.write_text(content, encoding="utf-8")
    return str(path)


@pytest.fixture
def requirement_mixed_path(tmp_path):
    """Create a mixed-format (.docx) requirement document."""
    doc = Document()
    doc.add_heading("毕业论文格式要求", level=1)
    doc.add_paragraph("以下为论文各部分的格式规范：")

    table = doc.add_table(rows=9, cols=5)
    table.style = "Table Grid"

    headers = ["项目", "字体", "字号", "对齐", "加粗"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    data = [
        ["论文题目", "黑体", "二号", "居中", "加粗"],
        ["一级标题", "黑体", "二号", "居中", "加粗"],
        ["二级标题", "黑体", "三号", "左对齐", "加粗"],
        ["三级标题", "黑体", "四号", "左对齐", "加粗"],
        ["正文", "宋体", "小四", "两端对齐", ""],
        ["摘要正文", "宋体", "小四", "", ""],
        ["关键词", "宋体", "小四", "", "加粗"],
        ["参考文献正文", "宋体", "五号", "", ""],
    ]
    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            table.rows[r + 1].cells[c].text = val

    doc.add_paragraph("")
    doc.add_paragraph("正文行距：1.5倍行距")
    doc.add_paragraph("参考文献行距：1.25倍行距")
    doc.add_paragraph("正文首行缩进2字符")
    doc.add_paragraph("页边距：上2.54cm，下2.54cm，左3.17cm，右3.17cm")

    path = tmp_path / "requirement_mixed.docx"
    doc.save(str(path))
    return str(path)


@pytest.fixture
def thesis_requirement_path(tmp_path):
    """Create a thesis format requirement document."""
    content = """XX大学毕业论文（设计）格式要求

一、论文题目
论文题目用黑体，二号字，居中，加粗。

二、作者信息
作者姓名用宋体，小四号，居中。
作者单位用宋体，五号，居中。

三、摘要
中文摘要标题用黑体，三号字，居中，加粗。
中文摘要正文用宋体，小四号，1.5倍行距。
英文摘要标题用Times New Roman，三号字，居中，加粗。
英文摘要正文用Times New Roman，小四号，1.5倍行距。

四、关键词
关键词标题加粗，关键词内容用宋体，小四号。

五、一级标题
一级标题用黑体，二号字，居中，加粗。

六、二级标题
二级标题用黑体，三号字，左对齐，加粗。

七、三级标题
三级标题用黑体，四号字，左对齐，加粗。

八、正文
正文用宋体，小四号（12pt），1.5倍行距，首行缩进2字符，两端对齐。

九、图标题
图标题用宋体，五号，居中。

十、表标题
表标题用宋体，五号，居中。

十一、参考文献
参考文献标题用黑体，三号字，居中，加粗。
参考文献正文用宋体，五号，1.25倍行距，悬挂缩进。

十二、致谢
致谢标题同一级标题格式。
致谢正文同正文格式。

十三、页边距
上边距2.54cm，下边距2.54cm，左边距3.17cm，右边距3.17cm。

十四、页眉页脚
页眉显示论文题目，宋体，五号，居中，有下划线。
页脚居中显示页码。
"""
    path = tmp_path / "thesis_requirement.txt"
    path.write_text(content, encoding="utf-8")
    return str(path)
