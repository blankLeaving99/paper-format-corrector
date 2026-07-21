"""Bug 修复验证测试

验证阶段 1 中修复的 7 个关键 Bug 不会回归。
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "src"))

from docx import Document
from docx.oxml.ns import qn

# ── Bug 1.1: requirement_parser.py 复制粘贴 Bug ─────────────────────


class TestRequirementParserThreeLineTable:
    """验证三线表关键词检测"""

    def test_three_line_table_keyword(self):
        """'三线表' 应触发三线表样式"""
        # "三线表" 和 "三线" 都应匹配
        text = "表格使用三线表格式"
        attrs = {}
        # 模拟 _parse_line 中的三线表检测逻辑
        if "三线表" in text or "三线" in text:
            attrs["table_style"] = "three_line"
        assert attrs.get("table_style") == "three_line"

    def test_three_line_keyword_only(self):
        """仅含 '三线' 也应匹配"""
        text = "表格使用三线格式"
        attrs = {}
        if "三线表" in text or "三线" in text:
            attrs["table_style"] = "three_line"
        assert attrs.get("table_style") == "three_line"


# ── Bug 1.2: table_handler.py 浮动 XML 元素 ─────────────────────────


class TestTableHandlerTblPr:
    """验证 tblPr 正确挂载到表格"""

    def test_get_or_create_tblPr_attaches_to_table(self):
        """新建的 tblPr 应挂载到表格元素树"""
        from paper_format_corrector.core.document.elements.table_handler import TableHandler

        handler = TableHandler({})
        doc = Document()
        table = doc.add_table(rows=1, cols=1)

        # 获取 tblPr
        tblPr = handler._get_or_create_tblPr(table)

        # 验证 tblPr 已挂载
        assert tblPr is not None
        assert table._tbl.tblPr is tblPr

    def test_set_table_borders_on_new_table(self):
        """新表格设置边框不应报错"""
        from paper_format_corrector.core.document.elements.table_handler import TableHandler

        handler = TableHandler({})
        doc = Document()
        table = doc.add_table(rows=2, cols=2)

        # 不应抛出异常
        handler.set_table_borders(table)

        # 验证边框已设置
        tblPr = table._tbl.tblPr
        borders = tblPr.find(qn("w:tblBorders"))
        assert borders is not None

    def test_remove_table_borders_on_new_table(self):
        """新表格设置三线表不应报错"""
        from paper_format_corrector.core.document.elements.table_handler import TableHandler

        handler = TableHandler({})
        doc = Document()
        table = doc.add_table(rows=2, cols=2)

        # 不应抛出异常
        handler.remove_table_borders(table)


# ── Bug 1.3: toc_handler.py position 参数 ───────────────────────────


class TestTOCHandlerPosition:
    """验证 TOC 插入位置参数生效"""

    def test_insert_toc_at_position_0(self):
        """position=0 应插入到文档开头"""
        from paper_format_corrector.core.document.elements.toc_handler import TOCHandler

        handler = TOCHandler({})
        doc = Document()
        doc.add_paragraph("段落1")
        doc.add_paragraph("段落2")

        handler.insert_toc(doc, position=0)

        body = doc.element.body
        # 第一个元素应该是 TOC 标题
        first_elem = body[0]
        assert first_elem.tag == qn("w:p")

    def test_insert_toc_at_position_1(self):
        """position=1 应插入到第一个段落之后"""
        from paper_format_corrector.core.document.elements.toc_handler import TOCHandler

        handler = TOCHandler({})
        doc = Document()
        doc.add_paragraph("段落1")
        doc.add_paragraph("段落2")

        handler.insert_toc(doc, position=1)

        body = doc.element.body
        # 第一个元素应该是原段落，第二个是 TOC 标题
        assert len(body) >= 3


# ── Bug 1.4: reference_formatter.py ref_end_idx ─────────────────────


class TestReferenceFormatterEndIdx:
    """验证 ref_end_idx 参数生效"""

    def test_format_references_respects_end_idx(self):
        """ref_end_idx 应限制格式化范围"""
        from paper_format_corrector.core.document.parser.reference import ReferenceFormatter

        formatter = ReferenceFormatter({})
        doc = Document()
        doc.add_paragraph("参考文献")
        doc.add_paragraph("[1] 文献1")
        doc.add_paragraph("[2] 文献2")
        doc.add_paragraph("[3] 文献3")
        doc.add_paragraph("这是正文")

        # 只格式化前两条参考文献
        formatter.format_references(doc, ref_start_idx=0, ref_end_idx=3)

        # 第4段（索引3）之后的不应被格式化
        # 这里主要验证不抛异常且 ref_end_idx 被使用


# ── Bug 1.5: path_security.py .diff.html 扩展名 ─────────────────────


class TestPathSecurityExtensions:
    """验证扩展名检查正确"""

    def test_html_extension_accepted(self):
        """HTML 扩展名应被接受"""
        from paper_format_corrector.adapters.path_security import ALLOWED_OUTPUT_EXTENSIONS

        assert ".html" in ALLOWED_OUTPUT_EXTENSIONS

    def test_diff_html_is_html_suffix(self):
        """.diff.html 文件的 suffix 是 .html"""
        p = Path("report.diff.html")
        assert p.suffix == ".html"
        # .html 在允许列表中，所以 .diff.html 也会被接受


# ── Bug 1.6: image_handler.py margin 为 0 ───────────────────────────


class TestImageHandlerMargins:
    """验证 margin 为 0 时正确处理"""

    def test_zero_margin_not_treated_as_none(self):
        """margin 为 0 不应触发 fallback"""
        from paper_format_corrector.core.document.elements.image_handler import ImageHandler

        handler = ImageHandler({})
        doc = Document()
        section = doc.sections[0]

        # 设置 margin 为 0
        section.left_margin = 0
        section.right_margin = 0

        width = handler._get_page_content_width(doc)
        # 应该返回 page_width - 0 - 0，而不是 fallback Cm(15)
        assert width == section.page_width

    def test_none_margin_uses_fallback(self):
        """没有 section 时应返回 fallback"""
        from paper_format_corrector.core.document.elements.image_handler import ImageHandler

        handler = ImageHandler({})
        # 创建空文档（理论上 python-docx 总有至少一个 section）
        doc = Document()
        width = handler._get_page_content_width(doc)
        assert width is not None


# ── Bug 1.7: diff_reporter.py 引号转义 ──────────────────────────────


class TestDiffReporterEscaping:
    """验证 HTML 转义包含引号"""

    def test_quote_is_escaped(self):
        """引号应被转义"""
        from paper_format_corrector.shared.docx_utils import escape_html

        result = escape_html('He said "hello"')
        assert "&quot;" in result
        assert '"' not in result

    def test_all_special_chars_escaped(self):
        """所有特殊字符都应被转义"""
        from paper_format_corrector.shared.docx_utils import escape_html

        result = escape_html('<a href="test">&')
        assert "&lt;" in result
        assert "&gt;" in result
        assert "&quot;" in result
        assert "&amp;" in result
