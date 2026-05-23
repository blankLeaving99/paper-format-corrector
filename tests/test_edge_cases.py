"""Edge case tests for the format corrector."""

import sys
import os

sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "src"))

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from paper_format_corrector.core.format_corrector import FormatCorrector


def _correct(doc, template_path, config, tmp_path, name="edge"):
    """Save a doc, correct it, return the report."""
    input_path = str(tmp_path / f"input_{name}.docx")
    output_path = str(tmp_path / f"output_{name}.docx")
    doc.save(input_path)

    corrector = FormatCorrector(template_path, config)
    corrector.load_template_styles()
    return corrector.correct_document(input_path, output_path), output_path


def test_empty_document(template_path, config, tmp_path):
    """A document with no paragraphs should not crash."""
    doc = Document()
    report, output_path = _correct(doc, template_path, config, tmp_path, "empty")

    assert report is not None
    assert report["paragraphs_corrected"] == 0
    assert Path(output_path).exists()


def test_blank_only_document(template_path, config, tmp_path):
    """A document with only blank paragraphs should not crash."""
    doc = Document()
    for _ in range(5):
        doc.add_paragraph("")
    report, output_path = _correct(doc, template_path, config, tmp_path, "blank")

    assert report is not None
    assert report["paragraphs_corrected"] == 0
    assert Path(output_path).exists()


def test_no_headings_all_body(template_path, config, tmp_path):
    """A document with no headings, only body text."""
    doc = Document()

    for i in range(5):
        p = doc.add_paragraph()
        run = p.add_run(f"这是第{i + 1}段正文内容，没有标题结构。")
        run.font.size = Pt(10)

    report, output_path = _correct(doc, template_path, config, tmp_path, "no_headings")

    assert report is not None
    assert report["paragraphs_corrected"] > 0
    assert report["headings_fixed"] == 0, "No headings to fix"
    assert report["body_fixed"] > 0


def test_single_paragraph(template_path, config, tmp_path):
    """A document with a single paragraph."""
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("这是一段单独的正文内容。")
    run.font.size = Pt(10)

    report, output_path = _correct(doc, template_path, config, tmp_path, "single")

    assert report is not None
    assert report["paragraphs_corrected"] == 1
    assert report["body_fixed"] == 1
    assert Path(output_path).exists()


def test_only_headings_no_body(template_path, config, tmp_path):
    """A document with only headings and no body text."""
    doc = Document()

    p = doc.add_paragraph()
    run = p.add_run("第一章 引言")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    run = p.add_run("1.1 背景")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("1.1.1 详细背景")
    run.font.size = Pt(12)

    report, output_path = _correct(doc, template_path, config, tmp_path, "headings_only")

    assert report is not None
    assert report["headings_fixed"] >= 3
    assert report["body_fixed"] == 0


def test_mixed_language_content(template_path, config, tmp_path):
    """A document with mixed Chinese and English content."""
    doc = Document()

    p = doc.add_paragraph()
    run = p.add_run("第一章 Introduction to Deep Learning")
    run.font.size = Pt(14)
    run.font.bold = False

    p = doc.add_paragraph()
    run = p.add_run("深度学习（Deep Learning）是机器学习的一个分支，近年来取得了重大突破。")
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run("Convolutional Neural Networks (CNN) are widely used in image recognition.")
    run.font.size = Pt(10)

    report, output_path = _correct(doc, template_path, config, tmp_path, "mixed_lang")

    assert report is not None
    assert report["paragraphs_corrected"] > 0


def test_reference_only_document(template_path, config, tmp_path):
    """A document that only contains references."""
    doc = Document()

    p = doc.add_paragraph()
    run = p.add_run("参考文献")
    run.font.size = Pt(14)
    run.font.bold = True

    refs = [
        "[1] Smith J. Deep Learning Review[J]. Nature, 2020, 584: 1-10.",
        "[2] 张三. 深度学习综述[J]. 计算机学报, 2021, 44(1): 1-20.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(10)

    report, output_path = _correct(doc, template_path, config, tmp_path, "refs_only")

    assert report is not None
    assert report["paragraphs_corrected"] > 0
    assert Path(output_path).exists()
