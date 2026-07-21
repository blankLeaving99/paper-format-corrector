"""共享工具模块测试

测试 utils/docx_utils.py 和 infra/external_tools.py 的功能。
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "src"))

from paper_format_corrector.infrastructure.external_tools import find_libreoffice
from paper_format_corrector.shared.docx_utils import (
    ALIGN_MAP,
    EMU_PER_CM,
    EMU_PER_INCH,
    escape_html,
    set_east_asian_font,
    set_paragraph_font,
    set_run_font,
)

# ── 常量测试 ────────────────────────────────────────────────────────


class TestConstants:
    """测试共享常量"""

    def test_emu_per_cm_value(self):
        assert EMU_PER_CM == 360000

    def test_emu_per_inch_value(self):
        assert EMU_PER_INCH == 914400

    def test_align_map_keys(self):
        assert "center" in ALIGN_MAP
        assert "left" in ALIGN_MAP
        assert "right" in ALIGN_MAP
        assert "justify" in ALIGN_MAP


# ── escape_html 测试 ────────────────────────────────────────────────


class TestEscapeHtml:
    """测试 HTML 转义函数"""

    def test_empty_string(self):
        assert escape_html("") == ""

    def test_none_input(self):
        assert escape_html(None) == ""

    def test_no_special_chars(self):
        assert escape_html("hello world") == "hello world"

    def test_ampersand(self):
        assert escape_html("a&b") == "a&amp;b"

    def test_less_than(self):
        assert escape_html("a<b") == "a&lt;b"

    def test_greater_than(self):
        assert escape_html("a>b") == "a&gt;b"

    def test_double_quote(self):
        assert escape_html('a"b') == "a&quot;b"

    def test_all_special_chars(self):
        result = escape_html('<script>alert("xss")&amp;</script>')
        assert "&lt;" in result
        assert "&gt;" in result
        assert "&quot;" in result
        assert "&amp;" in result

    def test_chinese_with_special_chars(self):
        assert escape_html("标题<h1>") == "标题&lt;h1&gt;"


# ── set_run_font 测试 ───────────────────────────────────────────────


class TestSetRunFont:
    """测试字体设置函数"""

    def test_set_run_font_basic(self, tmp_path):
        """测试基本字体设置"""
        from docx import Document

        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("test")

        set_run_font(run, "Arial", "黑体", 14, True)

        assert run.font.name == "Arial"
        assert run.font.size.pt == 14
        assert run.font.bold is True

    def test_set_run_font_defaults(self, tmp_path):
        """测试默认参数"""
        from docx import Document

        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("test")

        set_run_font(run)

        assert run.font.name == "Times New Roman"
        assert run.font.bold is False

    def test_set_run_font_no_size(self, tmp_path):
        """测试不设置字号"""
        from docx import Document

        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("test")

        set_run_font(run, size_pt=None)

        assert run.font.size is None


# ── set_paragraph_font 测试 ─────────────────────────────────────────


class TestSetParagraphFont:
    """测试段落字体批量设置"""

    def test_sets_all_runs(self, tmp_path):
        from docx import Document

        doc = Document()
        para = doc.add_paragraph()
        para.add_run("run1")
        para.add_run("run2")

        set_paragraph_font(para, "Arial", "黑体", 12, True)

        for run in para.runs:
            assert run.font.name == "Arial"
            assert run.font.size.pt == 12
            assert run.font.bold is True

    def test_empty_paragraph(self, tmp_path):
        from docx import Document

        doc = Document()
        para = doc.add_paragraph()

        # 不应抛出异常
        set_paragraph_font(para)


# ── set_east_asian_font 测试 ────────────────────────────────────────


class TestSetEastAsianFont:
    """测试东亚字体设置"""

    def test_sets_east_asian_font(self, tmp_path):
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("测试")

        set_east_asian_font(run, "黑体")

        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        assert rFonts is not None
        assert rFonts.get(qn("w:eastAsia")) == "黑体"

    def test_creates_rfonts_if_missing(self, tmp_path):
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("测试")

        # 确保 rFonts 不存在
        rpr = run._element.get_or_add_rPr()
        for child in list(rpr):
            if child.tag == qn("w:rFonts"):
                rpr.remove(child)

        set_east_asian_font(run, "宋体")

        rFonts = rpr.find(qn("w:rFonts"))
        assert rFonts is not None
        assert rFonts.get(qn("w:eastAsia")) == "宋体"


# ── find_libreoffice 测试 ───────────────────────────────────────────


class TestFindLibreoffice:
    """测试 LibreOffice 查找"""

    def test_returns_str_or_none(self):
        result = find_libreoffice()
        assert result is None or isinstance(result, str)

    def test_result_is_cached(self):
        """多次调用应返回相同结果"""
        result1 = find_libreoffice()
        result2 = find_libreoffice()
        assert result1 == result2
