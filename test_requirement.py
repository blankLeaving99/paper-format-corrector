"""测试需求文档解析器

生成样例需求文档（自然语言 + 表格两种格式），然后解析并验证。
"""

from pathlib import Path
from requirement_parser import RequirementParser


def create_natural_language_requirement(path="input/requirement_natural.txt"):
    """生成自然语言格式的需求文档"""
    content = """毕业论文格式要求

一、论文题目
论文题目用黑体，二号字，居中，加粗。

二、作者与单位
作者姓名用宋体，小四号，居中。
作者单位用宋体，五号，居中。

三、摘要
摘要标题用黑体，三号字，居中，加粗。
摘要正文用宋体，小四号，1.5倍行距，首行缩进2字符。

四、关键词
关键词用宋体，小四号，关键词三个字加粗。

五、一级标题
一级标题用黑体，二号字，居中，加粗。

六、二级标题
二级标题用黑体，三号字，左对齐，加粗。

七、三级标题
三级标题用黑体，四号字，左对齐，加粗。

八、正文
正文用宋体，小四号，1.5倍行距，首行缩进2字符，两端对齐。

九、参考文献
参考文献标题用黑体，三号字，居中，加粗。
参考文献正文用宋体，五号，1.25倍行距，悬挂缩进。

十、页边距
上边距2.54cm，下边距2.54cm，左边距3.17cm，右边距3.17cm。

十一、图标题
图标题用宋体，五号，居中。

十二、表标题
表标题用宋体，五号，居中。
"""
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    Path(path).write_text(content, encoding="utf-8")
    print(f"已生成自然语言需求文档: {path}")
    return path


def create_table_requirement(path="input/requirement_table.txt"):
    """生成表格格式的需求文档"""
    content = """毕业论文格式要求（表格版）

项目 | 字体 | 字号 | 对齐 | 加粗
论文题目 | 黑体 | 二号 | 居中 | 加粗
作者姓名 | 宋体 | 小四 | 居中 |
作者单位 | 宋体 | 五号 | 居中 |
摘要标题 | 黑体 | 三号 | 居中 | 加粗
摘要正文 | 宋体 | 小四 | |
一级标题 | 黑体 | 二号 | 居中 | 加粗
二级标题 | 黑体 | 三号 | 左对齐 | 加粗
三级标题 | 黑体 | 四号 | 左对齐 | 加粗
正文 | 宋体 | 小四 | 两端对齐 |
参考文献标题 | 黑体 | 三号 | 居中 | 加粗
参考文献正文 | 宋体 | 五号 | |
图标题 | 宋体 | 五号 | 居中 |
表标题 | 宋体 | 五号 | 居中 |

行距要求：
- 正文：1.5倍行距
- 摘要正文：1.5倍行距
- 参考文献正文：1.25倍行距

缩进要求：
- 正文：首行缩进2字符
- 摘要正文：首行缩进2字符

页边距：上下2.54cm，左右3.17cm
"""
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    Path(path).write_text(content, encoding="utf-8")
    print(f"已生成表格格式需求文档: {path}")
    return path


def create_mixed_requirement(path="input/requirement_mixed.docx"):
    """生成混合格式的 .docx 需求文档"""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    doc.add_heading("毕业论文格式要求", level=1)

    doc.add_paragraph("以下为论文各部分的格式规范：")

    # 表格
    table = doc.add_table(rows=9, cols=5)
    table.style = "Table Grid"

    headers = ["项目", "字体", "字号", "对齐", "加粗"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    data = [
        ["论文题目", "黑体", "二号", "居中", "加粗"],
        ["一级标题", "黑体", "二号", "居中", "加粗"],
        ["二级标题", "黑体", "三号", "左对齐", "加粗"],
        ["三级标题", "黑体", "四号", "左对齐", "加粗"],
        ["正文", "宋体", "小四", "两端对齐", ""],
        ["摘要正文", "宋体", "小四", "", ""],
        ["关键词", "宋体", "小四", "", "加粗"],
        ["参考文献正文", "宋体", "五号", "", ""],
    ]
    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            table.rows[r + 1].cells[c].text = val

    doc.add_paragraph("")
    doc.add_paragraph("正文行距：1.5倍行距")
    doc.add_paragraph("参考文献行距：1.25倍行距")
    doc.add_paragraph("正文首行缩进2字符")
    doc.add_paragraph("页边距：上2.54cm，下2.54cm，左3.17cm，右3.17cm")

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"已生成 .docx 需求文档: {path}")
    return path


def test_parse(path, label):
    """测试解析单个需求文档"""
    print(f"\n{'=' * 60}")
    print(f"测试: {label}")
    print(f"文件: {path}")
    print(f"{'=' * 60}")

    parser = RequirementParser()
    config = parser.parse(path)
    parser.print_parsed_rules()

    # 验证关键字段
    fr = config.get("format_rules", {})
    checks = []

    # 检查字体
    font = fr.get("font", {})
    if font.get("chinese"):
        checks.append(f"  中文字体: {font['chinese']}")

    # 检查标题
    for h in ("heading1", "heading2", "heading3"):
        h_cfg = fr.get("headings", {}).get(h, {})
        if h_cfg:
            checks.append(f"  {h}: size={h_cfg.get('font_size')}, bold={h_cfg.get('bold')}, align={h_cfg.get('align')}")

    # 检查正文
    body = fr.get("body_text", {})
    if body:
        checks.append(f"  正文: size={body.get('font_size')}, spacing={body.get('line_spacing')}, indent={body.get('first_line_indent')}")

    # 检查边距
    margins = fr.get("margins", {})
    if margins:
        checks.append(f"  边距: top={margins.get('top')}, bottom={margins.get('bottom')}, left={margins.get('left')}, right={margins.get('right')}")

    if checks:
        print("\n验证结果:")
        for c in checks:
            print(c)

    return config


def main():
    print("=" * 60)
    print("  需求文档解析器测试")
    print("=" * 60)

    # 创建样例需求文档
    p1 = create_natural_language_requirement()
    p2 = create_table_requirement()
    p3 = create_mixed_requirement()

    # 测试解析
    c1 = test_parse(p1, "自然语言格式 (.txt)")
    c2 = test_parse(p2, "表格格式 (.txt)")
    c3 = test_parse(p3, "混合格式 (.docx)")

    # 对比三种方式解析出的配置是否一致
    print(f"\n{'=' * 60}")
    print("对比三种解析结果:")
    print(f"{'=' * 60}")

    fr1 = c1.get("format_rules", {})
    fr2 = c2.get("format_rules", {})
    fr3 = c3.get("format_rules", {})

    # 正文字号
    print(f"  正文字号: 自然语言={fr1.get('body_text', {}).get('font_size')}, "
          f"表格={fr2.get('body_text', {}).get('font_size')}, "
          f"docx={fr3.get('body_text', {}).get('font_size')}")

    # 一级标题字号
    print(f"  一级标题字号: 自然语言={fr1.get('headings', {}).get('heading1', {}).get('font_size')}, "
          f"表格={fr2.get('headings', {}).get('heading1', {}).get('font_size')}, "
          f"docx={fr3.get('headings', {}).get('heading1', {}).get('font_size')}")

    print("\n测试完成！")


if __name__ == "__main__":
    main()
