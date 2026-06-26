"""页眉页脚处理器（增强版）

支持：
- 动态章名页眉
- 首页不同
- 奇偶页不同
- 页码分段（前置罗马、正文阿拉伯）
- 页眉下划线
"""

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


class HeaderFooterHandler:
    """页眉页脚处理器"""

    def __init__(self, config):
        hf_config = config.get("format_rules", {}).get("header_footer", {})
        self.enabled = hf_config.get("enabled", False)
        self.header_config = hf_config.get("header", {})
        self.footer_config = hf_config.get("footer", {})
        self.different_first_page = hf_config.get("different_first_page", False)
        self.different_odd_even = hf_config.get("different_odd_even", False)

        pn_config = config.get("format_rules", {}).get("page_numbering", {})
        self.pn_enabled = pn_config.get("enabled", False)
        self.front_matter_style = pn_config.get("front_matter", {}).get("style", "roman_lower")
        self.body_style = pn_config.get("body", {}).get("style", "arabic")
        self.body_start = pn_config.get("body", {}).get("start", 1)

    def apply(self, doc, chapter_map=None):
        """应用页眉页脚设置

        Args:
            doc: Document 对象
            chapter_map: 章节映射 {section_index: chapter_name}，用于动态章名页眉
        """
        if not self.enabled and not self.pn_enabled:
            return

        for idx, section in enumerate(doc.sections):
            section.different_first_page_header_footer = self.different_first_page

            # 奇偶页不同
            if self.different_odd_even:
                self._set_odd_even_pages(section)

            if self.enabled:
                chapter_name = chapter_map.get(idx) if chapter_map else None
                self._setup_header(section, chapter_name)
                self._setup_footer(section)

            if self.pn_enabled:
                is_front_matter = idx == 0  # 第一节视为前置部分
                self._setup_page_numbers(section, is_front_matter)

    def _set_odd_even_pages(self, section):
        """启用奇偶页不同页眉页脚"""
        sectPr = section._sectPr
        # 设置 evenAndOddHeaders 属性
        even_odd = sectPr.find(qn("w:evenAndOddHeaders"))
        if even_odd is None:
            even_odd = OxmlElement("w:evenAndOddHeaders")
            sectPr.append(even_odd)

    def _resolve_header_text(self, text, chapter_name=None):
        """替换页眉文本中的占位符

        Args:
            text: 原始页眉文本
            chapter_name: 当前章名称

        Returns:
            str: 替换后的文本
        """
        if not text:
            return text
        if chapter_name:
            text = text.replace("{chapter}", chapter_name)
        else:
            text = text.replace("{chapter}", "")
        return text

    def _setup_header(self, section, chapter_name=None):
        header = section.header
        header_text = self.header_config.get("text", "")
        header_text = self._resolve_header_text(header_text, chapter_name)

        if not header_text and not self.header_config.get("bottom_border", False):
            return

        if header.paragraphs:
            para = header.paragraphs[0]
        else:
            para = header.add_paragraph()

        if header_text:
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = header_text
            else:
                para.add_run(header_text)

            for run in para.runs:
                run.font.name = self.header_config.get("font_name", "宋体")
                run.font.size = Pt(self.header_config.get("font_size", 10.5))

            align = self.header_config.get("align", "center")
            align_map = {
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }
            para.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.CENTER)

        if self.header_config.get("bottom_border", False):
            self._add_bottom_border(para)

    def _setup_footer(self, section):
        footer = section.footer
        if footer.paragraphs:
            para = footer.paragraphs[0]
        else:
            para = footer.add_paragraph()

        align = self.footer_config.get("align", "center")
        align_map = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        para.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.CENTER)

    def _setup_page_numbers(self, section, is_front_matter=False):
        """设置页码

        Args:
            section: Section 对象
            is_front_matter: 是否为前置部分（使用罗马数字）
        """
        footer = section.footer
        if footer.paragraphs:
            para = footer.paragraphs[0]
        else:
            para = footer.add_paragraph()

        for run in para.runs:
            run.text = ""

        self._add_page_number_field(para)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in para.runs:
            run.font.size = Pt(self.footer_config.get("font_size", 10.5))

        # 设置页码格式
        if is_front_matter:
            self._set_section_page_number_type(section, "roman_lower")
        else:
            self._set_section_page_number_type(section, self.body_style, self.body_start)

    def _set_section_page_number_type(self, section, style, start=None):
        """设置节的页码格式

        Args:
            section: Section 对象
            style: 页码格式 (roman_lower, arabic, etc.)
            start: 起始页码（仅对正文部分有效）
        """
        sectPr = section._sectPr

        # 移除已有的 pgNumType
        existing = sectPr.find(qn("w:pgNumType"))
        if existing is not None:
            sectPr.remove(existing)

        pgNumType = OxmlElement("w:pgNumType")

        # 设置格式
        fmt_map = {
            "roman_lower": "lowerRoman",
            "roman_upper": "upperRoman",
            "arabic": "decimal",
            "alpha_lower": "lowerLetter",
            "alpha_upper": "upperLetter",
        }
        fmt = fmt_map.get(style, "decimal")
        pgNumType.set(qn("w:fmt"), fmt)

        # 设置起始页码
        if start is not None:
            pgNumType.set(qn("w:start"), str(start))

        sectPr.append(pgNumType)

    def _add_page_number_field(self, paragraph):
        run = paragraph.add_run()

        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")
        run._element.append(fldChar_begin)

        run2 = paragraph.add_run()
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = " PAGE "
        run2._element.append(instrText)

        run3 = paragraph.add_run()
        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")
        run3._element.append(fldChar_end)

    def _add_bottom_border(self, paragraph):
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is None:
            pBdr = OxmlElement("w:pBdr")
            pPr.append(pBdr)

        bottom = pBdr.find(qn("w:bottom"))
        if bottom is None:
            bottom = OxmlElement("w:bottom")
            pBdr.append(bottom)

        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")

    def remove_existing_headers_footers(self, doc):
        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                for run in para.runs:
                    run.text = ""
            footer = section.footer
            for para in footer.paragraphs:
                for run in para.runs:
                    run.text = ""
