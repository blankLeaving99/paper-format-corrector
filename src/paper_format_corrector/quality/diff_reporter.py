"""矫正前后对比报告模块

对比原始文档和矫正后文档的差异，生成HTML格式的报告。

功能：
- 段落级别diff
- 格式变更标注（字体、字号、对齐、行距）
- 图表编号变更
- 参考文献变更
- 统计汇总
"""

import difflib
from docx import Document
from docx.shared import Pt


class DiffReporter:
    """矫正前后对比报告器"""

    def __init__(self):
        self.changes = []

    def compare(self, original_path, corrected_path):
        """对比两个文档"""
        orig_doc = Document(original_path)
        corr_doc = Document(corrected_path)

        self.changes = []

        orig_texts = [p.text.strip() for p in orig_doc.paragraphs]
        corr_texts = [p.text.strip() for p in corr_doc.paragraphs]

        # 文本级别的diff
        text_changes = self._diff_texts(orig_texts, corr_texts)

        # 格式级别的diff
        format_changes = self._diff_formats(orig_doc, corr_doc)

        return {
            "text_changes": text_changes,
            "format_changes": format_changes,
            "orig_paragraphs": len(orig_texts),
            "corr_paragraphs": len(corr_texts),
        }

    def _diff_texts(self, orig, corr):
        """文本内容diff"""
        changes = []
        matcher = difflib.SequenceMatcher(None, orig, corr)

        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == "equal":
                continue
            elif tag == "replace":
                for i in range(min(i2 - i1, j2 - j1)):
                    changes.append({
                        "type": "modify",
                        "line": j1 + i + 1,
                        "original": orig[i1 + i] if i1 + i < len(orig) else "",
                        "corrected": corr[j1 + i] if j1 + i < len(corr) else "",
                    })
            elif tag == "insert":
                for j in range(j1, j2):
                    changes.append({
                        "type": "insert",
                        "line": j + 1,
                        "corrected": corr[j],
                    })
            elif tag == "delete":
                for i in range(i1, i2):
                    changes.append({
                        "type": "delete",
                        "line": i + 1,
                        "original": orig[i],
                    })

        return changes

    def _diff_formats(self, orig_doc, corr_doc):
        """格式差异对比"""
        changes = []
        orig_paras = orig_doc.paragraphs
        corr_paras = corr_doc.paragraphs

        for i in range(min(len(orig_paras), len(corr_paras))):
            orig_p = orig_paras[i]
            corr_p = corr_paras[i]

            if not orig_p.text.strip():
                continue

            # 对齐方式
            if orig_p.alignment != corr_p.alignment:
                changes.append({
                    "line": i + 1,
                    "type": "alignment",
                    "text": orig_p.text[:30],
                    "from": str(orig_p.alignment),
                    "to": str(corr_p.alignment),
                })

            # 行距
            orig_ls = orig_p.paragraph_format.line_spacing
            corr_ls = corr_p.paragraph_format.line_spacing
            if orig_ls != corr_ls and orig_ls is not None and corr_ls is not None:
                changes.append({
                    "line": i + 1,
                    "type": "line_spacing",
                    "text": orig_p.text[:30],
                    "from": str(orig_ls),
                    "to": str(corr_ls),
                })

            # 字体/字号 (检查第一个run)
            if orig_p.runs and corr_p.runs:
                orig_run = orig_p.runs[0]
                corr_run = corr_p.runs[0]

                if orig_run.font.size != corr_run.font.size:
                    if orig_run.font.size and corr_run.font.size:
                        changes.append({
                            "line": i + 1,
                            "type": "font_size",
                            "text": orig_p.text[:30],
                            "from": f"{orig_run.font.size.pt}pt",
                            "to": f"{corr_run.font.size.pt}pt",
                        })

                if orig_run.font.bold != corr_run.font.bold:
                    changes.append({
                        "line": i + 1,
                        "type": "bold",
                        "text": orig_p.text[:30],
                        "from": str(orig_run.font.bold),
                        "to": str(corr_run.font.bold),
                    })

        return changes

    def generate_html_report(self, original_path, corrected_path, output_path):
        """生成HTML对比报告"""
        result = self.compare(original_path, corrected_path)

        html = self._build_html(result)

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html)

        return output_path

    def _build_html(self, result):
        text_changes = result["text_changes"]
        format_changes = result["format_changes"]

        rows = []
        for change in text_changes[:200]:
            if change["type"] == "modify":
                diff = self._inline_diff(change["original"], change["corrected"])
                rows.append(f'<tr class="modify"><td>{change["line"]}</td><td>内容修改</td><td>{diff}</td></tr>')
            elif change["type"] == "insert":
                rows.append(f'<tr class="insert"><td>{change["line"]}</td><td>新增</td><td>{self._esc(change["corrected"])}</td></tr>')
            elif change["type"] == "delete":
                rows.append(f'<tr class="delete"><td>{change["line"]}</td><td>删除</td><td><del>{self._esc(change["original"])}</del></td></tr>')

        for change in format_changes[:100]:
            rows.append(
                f'<tr class="format"><td>{change["line"]}</td>'
                f'<td>{change["type"]}</td>'
                f'<td>{self._esc(change["text"])}: {change["from"]} → {change["to"]}</td></tr>'
            )

        table_rows = "\n".join(rows)

        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>格式矫正对比报告</title>
<style>
body {{ font-family: "Microsoft YaHei", sans-serif; margin: 20px; }}
h1 {{ color: #333; border-bottom: 2px solid #007bff; padding-bottom: 10px; }}
.stats {{ background: #f8f9fa; padding: 15px; border-radius: 5px; margin: 20px 0; }}
.stats span {{ margin-right: 30px; }}
table {{ width: 100%; border-collapse: collapse; margin-top: 20px; }}
th {{ background: #007bff; color: white; padding: 10px; text-align: left; }}
td {{ padding: 8px; border-bottom: 1px solid #ddd; }}
tr.modify td {{ background: #fff3cd; }}
tr.insert td {{ background: #d4edda; }}
tr.delete td {{ background: #f8d7da; }}
tr.format td {{ background: #d1ecf1; }}
del {{ color: #dc3545; }}
ins {{ color: #28a745; }}
</style>
</head>
<body>
<h1>格式矫正对比报告</h1>
<div class="stats">
  <span>原始文档: {result["orig_paragraphs"]} 段</span>
  <span>矫正文档: {result["corr_paragraphs"]} 段</span>
  <span>内容变更: {len(text_changes)} 处</span>
  <span>格式变更: {len(format_changes)} 处</span>
</div>
<table>
<tr><th>行号</th><th>变更类型</th><th>详情</th></tr>
{table_rows}
</table>
</body>
</html>"""

    def _inline_diff(self, old, new):
        """生成行内diff HTML"""
        old_escaped = self._esc(old)
        new_escaped = self._esc(new)
        return f'<del>{old_escaped}</del> → <ins>{new_escaped}</ins>'

    def _esc(self, text):
        if not text:
            return ""
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
