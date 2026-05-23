"""格式质量评分模块

对矫正后的文档进行全面质量检查，输出0-100分评分和详细报告。

评分维度：
- 标题格式 (20分)
- 正文格式 (20分)
- 图表编号 (15分)
- 参考文献 (15分)
- 页面设置 (10分)
- 摘要关键词 (10分)
- 页眉页脚 (5分)
- 目录 (5分)
"""

import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


class QualityScorer:
    """文档格式质量评分器"""

    def __init__(self, config):
        self.config = config
        self.issues = []
        self.scores = {}
        # 从 config 中读取标题检测模式
        detect = config.get("auto_detect", {})
        self._heading_patterns = [
            re.compile(detect.get("chapter_pattern", r"^第[一二三四五六七八九十百零\d]+[章部分篇]")),
            re.compile(detect.get("section_pattern", r"^\d+\.\d+")),
            re.compile(detect.get("subsection_pattern", r"^\d+\.\d+\.\d+")),
        ]

    def score(self, doc_path):
        """对文档评分，返回 (total_score, details)"""
        doc = Document(doc_path)
        self.issues = []
        self.scores = {}

        self.scores["headings"] = self._score_headings(doc)
        self.scores["body"] = self._score_body(doc)
        self.scores["figures_tables"] = self._score_figures_tables(doc)
        self.scores["references"] = self._score_references(doc)
        self.scores["page_setup"] = self._score_page_setup(doc)
        self.scores["abstract_keywords"] = self._score_abstract_keywords(doc)
        self.scores["header_footer"] = self._score_header_footer(doc)
        self.scores["toc"] = self._score_toc(doc)

        total = sum(self.scores.values())
        return total, self.scores, self.issues

    def _score_headings(self, doc):
        """标题格式评分 (满分20)"""
        score = 20
        rules = self.config.get("format_rules", {}).get("headings", {})
        body_size = self.config.get("format_rules", {}).get("body_text", {}).get("font_size", 12)

        heading_count = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            is_heading = any(p.match(text) for p in self._heading_patterns)

            if is_heading:
                heading_count += 1
                for run in para.runs:
                    if run.font.size and run.font.size.pt < body_size:
                        self.issues.append(f"标题 '{text[:20]}...' 字号小于正文")
                        score -= 1
                        break

        if heading_count == 0:
            self.issues.append("未检测到任何标题")
            score -= 5

        return max(0, score)

    def _score_body(self, doc):
        """正文格式评分 (满分20)"""
        score = 20
        body_config = self.config.get("format_rules", {}).get("body_text", {})
        expected_size = body_config.get("font_size", 12)
        expected_spacing = body_config.get("line_spacing", 1.5)

        body_paras = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if len(text) > 30 and not any(p.match(text) for p in self._heading_patterns):
                body_paras.append(para)

        if not body_paras:
            return score

        # 检查字号一致性
        size_issues = 0
        for para in body_paras[:20]:
            for run in para.runs:
                if run.font.size and abs(run.font.size.pt - expected_size) > 1:
                    size_issues += 1
                    break

        if size_issues > len(body_paras[:20]) * 0.3:
            self.issues.append(f"正文字号不一致，{size_issues}/{min(20, len(body_paras))} 段不符合要求")
            score -= 5

        # 检查首行缩进
        indent_config = body_config.get("first_line_indent", 0)
        if indent_config:
            no_indent = 0
            for para in body_paras[:20]:
                if para.paragraph_format.first_line_indent is None or para.paragraph_format.first_line_indent == 0:
                    no_indent += 1
            if no_indent > len(body_paras[:20]) * 0.3:
                self.issues.append(f"部分正文缺少首行缩进")
                score -= 3

        # 检查两端对齐
        align = body_config.get("align", "justify")
        if align == "justify":
            not_justified = 0
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            for para in body_paras[:20]:
                if para.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY and para.alignment is not None:
                    not_justified += 1
            if not_justified > len(body_paras[:20]) * 0.5:
                self.issues.append("部分正文未两端对齐")
                score -= 2

        return max(0, score)

    def _score_figures_tables(self, doc):
        """图表编号评分 (满分15)"""
        score = 15
        fig_pattern = re.compile(r"^图\s*(\d+[\-\.]\d+|\d+)")
        tab_pattern = re.compile(r"^表\s*(\d+[\-\.]\d+|\d+)")

        fig_nums = []
        tab_nums = []

        for para in doc.paragraphs:
            text = para.text.strip()
            m = fig_pattern.match(text)
            if m:
                fig_nums.append(m.group(1))
            m = tab_pattern.match(text)
            if m:
                tab_nums.append(m.group(1))

        # 检查编号连续性
        for name, nums in [("图", fig_nums), ("表", tab_nums)]:
            if len(nums) < 2:
                continue
            for i in range(1, len(nums)):
                if nums[i] == nums[i-1]:
                    self.issues.append(f"{name}编号重复: {name}{nums[i]}")
                    score -= 2

        return max(0, score)

    def _score_references(self, doc):
        """参考文献评分 (满分15)"""
        score = 15
        ref_start = None
        ref_keywords = self.config.get("auto_detect", {}).get("reference_keywords", ["参考文献"])

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text in ref_keywords or any(text.startswith(kw) for kw in ref_keywords):
                ref_start = i
                break

        if ref_start is None:
            self.issues.append("未找到参考文献章节")
            return 5

        ref_count = 0
        has_type_tag = 0
        has_numbering = 0

        for i in range(ref_start + 1, len(doc.paragraphs)):
            text = doc.paragraphs[i].text.strip()
            if not text:
                continue
            # 遇到新章节标题时停止
            if any(p.match(text) for p in self._heading_patterns):
                break
            if re.match(r"^\d+[\.\)]", text) or re.match(r"^\[\d+\]", text):
                ref_count += 1
                if re.search(r"\[[A-Z]", text):
                    has_type_tag += 1
                if re.match(r"^\[\d+\]", text):
                    has_numbering += 1

        if ref_count == 0:
            self.issues.append("参考文献列表为空")
            score -= 10
        else:
            if has_type_tag < ref_count * 0.8:
                self.issues.append(f"部分参考文献缺少文献类型标识 ({has_type_tag}/{ref_count})")
                score -= 3
            if has_numbering < ref_count * 0.8:
                self.issues.append(f"部分参考文献编号格式不规范")
                score -= 2

        return max(0, score)

    def _score_page_setup(self, doc):
        """页面设置评分 (满分10)"""
        score = 10
        margins = self.config.get("format_rules", {}).get("margins", {})
        if not margins:
            return score

        for section in doc.sections:
            for direction, expected in [("top", "top"), ("bottom", "bottom"), ("left", "left"), ("right", "right")]:
                actual = getattr(section, f"{direction}_margin")
                if actual is not None and actual > 0:
                    actual_cm = actual / 360000
                    expected_cm = margins.get(expected, 2.54)
                    if abs(actual_cm - expected_cm) > 0.3:
                        self.issues.append(f"页面{direction}边距偏差: {actual_cm:.1f}cm (应为{expected_cm}cm)")
                        score -= 1

        return max(0, score)

    def _score_abstract_keywords(self, doc):
        """摘要关键词评分 (满分10)"""
        score = 10
        has_abstract = False
        has_keywords = False

        for para in doc.paragraphs:
            text = para.text.strip()
            if re.match(r"^摘\s*要$", text) or re.match(r"^Abstract$|^ABSTRACT$", text):
                has_abstract = True
            if re.match(r"^关键词[:：]|^Key\s*[Ww]ords", text):
                has_keywords = True

        if not has_abstract:
            self.issues.append("未检测到摘要")
            score -= 5
        if not has_keywords:
            self.issues.append("未检测到关键词")
            score -= 5

        return max(0, score)

    def _score_header_footer(self, doc):
        """页眉页脚评分 (满分5)"""
        score = 5
        has_header = False
        has_footer = False

        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                if para.text.strip():
                    has_header = True
                    break
            footer = section.footer
            for para in footer.paragraphs:
                if para.text.strip() or para._element.findall(qn("w:fldChar")):
                    has_footer = True
                    break

        if not has_header:
            self.issues.append("页眉为空")
            score -= 2
        if not has_footer:
            self.issues.append("页脚无页码")
            score -= 3

        return max(0, score)

    def _score_toc(self, doc):
        """目录评分 (满分5)"""
        score = 5
        has_toc = False

        for para in doc.paragraphs:
            text = para.text.strip()
            if text in ("目录", "目 录", "目  录", "Table of Contents"):
                has_toc = True
                break

        for field in doc.element.findall(".//" + qn("w:instrText")):
            if field.text and "TOC" in field.text:
                has_toc = True
                break

        if not has_toc:
            self.issues.append("未检测到目录")
            score -= 3

        return max(0, score)

    def format_report(self, total, scores, issues):
        """格式化评分报告"""
        lines = []
        lines.append("=" * 50)
        lines.append("  文档格式质量评分报告")
        lines.append("=" * 50)
        lines.append(f"\n  总分: {total}/100")

        grade = "A+" if total >= 95 else "A" if total >= 90 else "B+" if total >= 85 else "B" if total >= 80 else "C" if total >= 70 else "D"
        lines.append(f"  等级: {grade}")

        lines.append("\n  各项得分:")
        labels = {
            "headings": ("标题格式", 20),
            "body": ("正文格式", 20),
            "figures_tables": ("图表编号", 15),
            "references": ("参考文献", 15),
            "page_setup": ("页面设置", 10),
            "abstract_keywords": ("摘要关键词", 10),
            "header_footer": ("页眉页脚", 5),
            "toc": ("目录", 5),
        }
        for key, (label, max_score) in labels.items():
            s = scores.get(key, 0)
            filled = int(s / max_score * 20)
            bar = "#" * filled + "-" * (20 - filled)
            lines.append(f"    {label:　<6} {bar} {s}/{max_score}")

        if issues:
            lines.append(f"\n  发现的问题 ({len(issues)}):")
            for issue in issues:
                lines.append(f"    - {issue}")

        lines.append("=" * 50)
        return "\n".join(lines)
