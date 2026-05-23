"""需求文档解析器

从 .docx / .txt / .md 格式的需求文档中提取论文格式要求，
生成可直接用于 FormatCorrector 的配置字典。

支持两种规则写法：
1. 自然语言：如 "一级标题用黑体，二号字，居中，加粗"
2. 表格/列表：如 "一级标题 | 黑体 | 二号 | 居中 | 加粗"
"""

import re
from pathlib import Path


# ========== 中文字号 -> pt 映射 ==========
FONT_SIZE_MAP = {
    "初号": 42, "小初": 36,
    "一号": 26, "小一": 24,
    "二号": 22, "小二": 18,
    "三号": 16, "小三": 15,
    "四号": 14, "小四": 12,
    "五号": 10.5, "小五": 9,
    "六号": 7.5, "小六": 6.5,
    "七号": 5.5, "八号": 5,
}

# 字体名标准化
FONT_NAME_MAP = {
    "宋体": "宋体", "simsun": "宋体", "sim sun": "宋体",
    "黑体": "黑体", "simhei": "黑体", "sim hei": "黑体",
    "楷体": "楷体", "楷体_gb2312": "楷体", "simkai": "楷体",
    "仿宋": "仿宋", "仿宋_gb2312": "仿宋", "fangsong": "仿宋",
    "微软雅黑": "微软雅黑", "microsoft yahei": "微软雅黑",
    "times new roman": "Times New Roman", "times": "Times New Roman",
    "tnr": "Times New Roman",
    "arial": "Arial",
    "calibri": "Calibri",
    "courier new": "Courier New",
}

# 段落类型关键词 -> config key 映射
SECTION_TYPE_MAP = {
    "论文题目": "title_page", "题目": "title_page", "题名": "title_page",
    "标题": "title_page", "论文标题": "title_page",
    "一级标题": "heading1", "章标题": "heading1",
    "二级标题": "heading2", "节标题": "heading2",
    "三级标题": "heading3", "小节标题": "heading3",
    "四级标题": "heading4",
    "正文": "body_text", "主体": "body_text",
    "摘要": "abstract", "中文摘要": "abstract",
    "英文摘要": "abstract_en",
    "关键词": "keywords",
    "参考文献": "references", "参考文献标题": "references",
    "页眉": "header", "页脚": "footer",
    "目录": "toc",
    "作者": "author", "作者姓名": "author",
    "作者单位": "affiliation", "单位": "affiliation",
    "图标题": "figure_caption", "图注": "figure_caption",
    "表标题": "table_caption", "表注": "table_caption",
    "公式": "formula",
    "脚注": "footnote",
    "附录": "appendix", "致谢": "acknowledgment",
}


class RequirementParser:
    """需求文档解析器"""

    def __init__(self):
        self.rules = {}      # 解析出的格式规则
        self.raw_lines = []  # 原始文本行
        self.warnings = []   # 解析警告

    def parse(self, file_path):
        """解析需求文档，返回配置字典"""
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"需求文档不存在: {file_path}")

        ext = file_path.suffix.lower()
        if ext == ".docx":
            self.raw_lines = self._read_docx(file_path)
        elif ext in (".txt", ".md", ".markdown"):
            self.raw_lines = self._read_text(file_path)
        elif ext == ".pdf":
            self.raw_lines = self._read_pdf(file_path)
        elif ext in (".doc", ".odt", ".rtf"):
            # 先转换为 docx，再读取
            from ..core.file_converter import FileConverter
            converter = FileConverter()
            import tempfile
            import shutil
            tmp_dir = tempfile.mkdtemp()
            try:
                converted_path = converter.convert(str(file_path), tmp_dir)
                self.raw_lines = self._read_docx(Path(converted_path))
            except Exception as e:
                raise ValueError(f"无法转换 {ext} 格式的需求文档: {e}")
            finally:
                shutil.rmtree(tmp_dir, ignore_errors=True)
        else:
            raise ValueError(f"不支持的文件格式: {ext}，支持 .docx/.doc/.odt/.rtf/.pdf/.txt/.md")

        # 解析每一行
        for line in self.raw_lines:
            line = line.strip()
            if not line:
                continue
            self._parse_line(line)

        # 合成配置
        config = self._build_config()

        if self.warnings:
            config["_warnings"] = self.warnings

        return config

    def get_warnings(self):
        return self.warnings

    # ========== 文档读取 ==========

    def _read_docx(self, path):
        from docx import Document
        doc = Document(str(path))
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
        # 也读取表格内容
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    lines.append("\t".join(cells))
        return lines

    def _read_text(self, path):
        # 尝试常见编码
        for enc in ("utf-8-sig", "utf-8", "gbk", "gb18030", "big5", "latin-1"):
            try:
                text = path.read_text(encoding=enc)
                return text.splitlines()
            except (UnicodeDecodeError, UnicodeError):
                continue
        # 最终回退
        text = path.read_text(encoding="utf-8", errors="replace")
        return text.splitlines()

    def _read_pdf(self, path):
        """读取 PDF 文件内容"""
        # 尝试 pdfplumber
        try:
            import pdfplumber
            lines = []
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        lines.extend(text.splitlines())
            return lines
        except ImportError:
            pass

        # 尝试 PyMuPDF
        try:
            import fitz
            doc = fitz.open(str(path))
            lines = []
            for page in doc:
                text = page.get_text()
                if text:
                    lines.extend(text.splitlines())
            doc.close()
            return lines
        except ImportError:
            pass

        # 尝试 PyPDF2
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(str(path))
            lines = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    lines.extend(text.splitlines())
            return lines
        except ImportError:
            pass

        raise RuntimeError(
            "读取 PDF 需要安装以下库之一：pip install pdfplumber / PyMuPDF / PyPDF2"
        )

    # ========== 行级解析 ==========

    def _parse_line(self, line):
        """解析单行规则"""
        # 优先尝试表格格式（用 | 或 \t 分隔）
        if self._try_parse_table_row(line):
            return

        # 尝试列表格式（以 - 或 * 或 数字. 开头）
        if self._try_parse_list_item(line):
            return

        # 自然语言解析
        self._try_parse_natural_language(line)

    def _try_parse_table_row(self, line):
        """尝试解析表格行：项目 | 字体 | 字号 | 对齐 | 加粗"""
        # 用 | 或 \t 分隔
        if "|" in line:
            cells = [c.strip() for c in line.split("|") if c.strip()]
        elif "\t" in line:
            cells = [c.strip() for c in line.split("\t") if c.strip()]
        else:
            return False

        if len(cells) < 2:
            return False

        # 第一列应该是段落类型
        section_type = self._match_section_type(cells[0])
        if not section_type:
            return False

        # 解析后续列的属性
        attrs = {}
        for cell in cells[1:]:
            self._extract_attributes(cell, attrs)

        if attrs:
            self._store_rule(section_type, attrs, source="table")
            return True

        return False

    def _try_parse_list_item(self, line):
        """尝试解析列表项：- 一级标题：黑体，二号，居中"""
        m = re.match(r"^[\-\*•·]\s*(.+)", line)
        if not m:
            m = re.match(r"^\d+[\.\)、]\s*(.+)", line)
        if not m:
            return False

        content = m.group(1)

        # 尝试用冒号/逗号分隔 "类型：属性"
        for sep in ["：", ":", "——", "—"]:
            if sep in content:
                parts = content.split(sep, 1)
                section_type = self._match_section_type(parts[0].strip())
                if section_type:
                    attrs = {}
                    self._extract_attributes(parts[1], attrs)
                    if attrs:
                        self._store_rule(section_type, attrs, source="list")
                        return True

        return False

    def _try_parse_natural_language(self, line):
        """自然语言解析：找出段落类型，然后提取属性"""
        # 尝试匹配段落类型关键词
        section_type = None
        type_match = None

        for keyword, stype in SECTION_TYPE_MAP.items():
            idx = line.find(keyword)
            if idx >= 0:
                if section_type is None or idx < type_match.start():
                    section_type = stype
                    type_match = re.search(re.escape(keyword), line)

        if not section_type:
            return False

        # 提取该行中的所有属性
        attrs = {}
        self._extract_attributes(line, attrs)

        if attrs:
            self._store_rule(section_type, attrs, source="nl")

    # ========== 属性提取 ==========

    def _extract_attributes(self, text, attrs):
        """从文本中提取格式属性"""
        self._extract_font_name(text, attrs)
        self._extract_font_size(text, attrs)
        self._extract_alignment(text, attrs)
        self._extract_bold_italic(text, attrs)
        self._extract_line_spacing(text, attrs)
        self._extract_indent(text, attrs)
        self._extract_margin(text, attrs)
        self._extract_special_features(text, attrs)

    def _extract_font_name(self, text, attrs):
        """提取字体名"""
        text_lower = text.lower()
        for alias, canonical in FONT_NAME_MAP.items():
            if alias in text_lower:
                # 判断是中文字体还是英文字体
                if canonical in ("宋体", "黑体", "楷体", "仿宋", "微软雅黑"):
                    attrs["chinese_font"] = canonical
                else:
                    attrs["english_font"] = canonical
                return

    def _extract_font_size(self, text, attrs):
        """提取字号"""
        # 先匹配中文字号（优先最长匹配，同长时优先最靠前的）
        best_name = None
        best_pt = None
        best_len = 0
        best_idx = len(text)
        for name, pt in FONT_SIZE_MAP.items():
            idx = text.find(name)
            if idx >= 0:
                if len(name) > best_len or (len(name) == best_len and idx < best_idx):
                    best_name = name
                    best_pt = pt
                    best_len = len(name)
                    best_idx = idx
        if best_name:
            attrs["font_size"] = best_pt
            attrs["font_size_label"] = best_name
            return

        # 匹配数字+pt/磅/号
        m = re.search(r"(\d+(?:\.\d+)?)\s*(?:pt|磅|号)", text)
        if m:
            attrs["font_size"] = float(m.group(1))
            return

        # 匹配 "X号字"
        m = re.search(r"(\d+(?:\.\d+)?)\s*号\s*字?", text)
        if m:
            attrs["font_size"] = float(m.group(1))

    def _extract_alignment(self, text, attrs):
        """提取对齐方式"""
        if "居中" in text:
            attrs["align"] = "center"
        elif "右对齐" in text or "右端对齐" in text:
            attrs["align"] = "right"
        elif "两端对齐" in text:
            attrs["align"] = "justify"
        elif "左对齐" in text:
            attrs["align"] = "left"

    def _extract_bold_italic(self, text, attrs):
        """提取加粗/斜体"""
        if "加粗" in text or "粗体" in text or "黑体" in text:
            # 黑体通常暗示加粗
            attrs["bold"] = True
        if "不加粗" in text or "不粗" in text or "普通" in text:
            attrs["bold"] = False
        if "斜体" in text:
            attrs["italic"] = True

    def _extract_line_spacing(self, text, attrs):
        """提取行距"""
        # "单倍行距"
        if "单倍行距" in text:
            attrs["line_spacing"] = 1.0
            attrs["line_spacing_type"] = "multiple"
            return

        # "1.5倍行距" "1.5倍"
        m = re.search(r"(\d+(?:\.\d+)?)\s*倍\s*(?:行距)?", text)
        if m:
            attrs["line_spacing"] = float(m.group(1))
            attrs["line_spacing_type"] = "multiple"
            return

        # "2倍行距"
        if "2倍行距" in text:
            attrs["line_spacing"] = 2.0
            attrs["line_spacing_type"] = "multiple"
            return

        # "行距20磅" "固定值20磅" "行距固定值20磅"
        m = re.search(r"(?:行距|固定值?)\s*(\d+(?:\.\d+)?)\s*(?:磅|pt)", text)
        if m:
            attrs["line_spacing"] = float(m.group(1))
            attrs["line_spacing_type"] = "exact"
            return

        # "最小值20磅"
        m = re.search(r"最小值\s*(\d+(?:\.\d+)?)\s*(?:磅|pt)", text)
        if m:
            attrs["line_spacing"] = float(m.group(1))
            attrs["line_spacing_type"] = "atLeast"
            return

        # "行距1.5"
        m = re.search(r"行距\s*(\d+(?:\.\d+)?)", text)
        if m:
            attrs["line_spacing"] = float(m.group(1))

    def _extract_indent(self, text, attrs):
        """提取缩进"""
        # "首行缩进2字符" "缩进2字符"
        m = re.search(r"(?:首行)?缩进\s*(\d+(?:\.\d+)?)\s*(?:字符|字)", text)
        if m:
            attrs["first_line_indent"] = float(m.group(1))
            return

        # "首行缩进0.75cm"
        m = re.search(r"(?:首行)?缩进\s*(\d+(?:\.\d+)?)\s*(?:cm|厘米)", text)
        if m:
            attrs["first_line_indent_cm"] = float(m.group(1))
            return

        # "不缩进" "无缩进"
        if "不缩进" in text or "无缩进" in text:
            attrs["first_line_indent"] = 0

        # "悬挂缩进"
        if "悬挂缩进" in text:
            attrs["hanging_indent"] = True

    def _extract_margin(self, text, attrs):
        """提取页边距"""
        # "上边距2.54cm"
        for direction, key in [("上", "top"), ("下", "bottom"), ("左", "left"), ("右", "right")]:
            m = re.search(rf"{direction}(?:边距|页距|距)\s*(\d+(?:\.\d+)?)\s*(?:cm|厘米)?", text)
            if m:
                attrs.setdefault("margins", {})[key] = float(m.group(1))

        # "上下2.54cm，左右3.17cm"
        m = re.search(r"上下\s*(\d+(?:\.\d+)?)", text)
        if m:
            val = float(m.group(1))
            attrs.setdefault("margins", {})["top"] = val
            attrs.setdefault("margins", {})["bottom"] = val
        m = re.search(r"左右\s*(\d+(?:\.\d+)?)", text)
        if m:
            val = float(m.group(1))
            attrs.setdefault("margins", {})["left"] = val
            attrs.setdefault("margins", {})["right"] = val

    def _extract_special_features(self, text, attrs):
        """提取特殊功能特征"""
        # 页眉内容
        m = re.search(r"页眉\s*[:：]?\s*[\"""“](.+?)[\"""”]", text)
        if m:
            attrs["header_text"] = m.group(1)
        elif "页眉" in text and "章名" in text:
            attrs["header_text"] = "{chapter}"
        elif "页眉" in text and "显示" in text:
            # "页眉显示论文题目"
            m = re.search(r"页眉\s*(?:显示|为|写)\s*(.+)", text)
            if m:
                attrs["header_text"] = m.group(1).strip()

        # 首页不同
        if "首页不同" in text or "首页无页眉" in text:
            attrs["different_first_page"] = True

        # 三线表
        if "三线表" in text or "三线表" in text:
            attrs["table_style"] = "three_line"

        # 表格边框
        if "无边框" in text and "表" in text:
            attrs["table_style"] = "no_border"
        elif "有边框" in text and "表" in text:
            attrs["table_style"] = "full_border"

        # 图片居中
        if "图片居中" in text or "插图居中" in text:
            attrs["image_center"] = True

    # ========== 规则存储与合并 ==========

    def _match_section_type(self, text):
        """匹配段落类型关键词（优先最长匹配）"""
        text = text.strip().lower()
        best_match = None
        best_len = 0
        for keyword, stype in SECTION_TYPE_MAP.items():
            if keyword in text and len(keyword) > best_len:
                best_match = stype
                best_len = len(keyword)
        return best_match

    def _store_rule(self, section_type, attrs, source="unknown"):
        """存储解析出的规则"""
        if section_type not in self.rules:
            self.rules[section_type] = {}
        existing = self.rules[section_type]
        for k, v in attrs.items():
            # 表格/列表解析的值优先，不被自然语言覆盖
            if k in existing and existing.get("_source") == "table":
                continue
            existing[k] = v
        existing["_source"] = source

    # ========== 配置生成 ==========

    def _build_config(self):
        """根据解析出的规则生成配置字典"""
        config = {"format_rules": {}}
        rules = self.rules

        # 字体
        font = {}
        if "chinese_font" in self._get_all_attrs(rules):
            font["chinese"] = self._get_most_common(rules, "chinese_font", "宋体")
        if "english_font" in self._get_all_attrs(rules):
            font["english"] = self._get_most_common(rules, "english_font", "Times New Roman")

        # 从 heading 规则中提取 heading_chinese
        for hk in ("heading1", "heading2", "heading3"):
            if hk in rules and "chinese_font" in rules[hk]:
                font["heading_chinese"] = rules[hk]["chinese_font"]
                break

        if font:
            config["format_rules"]["font"] = font

        # 标题
        headings = {}
        for level, config_key in [(1, "heading1"), (2, "heading2"), (3, "heading3")]:
            if config_key in rules:
                r = rules[config_key]
                h = {}
                if "font_size" in r:
                    h["font_size"] = r["font_size"]
                if "bold" in r:
                    h["bold"] = r["bold"]
                if "align" in r:
                    h["align"] = r["align"]
                if "line_spacing" in r:
                    h["line_spacing"] = self._build_line_spacing(r)
                if h:
                    headings[config_key] = h
        if headings:
            config["format_rules"]["headings"] = headings

        # 正文
        if "body_text" in rules:
            r = rules["body_text"]
            body = {}
            if "font_size" in r:
                body["font_size"] = r["font_size"]
            if "line_spacing" in r:
                body["line_spacing"] = self._build_line_spacing(r)
            if "first_line_indent" in r:
                body["first_line_indent"] = r["first_line_indent"]
            if "align" in r:
                body["align"] = r["align"]
            if body:
                config["format_rules"]["body_text"] = body

        # 摘要
        if "abstract" in rules:
            r = rules["abstract"]
            abs_cfg = {}
            if "font_size" in r:
                abs_cfg["content_font_size"] = r["font_size"]
            if "bold" in r:
                abs_cfg["title_bold"] = r["bold"]
            if "line_spacing" in r:
                abs_cfg["content_line_spacing"] = self._build_line_spacing(r)
            if "first_line_indent" in r:
                abs_cfg["content_first_line_indent"] = r["first_line_indent"]
            if abs_cfg:
                config["format_rules"]["abstract"] = abs_cfg

        # 关键词
        if "keywords" in rules:
            r = rules["keywords"]
            kw = {}
            if "font_size" in r:
                kw["font_size"] = r["font_size"]
            if "bold" in r:
                kw["bold_label"] = r["bold"]
            if kw:
                config["format_rules"]["keywords"] = kw

        # 参考文献
        if "references" in rules:
            r = rules["references"]
            ref = {}
            if "font_size" in r:
                ref["font_size"] = r["font_size"]
            if "line_spacing" in r:
                ref["line_spacing"] = r["line_spacing"]
            if "hanging_indent" in r:
                ref["hanging_indent"] = r["hanging_indent"]
            if ref:
                config["format_rules"]["references"] = ref

        # 题目页
        if "title_page" in rules:
            r = rules["title_page"]
            tp = {}
            if "font_size" in r:
                tp["title_font_size"] = r["font_size"]
            if "bold" in r:
                tp["title_bold"] = r["bold"]
            if "align" in r:
                tp["title_align"] = r["align"]
            if tp:
                config["format_rules"]["title_page"] = tp

        # 作者
        if "author" in rules:
            r = rules["author"]
            if "font_size" in r:
                config.setdefault("format_rules", {}).setdefault("title_page", {})[
                    "author_font_size"
                ] = r["font_size"]

        # 页边距
        all_margins = {}
        for stype, attrs in rules.items():
            if "margins" in attrs:
                all_margins.update(attrs["margins"])
        if all_margins:
            config["format_rules"]["margins"] = all_margins

        # 图表标题
        for kind, cfg_key in [("figure_caption", "figures"), ("table_caption", "tables")]:
            if kind in rules:
                r = rules[kind]
                cfg = {}
                if "font_size" in r:
                    cfg["font_size"] = r["font_size"]
                if "align" in r:
                    cfg["align"] = r["align"]
                if cfg:
                    config["format_rules"][cfg_key] = cfg

        # 页眉
        all_attrs = self._get_all_attrs(rules)
        header_cfg = {}
        if "header_text" in all_attrs:
            header_cfg["text"] = self._get_most_common(rules, "header_text", "")
        if "different_first_page" in all_attrs:
            config.setdefault("format_rules", {}).setdefault("header_footer", {})["different_first_page"] = True

        if "header" in rules:
            r = rules["header"]
            if "font_size" in r:
                header_cfg["font_size"] = r["font_size"]

        if header_cfg:
            config.setdefault("format_rules", {}).setdefault("header_footer", {}).setdefault("header", {}).update(header_cfg)

        # 页脚
        if "footer" in rules:
            r = rules["footer"]
            cfg = {}
            if "font_size" in r:
                cfg["font_size"] = r["font_size"]
            if cfg:
                config.setdefault("format_rules", {}).setdefault("header_footer", {}).setdefault("footer", {}).update(cfg)

        # 表格样式
        if "table_style" in all_attrs:
            style = self._get_most_common(rules, "table_style", "full_border")
            config.setdefault("format_rules", {})["_table_style"] = style

        return config

    def _build_line_spacing(self, r):
        """构建行距配置（支持简单值和字典格式）"""
        if "line_spacing_type" in r:
            return {"type": r["line_spacing_type"], "value": r["line_spacing"]}
        return r.get("line_spacing", 1.5)

    def _get_all_attrs(self, rules):
        """获取所有规则中的属性名集合"""
        attrs = set()
        for r in rules.values():
            attrs.update(r.keys())
        return attrs

    def _get_most_common(self, rules, attr, default):
        """获取某个属性出现最多的值"""
        values = [r[attr] for r in rules.values() if attr in r]
        if not values:
            return default
        from collections import Counter
        return Counter(values).most_common(1)[0][0]

    # ========== 调试输出 ==========

    def print_parsed_rules(self):
        """打印解析出的规则（调试用）"""
        print("=" * 50)
        print("解析出的格式规则:")
        print("=" * 50)
        for section, attrs in self.rules.items():
            print(f"\n  [{section}]")
            for k, v in attrs.items():
                print(f"    {k}: {v}")
        if self.warnings:
            print(f"\n警告 ({len(self.warnings)}):")
            for w in self.warnings:
                print(f"  - {w}")
