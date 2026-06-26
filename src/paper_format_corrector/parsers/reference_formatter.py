import re

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..utils.docx_utils import set_east_asian_font


class ReferenceFormatter:
    """参考文献格式矫正器 (GB/T 7714)"""

    # 常见文献类型标识
    TYPE_MAP = {
        "J": "journal",       # 期刊
        "M": "book",          # 专著
        "C": "conference",    # 会议
        "D": "thesis",        # 学位论文
        "R": "report",        # 报告
        "N": "newspaper",     # 报纸
        "S": "standard",      # 标准
        "P": "patent",        # 专利
        "EB/OL": "online",    # 网络文献
        "DB/OL": "online",
    }

    def __init__(self, config):
        ref_config = config.get("format_rules", {}).get("references", {})
        self.font_size = ref_config.get("font_size", 10.5)
        self.line_spacing = ref_config.get("line_spacing", 1.25)
        self.hanging_indent = ref_config.get("hanging_indent", True)
        self.numbering = ref_config.get("numbering", "sequential")
        self.font_rules = config.get("format_rules", {}).get("font", {})
        self.templates = ref_config.get("templates", {})

    # 引用风格常量
    CITATION_BRACKET = "bracket"        # [1], [2], [1-3]
    CITATION_SUPERSCRIPT = "superscript"  # ¹, ², ³
    CITATION_AUTHOR_YEAR = "author_year"  # (Smith, 2020), (Smith and Jones, 2020)
    CITATION_UNKNOWN = "unknown"

    # 上标数字 Unicode 映射
    _SUPERSCRIPT_DIGITS = set("⁰¹²³⁴⁵⁶⁷⁸⁹")

    def detect_citation_style(self, doc, ref_start_idx):
        """自动检测文档使用的引用风格

        Args:
            doc: Document 对象
            ref_start_idx: 参考文献起始段落索引

        Returns:
            str: CITATION_BRACKET / CITATION_SUPERSCRIPT / CITATION_AUTHOR_YEAR / CITATION_UNKNOWN
        """
        bracket_count = 0
        superscript_count = 0
        author_year_count = 0

        # 只扫描正文部分（参考文献之前）
        for i in range(min(ref_start_idx, len(doc.paragraphs))):
            text = doc.paragraphs[i].text.strip()
            if not text:
                continue

            # 方括号引用: [1], [2], [1-3], [1,2]
            bracket_matches = re.findall(r"\[\d+(?:[,\s\-–—]*\d+)*\]", text)
            bracket_count += len(bracket_matches)

            # 上标引用: ¹²³
            for char in text:
                if char in self._SUPERSCRIPT_DIGITS:
                    superscript_count += 1

            # 作者-年份引用: (Smith, 2020), (Smith and Jones, 2020), (张三, 2020)
            ay_matches = re.findall(
                r"\([A-Z][a-z一-鿿]+(?:\s+(?:and|&|和|、)\s+[A-Z][a-z一-鿿]+)*,?\s*\d{4}[a-z]?\)",
                text
            )
            author_year_count += len(ay_matches)

        # 判断主要引用风格
        max_count = max(bracket_count, superscript_count, author_year_count)
        if max_count == 0:
            return self.CITATION_UNKNOWN

        if bracket_count == max_count:
            return self.CITATION_BRACKET
        elif superscript_count == max_count:
            return self.CITATION_SUPERSCRIPT
        else:
            return self.CITATION_AUTHOR_YEAR

    def get_citation_style_name(self, style):
        """返回引用风格的中文名称"""
        names = {
            self.CITATION_BRACKET: "方括号编号 [1]",
            self.CITATION_SUPERSCRIPT: "上标数字 ¹",
            self.CITATION_AUTHOR_YEAR: "作者-年份 (Smith, 2020)",
            self.CITATION_UNKNOWN: "未知",
        }
        return names.get(style, "未知")

    def format_references(self, doc, ref_start_idx, ref_end_idx=None):
        """格式化参考文献区域

        Args:
            doc: python-docx Document 对象
            ref_start_idx: 参考文献标题段落索引
            ref_end_idx: 参考文献区域结束索引（不含），None 表示自动检测
        """
        paragraphs = doc.paragraphs
        if ref_start_idx >= len(paragraphs):
            return

        # 确定上界
        end_bound = ref_end_idx if ref_end_idx is not None else len(paragraphs)
        end_bound = min(end_bound, len(paragraphs))

        # 格式化参考文献标题
        title_para = paragraphs[ref_start_idx]
        self._format_ref_title(title_para)

        # 格式化每条参考文献
        ref_num = 1
        for i in range(ref_start_idx + 1, end_bound):
            para = paragraphs[i]
            text = para.text.strip()
            if not text:
                continue

            # 如果遇到新的顶级标题，停止（仅在自动检测模式下）
            if ref_end_idx is None and self._is_new_section(text) and i > ref_start_idx + 1:
                break

            self._format_ref_item(para, ref_num)
            self._reformat_reference(para, ref_num)
            ref_num += 1

    def _format_ref_title(self, paragraph):
        """格式化参考文献标题"""
        rules = self.font_rules
        for run in paragraph.runs:
            run.font.name = rules.get("english", "Times New Roman")
            set_east_asian_font(run, rules.get("chinese", "宋体"))
            run.font.size = Pt(self.font_size)
            run.font.bold = True

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)

    def _format_ref_item(self, paragraph, num):
        """格式化单条参考文献"""
        rules = self.font_rules
        for run in paragraph.runs:
            run.font.name = rules.get("english", "Times New Roman")
            set_east_asian_font(run, rules.get("chinese", "宋体"))
            run.font.size = Pt(self.font_size)
            run.font.bold = False

        paragraph.paragraph_format.line_spacing = self.line_spacing
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

        if self.hanging_indent:
            paragraph.paragraph_format.first_line_indent = Pt(-24)  # 悬挂缩进
            paragraph.paragraph_format.left_indent = Pt(24)

    def _parse_reference_fields(self, text):
        """解析参考文献条目的结构化字段

        Args:
            text: 单条参考文献文本

        Returns:
            dict: 包含 num, type, authors, title, journal, year, volume, pages 等字段
        """
        fields = {"raw": text}

        # 提取编号 [N]
        m = re.match(r"^\[(\d+)\]", text)
        if m:
            fields["num"] = int(m.group(1))
            text_after_num = text[m.end():].strip()
        else:
            text_after_num = text.strip()

        # 提取文献类型标识 [J], [M], [EB/OL] 等
        type_match = re.search(r"\[([A-Z](?:/[A-Z]+)?)\]", text_after_num)
        if type_match:
            fields["type"] = type_match.group(1)
            fields["type_name"] = self.TYPE_MAP.get(fields["type"], "unknown")

        # 提取作者（类型标识之前的文本）
        if type_match:
            before_type = text_after_num[:type_match.start()].strip()
            # 去掉末尾的句号
            before_type = before_type.rstrip(".")
            if before_type:
                fields["authors"] = before_type

        # 提取标题（类型标识之后到下一个句号或逗号之间的文本）
        if type_match:
            after_type = text_after_num[type_match.end():].strip()
            # 标题通常是第一个句号之前的内容
            title_end = after_type.find(".")
            if title_end > 0:
                fields["title"] = after_type[:title_end].strip()
                rest = after_type[title_end + 1:].strip()
            else:
                fields["title"] = after_type
                rest = ""

            # 从剩余文本中提取期刊/年份/卷/页
            if rest:
                # 提取期刊名（逗号之前）
                journal_match = re.match(r"([^,]+)", rest)
                if journal_match:
                    fields["journal"] = journal_match.group(1).strip()

                # 提取年份
                year_match = re.search(r"(\d{4})", rest)
                if year_match:
                    fields["year"] = year_match.group(1)

                # 提取卷(期): 页码
                vol_match = re.search(r"(\d+)\((\d+)\):\s*(\d+(?:-\d+)?)", rest)
                if vol_match:
                    fields["volume"] = vol_match.group(1)
                    fields["number"] = vol_match.group(2)
                    fields["pages"] = vol_match.group(3)

        return fields

    def _format_ref_by_template(self, ref_type, fields, num):
        """使用模板渲染参考文献

        Args:
            ref_type: 文献类型 (journal, book, conference, etc.)
            fields: 解析出的字段字典
            num: 编号

        Returns:
            str: 渲染后的参考文献文本，无模板时返回 None
        """
        template = self.templates.get(ref_type)
        if not template:
            return None

        # 填充模板
        result = template
        result = result.replace("{num}", str(num))
        for key, value in fields.items():
            if key not in ("raw", "type", "type_name", "num"):
                result = result.replace(f"{{{key}}}", str(value))

        # 清理未填充的占位符
        result = re.sub(r"\{[a-z_]+\}", "", result)
        # 清理多余的空格和标点
        result = re.sub(r"\s+", " ", result)
        result = re.sub(r"\.\.", ".", result)
        result = re.sub(r"\.\s*\.", ".", result)
        result = result.strip()

        return result

    def _reformat_reference(self, paragraph, num):
        """解析并重排单条参考文献

        如果有模板，使用模板重排；否则只做格式化。

        Args:
            paragraph: 参考文献段落
            num: 编号
        """
        if not self.templates:
            return  # 无模板，跳过重排

        text = paragraph.text.strip()
        fields = self._parse_reference_fields(text)

        # 确定文献类型
        ref_type = fields.get("type_name", "unknown")
        formatted = self._format_ref_by_template(ref_type, fields, num)

        if formatted and formatted != text:
            # 重写段落内容
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = formatted
            else:
                paragraph.add_run(formatted)

    def validate_references(self, doc, ref_start_idx):
        """验证参考文献格式，返回问题列表"""
        issues = []
        paragraphs = doc.paragraphs

        ref_num = 0
        for i in range(ref_start_idx + 1, len(paragraphs)):
            text = paragraphs[i].text.strip()
            if not text:
                continue
            if self._is_new_section(text):
                break

            ref_num += 1
            problems = self._check_single_ref(text, ref_num)
            issues.extend(problems)

        return issues

    def _check_single_ref(self, text, expected_num):
        """检查单条参考文献格式"""
        issues = []

        # 检查编号格式 [1]
        m = re.match(r"^\[(\d+)\]", text)
        if m:
            num = int(m.group(1))
            if num != expected_num:
                issues.append(
                    f"参考文献 [{num}] 编号不连续，应为 [{expected_num}]"
                )
        else:
            # 没有 [N] 格式的编号
            m2 = re.match(r"^(\d+)[\.\)]", text)
            if m2:
                issues.append(
                    f"参考文献 {m2.group(1)} 建议使用 [N] 格式编号"
                )
            else:
                issues.append(f"参考文献 {expected_num} 缺少编号")

        # 检查是否包含文献类型标识
        if not re.search(r"\[[A-Z]", text):
            issues.append(f"参考文献 [{expected_num}] 缺少文献类型标识 (如 [J], [M])")

        # 检查作者格式 (中文文献作者间应用逗号)
        if re.search(r"[一-鿿]{2,4}\s+[一-鿿]{2,4}", text):
            if "，" not in text and "," not in text[:text.find(".")]:
                issues.append(f"参考文献 [{expected_num}] 多位作者间建议用逗号分隔")

        return issues

    def _is_new_section(self, text):
        """判断是否进入新的章节"""
        if re.match(r"^第[一二三四五六七八九十\d]+[章部分篇]", text):
            return True
        if re.match(r"^\d+\.\d+", text):
            return True
        return False

    def check_citation_consistency(self, doc, ref_start_idx):
        """检查正文引用与参考文献列表的一致性

        Returns:
            list of dict: [{"type": "orphan"|"missing"|"duplicate", "message": str}]
        """
        issues = []
        paragraphs = doc.paragraphs

        # 1. 收集正文中所有引用编号
        cited_nums = set()
        for i in range(ref_start_idx):
            text = paragraphs[i].text.strip()
            if not text:
                continue
            # 匹配 [1], [2], [1-3], [1,2,3], [1, 2, 3] 等
            for m in re.finditer(r"\[(\d+)(?:[,\s\-–—]*(\d+))*\]", text):
                for num_str in re.findall(r"\d+", m.group(0)):
                    cited_nums.add(int(num_str))

        # 2. 收集参考文献列表中的编号
        ref_nums = {}
        for i in range(ref_start_idx + 1, len(paragraphs)):
            text = paragraphs[i].text.strip()
            if not text:
                continue
            if self._is_new_section(text):
                break
            m = re.match(r"^\[(\d+)\]", text)
            if m:
                num = int(m.group(1))
                if num in ref_nums:
                    issues.append({"type": "duplicate", "message": f"参考文献 [{num}] 重复出现"})
                ref_nums[num] = text[:50]

        # 3. 检查孤立引用（正文引用了但参考文献列表中没有）
        for num in sorted(cited_nums):
            if num not in ref_nums:
                issues.append({"type": "orphan", "message": f"正文引用 [{num}] 在参考文献列表中不存在"})

        # 4. 检查未引用文献（参考文献列表中有但正文未引用）
        for num in sorted(ref_nums.keys()):
            if num not in cited_nums:
                issues.append({"type": "missing", "message": f"参考文献 [{num}] 未在正文中被引用"})

        return issues
