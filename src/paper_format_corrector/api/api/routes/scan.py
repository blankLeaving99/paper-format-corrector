"""Document scan / structure analysis endpoint.

POST /api/v1/scan — analyze document structure and return element counts.
"""

from __future__ import annotations

import tempfile
from pathlib import Path
from typing import Any

from fastapi import APIRouter, File, HTTPException, UploadFile

router = APIRouter(tags=["scan"])


def _analyze_document(doc_path: str) -> dict[str, Any]:
    """Run the DocumentAnalyzer and return a serializable dict."""
    from paper_format_corrector.core.document.analyzer import DocumentAnalyzer

    analyzer = DocumentAnalyzer()
    analysis = analyzer.analyze(Path(doc_path))

    # Count headings by level
    heading_counts: dict[str, int] = {}
    for para in analysis.paragraphs:
        ptype = para.paragraph_type.value
        if ptype.startswith("heading"):
            heading_counts[ptype] = heading_counts.get(ptype, 0) + 1

    # Count other element types
    structure = analysis.structure

    # Separate figure/table captions from images/tables in docx
    figure_captions = structure.get("figure_caption", 0)
    table_captions = structure.get("table_caption", 0)

    # Count actual images and tables from docx XML
    from docx import Document

    doc = Document(doc_path)
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1
    table_count = len(doc.tables)

    # Count formulas and code blocks
    formula_count = structure.get("formula", 0)
    code_count = structure.get("code", 0)

    return {
        "total_paragraphs": analysis.total_paragraphs,
        "structure": structure,
        "headings": heading_counts,
        "summary": {
            "paragraph_count": analysis.total_paragraphs,
            "heading_count": sum(heading_counts.values()),
            "image_count": image_count,
            "table_count": table_count,
            "formula_count": formula_count,
            "code_count": code_count,
            "figure_caption_count": figure_captions,
            "table_caption_count": table_captions,
        },
        "fonts_used": analysis.fonts_used,
        "font_sizes_used": {
            str(k): v for k, v in analysis.font_sizes_used.items()
        },
        "metadata": analysis.metadata,
        "issues": analysis.issues,
    }


@router.post(
    "/api/v1/scan",
    summary="扫描文档结构",
    description=(
        "上传 .docx 文件并分析其文档结构。返回标题层级、段落数量、"
        "图片数量、表格数量、公式数量等结构化信息。"
    ),
    response_model=dict[str, Any],
    responses={
        200: {
            "description": "文档结构分析结果",
            "content": {
                "application/json": {
                    "example": {
                        "total_paragraphs": 120,
                        "headings": {
                            "heading1": 5,
                            "heading2": 15,
                            "heading3": 8,
                        },
                        "summary": {
                            "paragraph_count": 120,
                            "heading_count": 28,
                            "image_count": 3,
                            "table_count": 2,
                            "formula_count": 10,
                            "code_count": 0,
                            "figure_caption_count": 3,
                            "table_caption_count": 2,
                        },
                        "fonts_used": {"宋体": 80, "Times New Roman": 40},
                        "metadata": {
                            "margins": {
                                "top": 2.54,
                                "bottom": 2.54,
                                "left": 3.17,
                                "right": 3.17,
                            }
                        },
                    }
                }
            },
        },
        400: {"description": "文件格式不支持"},
    },
)
async def scan_document(
    file: UploadFile = File(..., description="待扫描的 .docx 文件"),
) -> dict[str, Any]:
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="仅支持 .docx 格式文件")

    with tempfile.TemporaryDirectory() as tmp_dir:
        input_path = Path(tmp_dir) / file.filename
        content = await file.read()
        input_path.write_bytes(content)

        try:
            return _analyze_document(str(input_path))
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"扫描失败: {exc}")
