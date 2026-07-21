"""Services behind the interactive paper-format workbench.

Enhanced with confidence scoring, dry-run correction plans, and detailed
modification reports as specified in zhinan.md sections 3.1 and 3.2.
"""

from __future__ import annotations

import re
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document

from ...infrastructure.parsers.document_analyzer import DocumentAnalyzer, ParagraphInfo

# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------

@dataclass
class ConfidenceItem:
    element: str
    confidence: str  # "high", "medium", "low"
    score: float  # 0.0 ~ 1.0
    reason: str
    samples: list[dict[str, Any]] = field(default_factory=list)


@dataclass
class CorrectionPlanItem:
    element_type: str
    element_count: int
    action: str
    target_rule: dict[str, Any]
    source: str  # "manual", "sample", "template", "default"
    confidence: str = "high"
    warnings: list[str] = field(default_factory=list)


@dataclass
class CorrectionPlan:
    items: list[CorrectionPlanItem]
    total_affected: int = 0
    risk_items: list[dict[str, Any]] = field(default_factory=list)


@dataclass
class ModificationReport:
    applied: dict[str, int]
    skipped: list[dict[str, str]]
    warnings: list[str]
    risk_items: list[dict[str, Any]]
    rule_sources: dict[str, str]
    total_elements: int = 0
    modified_elements: int = 0


# ---------------------------------------------------------------------------
# Core scanning
# ---------------------------------------------------------------------------

def scan_document(path: str | Path) -> dict[str, Any]:
    """Return a comprehensive, UI-friendly inventory of a Word document.

    Enhanced with confidence scores for each detected element type.
    """
    document_path = Path(path)
    if document_path.suffix.lower() != ".docx":
        raise ValueError("格式工作台目前仅支持 .docx 文档")

    analysis = DocumentAnalyzer().analyze(document_path)
    document = Document(str(document_path))

    samples: dict[str, list[dict[str, Any]]] = {}
    for paragraph in analysis.paragraphs:
        key = paragraph.paragraph_type.value
        samples.setdefault(key, [])
        if paragraph.text and len(samples[key]) < 5:
            samples[key].append(_paragraph_summary(paragraph))

    confidence_items = _compute_confidence(analysis.paragraphs, document)

    elements = {
        **analysis.structure,
        "table": len(document.tables),
        "image": _count_images(document),
        "total_paragraphs": len(document.paragraphs),
    }

    return {
        "elements": elements,
        "samples": samples,
        "margins": analysis.metadata.get("margins", {}),
        "issues": analysis.issues,
        "confidence": [_confidence_to_dict(c) for c in confidence_items],
        "page_setup": _extract_page_setup(document),
        "header_footer": _extract_header_footer(document),
        "font_summary": _summarize_fonts(analysis),
    }


# ---------------------------------------------------------------------------
# Sample learning
# ---------------------------------------------------------------------------

def learn_style_profile(path: str | Path) -> dict[str, Any]:
    """Infer a reusable format configuration from a well-formatted sample.

    Enhanced to learn more element types and provide confidence info.
    """
    analysis = DocumentAnalyzer().analyze(path)
    grouped: dict[str, list[ParagraphInfo]] = {}
    for paragraph in analysis.paragraphs:
        if paragraph.text:
            grouped.setdefault(paragraph.paragraph_type.value, []).append(paragraph)

    config: dict[str, Any] = {"format_rules": {"margins": analysis.metadata.get("margins", {})}}
    rules = config["format_rules"]

    # 学习正文字体和段落规则
    body = _dominant_style(grouped.get("body", []))
    if body:
        rules["body_text"] = _to_paragraph_rule(body, include_indent=True)
        rules["font"] = {
            "chinese": body.get("font_name") or "宋体",
            "english": body.get("font_name") or "Times New Roman",
        }

    # 学习标题层级
    headings: dict[str, Any] = {}
    for element_name in ("heading1", "heading2", "heading3"):
        style = _dominant_style(grouped.get(element_name, []))
        if style:
            headings[element_name] = _to_paragraph_rule(style)
    if headings:
        rules["headings"] = headings

    # 学习摘要格式
    abstract_style = _dominant_style(grouped.get("abstract", []))
    if abstract_style:
        rules["abstract"] = {
            "content": _to_paragraph_rule(abstract_style, include_indent=True),
        }

    # 学习关键词格式
    keywords_style = _dominant_style(grouped.get("keywords_cn", []) + grouped.get("keywords_en", []))
    if keywords_style:
        rules["keywords"] = _to_paragraph_rule(keywords_style)

    # 学习表格格式
    doc = Document(str(path))
    table_rule = _learn_table_rule(doc)
    if table_rule:
        rules["tables"] = table_rule

    # 学习图片格式
    image_rule = _learn_image_rule(doc)
    if image_rule:
        rules["images"] = image_rule

    # 学习参考文献格式
    ref_style = _dominant_style(grouped.get("reference", []))
    if ref_style:
        rules["references"] = _to_paragraph_rule(ref_style)

    # 学习代码块格式
    code_style = _dominant_style(grouped.get("code", []))
    if code_style:
        rules["code"] = {
            "mono_font": code_style.get("font_name") or "Consolas",
            "mono_font_size": code_style.get("font_size") or 10,
        }

    # 学习公式格式
    formula_style = _dominant_style(grouped.get("formula", []) + grouped.get("formula_content", []))
    if formula_style:
        rules["formulas"] = {
            "font_size": formula_style.get("font_size") or 12,
            "center": formula_style.get("align") == "center",
        }

    # 学习页眉页脚
    hf_rule = _learn_header_footer(doc)
    if hf_rule:
        rules["header_footer"] = hf_rule

    return config


def explain_style_profile(path: str | Path) -> dict[str, Any]:
    """Return a reviewable learning result with confidence information."""
    profile = learn_style_profile(path)
    analysis = DocumentAnalyzer().analyze(path)
    rules = profile["format_rules"]
    learned: list[dict[str, Any]] = []
    for source, label in (
        ("body_text", "正文"), ("headings", "标题"), ("tables", "表格"),
        ("images", "图片"), ("abstract", "摘要"), ("references", "参考文献"),
        ("keywords", "关键词"), ("code", "代码块"), ("formulas", "公式"),
        ("header_footer", "页眉页脚"),
    ):
        if source in rules:
            learned.append({"element": label, "rule": rules[source]})

    confidence = "高" if rules.get("body_text") and rules.get("headings") else "中" if rules.get("body_text") else "低"
    return {
        "learned": learned,
        "margins": rules.get("margins", {}),
        "confidence": confidence,
        "notice": "只会采用重复出现的段落格式；单个例外格式不会作为全局规则。",
        "source_issues": analysis.issues,
        "elements_scanned": len(analysis.paragraphs),
    }


# ---------------------------------------------------------------------------
# Dry-run correction plan
# ---------------------------------------------------------------------------

def build_correction_plan(
    path: str | Path,
    format_rules: dict[str, Any],
    source: str = "manual",
) -> CorrectionPlan:
    """Generate a correction plan without applying changes.

    Shows what would be modified and identifies risk items.
    """
    inventory = scan_document(path)
    elements = inventory["elements"]
    items: list[CorrectionPlanItem] = []
    risk_items: list[dict[str, Any]] = []

    body_config = format_rules.get("body_text", {})
    if body_config:
        body_count = elements.get("body", 0)
        items.append(CorrectionPlanItem(
            element_type="正文段落", element_count=body_count,
            action="统一字体、字号、行距、缩进、对齐",
            target_rule=body_config, source=source,
        ))

    headings = format_rules.get("headings", {})
    for level in ("heading1", "heading2", "heading3"):
        if level in headings:
            count = elements.get(level, 0)
            level_names = {"heading1": "一级标题", "heading2": "二级标题", "heading3": "三级标题"}
            items.append(CorrectionPlanItem(
                element_type=level_names.get(level, level), element_count=count,
                action="统一字体、字号、加粗、对齐",
                target_rule=headings[level], source=source,
            ))

    tables_config = format_rules.get("tables", {})
    if tables_config:
        table_count = elements.get("table", 0)
        style = tables_config.get("style", "未指定")
        items.append(CorrectionPlanItem(
            element_type="表格", element_count=table_count,
            action=f"应用表格样式: {style}",
            target_rule=tables_config, source=source,
        ))

    images_config = format_rules.get("images", {})
    if images_config:
        image_count = elements.get("image", 0)
        items.append(CorrectionPlanItem(
            element_type="图片", element_count=image_count,
            action="统一宽度、居中对齐",
            target_rule=images_config, source=source,
        ))

    for confidence_item in inventory.get("confidence", []):
        if confidence_item.get("confidence") == "low":
            risk_items.append({
                "element": confidence_item["element"],
                "reason": confidence_item.get("reason", ""),
                "samples": confidence_item.get("samples", [])[:2],
            })

    protected = {"code": "代码段", "formula": "公式"}
    for key, label in protected.items():
        count = elements.get(key, 0)
        if count:
            items.append(CorrectionPlanItem(
                element_type=label, element_count=count,
                action="保留原格式，不做修改",
                target_rule={}, source="保护策略",
                confidence="high",
                warnings=[f"{label}默认保留原格式"],
            ))

    total_affected = sum(item.element_count for item in items if item.action != "保留原格式，不做修改")

    return CorrectionPlan(
        items=items, total_affected=total_affected, risk_items=risk_items,
    )


def plan_to_dict(plan: CorrectionPlan) -> dict[str, Any]:
    """Convert a CorrectionPlan to a JSON-serializable dict."""
    return {
        "items": [
            {
                "element_type": item.element_type,
                "element_count": item.element_count,
                "action": item.action,
                "source": item.source,
                "confidence": item.confidence,
                "warnings": item.warnings,
            }
            for item in plan.items
        ],
        "total_affected": plan.total_affected,
        "risk_items": plan.risk_items,
    }


# ---------------------------------------------------------------------------
# Application report
# ---------------------------------------------------------------------------

def build_application_report(path: str | Path, report: dict[str, Any]) -> dict[str, Any]:
    """Build a detailed modification report explaining what changed and what didn't."""
    inventory = scan_document(path)
    elements = inventory["elements"]

    applied = {
        "paragraphs": report.get("paragraphs_corrected", 0),
        "headings": report.get("headings_fixed", 0),
        "body": report.get("body_fixed", 0),
        "tables": report.get("tables_formatted", 0),
        "images": report.get("images_centered", 0),
    }

    unchanged: list[dict[str, str]] = []
    protected = {
        "code": "代码段默认保留原格式，防止破坏缩进和语法展示。",
        "formula": "公式内容默认保留，仅处理可识别的编号。",
        "unknown": "无法识别的段落未按正文或标题强制修改，请人工确认其类型。",
    }
    for element, reason in protected.items():
        count = elements.get(element, 0)
        if count:
            unchanged.append({"element": element, "count": str(count), "reason": reason})

    if elements.get("image", 0) > report.get("images_centered", 0):
        unchanged.append({
            "element": "image",
            "count": str(elements["image"] - report.get("images_centered", 0)),
            "reason": "未识别为行内图片的对象保持原样。",
        })

    if elements.get("table", 0) > report.get("tables_formatted", 0):
        unchanged.append({
            "element": "table",
            "count": str(elements["table"] - report.get("tables_formatted", 0)),
            "reason": "未识别的表格保持原样。",
        })

    total_elements = sum(elements.get(k, 0) for k in ("body", "heading1", "heading2", "heading3", "table", "image"))
    modified_elements = sum(applied.values())

    return {
        "applied": applied,
        "needs_review": unchanged,
        "detected_issues": inventory["issues"],
        "total_elements": total_elements,
        "modified_elements": modified_elements,
        "coverage_rate": f"{modified_elements / total_elements * 100:.1f}%" if total_elements > 0 else "0%",
        "rule_sources": report.get("rule_sources", {}),
        "risk_items": _identify_risks(inventory),
    }


# ---------------------------------------------------------------------------
# Manual style configuration
# ---------------------------------------------------------------------------

def manual_style_config(
    body_font: str, body_size: float, body_line_spacing: float, body_indent: float,
    heading1_size: float, heading2_size: float, heading3_size: float, heading_font: str,
    table_style: str, table_font_size: float, image_max_width: str,
    body_en_font: str = "Times New Roman",
    heading1_bold: bool = True, heading2_bold: bool = True, heading3_bold: bool = True,
    heading1_align: str = "center", heading2_align: str = "left", heading3_align: str = "left",
    abstract_size: float = 12, abstract_indent: float = 0,
    ref_size: float = 10.5, ref_line_spacing: float = 1.25,
) -> dict[str, Any]:
    """Translate workbench form controls into the configuration schema."""
    return {"format_rules": {
        "font": {"chinese": body_font, "english": body_en_font, "heading_chinese": heading_font},
        "body_text": {
            "font_size": float(body_size), "line_spacing": float(body_line_spacing),
            "first_line_indent": float(body_indent), "align": "justify",
        },
        "headings": {
            "heading1": {"font_size": float(heading1_size), "bold": heading1_bold, "align": heading1_align},
            "heading2": {"font_size": float(heading2_size), "bold": heading2_bold, "align": heading2_align},
            "heading3": {"font_size": float(heading3_size), "bold": heading3_bold, "align": heading3_align},
        },
        "abstract": {"title": {"font_size": float(heading1_size), "bold": True, "align": "center"},
                      "content": {"font_size": float(abstract_size), "line_spacing": float(body_line_spacing),
                                  "first_line_indent": float(abstract_indent), "align": "justify"}},
        "tables": {"font_size": float(table_font_size), "style": table_style},
        "images": {"max_width": image_max_width, "alignment": "center"},
        "references": {"font_size": float(ref_size), "line_spacing": float(ref_line_spacing)},
    }}


# ---------------------------------------------------------------------------
# Private helpers
# ---------------------------------------------------------------------------

def _paragraph_summary(paragraph: ParagraphInfo) -> dict[str, Any]:
    return {
        "position": paragraph.index + 1,
        "text": paragraph.text[:100],
        "font": paragraph.font_name,
        "size": paragraph.font_size,
        "bold": paragraph.is_bold,
        "alignment": paragraph.alignment,
    }


def _dominant_style(paragraphs: list[ParagraphInfo]) -> dict[str, Any]:
    if not paragraphs:
        return {}
    signatures = Counter(
        (p.font_name, p.font_size, p.is_bold, p.alignment, p.line_spacing, p.first_line_indent)
        for p in paragraphs
    )
    signature, count = signatures.most_common(1)[0]
    if len(paragraphs) > 1 and count / len(paragraphs) < 0.5:
        return {}
    return dict(zip(
        ("font_name", "font_size", "bold", "align", "line_spacing", "first_line_indent"),
        signature,
    ))


def _to_paragraph_rule(style: dict[str, Any], include_indent: bool = False) -> dict[str, Any]:
    rule: dict[str, Any] = {
        "font_size": style.get("font_size") or 12,
        "bold": bool(style.get("bold")),
        "align": style.get("align") or "left",
        "line_spacing": style.get("line_spacing") or 1.0,
    }
    if include_indent:
        indent_raw = float(style.get("first_line_indent") or 0)
        rule["first_line_indent"] = round(indent_raw / 152400, 1) if indent_raw > 15240 else round(indent_raw, 1)
    return rule


def _learn_table_rule(document: Document) -> dict[str, Any]:
    if not document.tables:
        return {}
    table = document.tables[0]
    header_runs = table.rows[0].cells[0].paragraphs[0].runs if table.rows else []
    return {"header_bold": any(run.bold for run in header_runs), "style": "three_line"}


def _learn_image_rule(document: Document) -> dict[str, Any]:
    widths: list[float] = []
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            drawings = run._element.xpath(".//w:drawing")
            for drawing in drawings:
                extents = drawing.xpath(".//wp:extent")
                for extent in extents:
                    cx = int(extent.get("cx", 0))
                    if cx > 0:
                        widths.append(cx / 914400)
    if not widths:
        return {}
    avg_width = sum(widths) / len(widths)
    return {"max_width": f"{avg_width:.1f}cm", "alignment": "center"}


def _learn_header_footer(document: Document) -> dict[str, Any]:
    """学习页眉页脚格式"""
    if not document.sections:
        return {}
    section = document.sections[0]
    result: dict[str, Any] = {}

    # 检查页眉
    try:
        header = section.header
        if header and header.paragraphs:
            for para in header.paragraphs:
                if para.text.strip():
                    result["header_text"] = para.text.strip()
                    if para.runs:
                        run = para.runs[0]
                        if run.font.size:
                            result["header_font_size"] = run.font.size.pt
                        if run.font.bold is not None:
                            result["header_bold"] = run.font.bold
                    break
    except Exception:
        pass

    # 检查页脚
    try:
        footer = section.footer
        if footer and footer.paragraphs:
            for para in footer.paragraphs:
                if para.text.strip():
                    result["footer_text"] = para.text.strip()
                    break
    except Exception:
        pass

    # 检查页码位置
    try:
        from docx.oxml.ns import qn
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn("w:pgNumType"))
        if pgNumType is not None:
            start = pgNumType.get(qn("w:start"))
            if start:
                result["page_number_start"] = int(start)
    except Exception:
        pass

    return result if result else {}


def _count_images(document: Document) -> int:
    return sum(
        1 for paragraph in document.paragraphs
        for run in paragraph.runs
        if run._element.xpath(".//w:drawing")
    )


def _compute_confidence(paragraphs: list[ParagraphInfo], document: Document) -> list[ConfidenceItem]:
    items: list[ConfidenceItem] = []
    grouped: dict[str, list[ParagraphInfo]] = {}
    for p in paragraphs:
        grouped.setdefault(p.paragraph_type.value, []).append(p)

    for ptype, group in grouped.items():
        if not group:
            continue
        total = len(group)
        if total >= 3:
            signatures = Counter((p.font_name, p.font_size, p.is_bold) for p in group)
            top_count = signatures.most_common(1)[0][1]
            ratio = top_count / total
            if ratio >= 0.8:
                conf, score = "high", min(ratio, 1.0)
                reason = f"该类型{total}个段落中{top_count}个格式一致"
            elif ratio >= 0.5:
                conf, score = "medium", ratio
                reason = f"该类型{total}个段落中{top_count}个格式一致，部分存在差异"
            else:
                conf, score = "low", ratio
                reason = f"该类型{total}个段落格式不统一"
        elif total == 1:
            conf, score = "medium", 0.6
            reason = "仅有一个样本，无法确认是否为通用样式"
        else:
            conf, score = "low", 0.3
            reason = "样本过少"

        samples = [_paragraph_summary(p) for p in group[:3]]
        items.append(ConfidenceItem(element=ptype, confidence=conf, score=score, reason=reason, samples=samples))

    table_count = len(document.tables)
    if table_count > 0:
        items.append(ConfidenceItem(
            element="table", confidence="high" if table_count >= 2 else "medium",
            score=0.8 if table_count >= 2 else 0.6,
            reason=f"检测到{table_count}个表格",
        ))

    return items


def _extract_page_setup(document: Document) -> dict[str, Any]:
    if not document.sections:
        return {}
    section = document.sections[0]
    return {
        "page_width_cm": round(section.page_width / 360000, 2),
        "page_height_cm": round(section.page_height / 360000, 2),
        "left_margin_cm": round(section.left_margin / 360000, 2),
        "right_margin_cm": round(section.right_margin / 360000, 2),
        "top_margin_cm": round(section.top_margin / 360000, 2),
        "bottom_margin_cm": round(section.bottom_margin / 360000, 2),
    }


def _extract_header_footer(document: Document) -> dict[str, Any]:
    """提取页眉页脚信息"""
    result: dict[str, Any] = {}
    if not document.sections:
        return result
    section = document.sections[0]

    # 页眉
    try:
        header = section.header
        if header and not header.is_linked_to_previous and header.paragraphs:
            texts = [p.text.strip() for p in header.paragraphs if p.text.strip()]
            if texts:
                result["header"] = texts[0][:50]
    except Exception:
        pass

    # 页脚
    try:
        footer = section.footer
        if footer and not footer.is_linked_to_previous and footer.paragraphs:
            texts = [p.text.strip() for p in footer.paragraphs if p.text.strip()]
            if texts:
                result["footer"] = texts[0][:50]
    except Exception:
        pass

    return result


def _summarize_fonts(analysis: DocumentAnalysis) -> dict[str, Any]:
    """字体使用摘要"""
    total = sum(analysis.fonts_used.values()) or 1
    top_fonts = sorted(analysis.fonts_used.items(), key=lambda x: -x[1])[:5]
    return {
        "total_fonts": len(analysis.fonts_used),
        "top_fonts": [{"name": name, "count": count, "ratio": f"{count/total*100:.1f}%"} for name, count in top_fonts],
        "total_sizes": len(analysis.font_sizes_used),
        "sizes_used": sorted(analysis.font_sizes_used.keys()),
    }


def _identify_risks(inventory: dict[str, Any]) -> list[dict[str, Any]]:
    risks: list[dict[str, Any]] = []
    for item in inventory.get("confidence", []):
        if item.get("confidence") == "low":
            risks.append({
                "type": "识别不确定",
                "element": item.get("element", ""),
                "detail": item.get("reason", ""),
            })
    for issue in inventory.get("issues", []):
        risks.append({"type": "格式问题", "detail": str(issue)})
    return risks


def _confidence_to_dict(c: ConfidenceItem) -> dict[str, Any]:
    return {
        "element": c.element,
        "confidence": c.confidence,
        "score": round(c.score, 2),
        "reason": c.reason,
        "samples": c.samples,
    }
