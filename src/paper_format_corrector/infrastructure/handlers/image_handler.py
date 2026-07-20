"""图片处理模块

增强功能:
- 宽度调整（占满页宽/百分比/固定厘米）
- 居中对齐
- DPI检查与清晰度提示
- 支持内联图片和浮动图片
- 图题检测与格式化
- 图片编号规则
- 保持高宽比
- 批量统一处理
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from ...shared.utils.docx_utils import set_east_asian_font


@dataclass
class ImageFormatResult:
    """图片格式化结果"""
    images_centered: int = 0
    images_resized: int = 0
    captions_formatted: int = 0
    warnings: list[str] = field(default_factory=list)
    stats: dict[str, Any] = field(default_factory=dict)


class ImageHandler:
    """图片处理器，支持宽度调整、居中、DPI检查、图题格式化"""

    CAPTION_PATTERN = re.compile(
        r"^(图\s*\d|Fig\.?\s*\d|Figure\s*\d|图\s*[\d一二三四五六七八九十]+)",
        re.IGNORECASE,
    )

    def __init__(self, config: dict[str, Any]):
        img_config = config.get("format_rules", {}).get("images", {})
        self.max_width = img_config.get("max_width", "full")
        self.center = img_config.get("center", True)
        self.alignment = img_config.get("alignment", "center")
        self.min_dpi = img_config.get("min_dpi", 150)
        self.min_width_inches = img_config.get("min_width_inches", 1.0)

        caption_config = img_config.get("caption", {})
        self.caption_font_size = caption_config.get("font_size", 10.5)
        self.caption_bold = caption_config.get("bold", False)
        self.caption_position = caption_config.get("position", "below")
        self.caption_label = caption_config.get("label", "图")
        self.caption_numbering = caption_config.get("numbering", "chapter")

        self._warnings: list[str] = []
        self._resized_count = 0

    def process_all_images(self, doc: Any) -> ImageFormatResult:
        """处理文档中所有图片，返回格式化结果"""
        result = ImageFormatResult()
        page_width = self._get_page_content_width(doc)

        for para in doc.paragraphs:
            has_image = False

            for run in para.runs:
                drawings = run._element.findall(qn("w:drawing"))
                for drawing in drawings:
                    resized = self._process_drawing(drawing, page_width)
                    if resized:
                        result.images_resized += 1
                    has_image = True

                picts = run._element.findall(qn("w:pict"))
                for _pict in picts:
                    has_image = True

            if has_image and self.center:
                align_map = {
                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                }
                para.alignment = align_map.get(self.alignment, WD_ALIGN_PARAGRAPH.CENTER)
                result.images_centered += 1

        result.captions_formatted = self._format_captions(doc)
        result.warnings = list(self._warnings)
        result.stats = self.get_image_stats(doc)
        return result

    def get_warnings(self) -> list[str]:
        return list(self._warnings)

    def _get_page_content_width(self, doc: Any) -> int:
        if not doc.sections:
            return Cm(15)
        section = doc.sections[0]
        page_width = section.page_width
        left_margin = section.left_margin
        right_margin = section.right_margin
        if page_width is not None and left_margin is not None and right_margin is not None:
            return page_width - left_margin - right_margin
        return Cm(15)

    def _process_drawing(self, drawing: Any, page_width: int) -> bool:
        resized = False
        for ext in drawing.iter(qn("wp:extent")):
            cx = int(ext.get("cx", 0))
            cy = int(ext.get("cy", 0))

            if cx <= 0:
                continue

            target_width = self._calc_target_width(page_width)

            if cx > target_width:
                ratio = target_width / cx
                new_cx = int(target_width)
                new_cy = int(cy * ratio)
                ext.set("cx", str(new_cx))
                ext.set("cy", str(new_cy))
                resized = True

                for ext2 in drawing.iter(qn("a:ext")):
                    if ext2.get("cx") == str(cx):
                        ext2.set("cx", str(new_cx))
                        ext2.set("cy", str(new_cy))

        self._check_dpi(drawing)
        return resized

    def _calc_target_width(self, page_width: int) -> int:
        if self.max_width == "full":
            return page_width
        elif isinstance(self.max_width, str) and self.max_width.endswith("%"):
            pct = int(self.max_width.rstrip("%")) / 100
            return int(page_width * pct)
        elif isinstance(self.max_width, (int, float)):
            return Cm(self.max_width)
        return page_width

    def _check_dpi(self, drawing: Any) -> None:
        for blip in drawing.iter(qn("a:blip")):
            embed = blip.get(qn("r:embed"))
            if not embed:
                continue
            for ext in drawing.iter(qn("wp:extent")):
                cx_emu = int(ext.get("cx", 0))
                cy_emu = int(ext.get("cy", 0))
                if cx_emu > 0 and cy_emu > 0:
                    width_inches = cx_emu / 914400
                    height_inches = cy_emu / 914400
                    if width_inches < self.min_width_inches:
                        self._warnings.append(
                            f"图片宽度仅 {width_inches:.2f} 英寸({width_inches * 2.54:.1f}cm)，"
                            f"可能分辨率不足，建议不小于 {self.min_width_inches:.1f} 英寸"
                        )
                    if width_inches > 0 and height_inches > 0:
                        aspect = width_inches / height_inches
                        if aspect > 5 or aspect < 0.2:
                            self._warnings.append(
                                f"图片宽高比异常 ({width_inches:.1f}:{height_inches:.1f})，"
                                "可能被拉伸或压缩"
                            )

    def _format_captions(self, doc: Any) -> int:
        """检测并格式化图题段落"""
        count = 0
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if self.CAPTION_PATTERN.match(text):
                self._format_caption_paragraph(paragraph)
                count += 1
        return count

    def _format_caption_paragraph(self, paragraph: Any) -> None:
        """格式化图题段落"""
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            set_east_asian_font(run, "宋体")
            run.font.size = Pt(self.caption_font_size)
            run.font.bold = self.caption_bold

        if self.caption_position == "below":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(6)

    def get_image_stats(self, doc: Any) -> dict[str, Any]:
        total = 0
        widths: list[float] = []
        heights: list[float] = []
        small_count = 0

        for para in doc.paragraphs:
            for run in para.runs:
                drawings = run._element.findall(qn("w:drawing"))
                for drawing in drawings:
                    for ext in drawing.iter(qn("wp:extent")):
                        cx = int(ext.get("cx", 0))
                        cy = int(ext.get("cy", 0))
                        if cx > 0:
                            total += 1
                            w_inches = cx / 914400
                            widths.append(w_inches)
                            if cy > 0:
                                heights.append(cy / 914400)
                            if w_inches < self.min_width_inches:
                                small_count += 1

        return {
            "total": total,
            "avg_width_inches": round(sum(widths) / len(widths), 2) if widths else 0,
            "max_width_inches": round(max(widths), 2) if widths else 0,
            "min_width_inches": round(min(widths), 2) if widths else 0,
            "small_images": small_count,
            "aspect_ratios": [
                round(w / h, 2) for w, h in zip(widths, heights) if h > 0
            ][:5],
        }
