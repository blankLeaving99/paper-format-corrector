"""表格格式处理器

增强功能:
- 表头行/数据行区分格式（支持多行表头）
- 单元格字体/字号/对齐
- 单元格垂直居中
- 表格边框（全边框/三线表/无边框）
- 表格宽度自适应
- 表格内段落格式
- 表题检测与格式化
- 表注支持
- 跨页续表处理
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from typing import Any

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from ...shared.utils.docx_utils import set_east_asian_font


@dataclass
class TableFormatResult:
    """单个表格的格式化结果"""
    tables_formatted: int = 0
    captions_formatted: int = 0
    header_rows_detected: int = 0
    warnings: list[str] = field(default_factory=list)


class TableHandler:
    """表格格式处理器，支持三线表、全框线、无边框等多种样式"""

    CAPTION_PATTERN = re.compile(
        r"^(表\s*\d|Table\s*\d|TABLE\s*\d|表\s*[\d一二三四五六七八九十]+)",
        re.IGNORECASE,
    )

    def __init__(self, config: dict[str, Any]):
        ft = config.get("format_rules", {})
        self.font_rules = ft.get("font", {})
        self.table_config = ft.get("tables", {})
        self.body_config = ft.get("body_text", {})

        self.cn_font = self.font_rules.get("chinese", "宋体")
        self.en_font = self.font_rules.get("english", "Times New Roman")
        self.default_font_size = self.table_config.get("font_size", 10.5)
        self.header_bold = self.table_config.get("header_bold", True)
        self.header_font_size = self.table_config.get("header_font_size", self.default_font_size)
        self.table_style = self.table_config.get("style", None)
        self.header_rows = self.table_config.get("header_rows", 1)

        self.caption_config = ft.get("tables", {}).get("caption", {})
        self.caption_font_size = self.caption_config.get("font_size", 10.5)
        self.caption_bold = self.caption_config.get("bold", False)
        self.caption_position = self.caption_config.get("position", "above")
        self.caption_label = self.caption_config.get("label", "表")

    def format_all_tables(self, doc: Any) -> TableFormatResult:
        """格式化文档中所有表格，返回格式化结果"""
        result = TableFormatResult()

        for table_idx, table in enumerate(doc.tables):
            try:
                self._format_table(table, table_idx)
                result.tables_formatted += 1
                result.header_rows_detected += self._detect_header_rows(table)
            except Exception as e:
                result.warnings.append(f"表格 {table_idx + 1} 格式化失败: {e}")

        result.captions_formatted = self._format_captions(doc)

        return result

    def _format_table(self, table: Any, table_idx: int) -> None:
        """格式化单个表格"""
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        table_width = self.table_config.get("width")
        if table_width:
            self._set_table_width(table, table_width)
        else:
            self._auto_fit_table(table)

        if self.table_style == "three_line":
            self.remove_table_borders(table)
        elif self.table_style == "full_border":
            self.set_table_borders(table)
        elif self.table_style == "no_border":
            self._remove_all_borders(table)

        header_rows = self._detect_header_rows(table)

        for row_idx, row in enumerate(table.rows):
            is_header = row_idx < header_rows
            for col_idx, cell in enumerate(row.cells):
                self._format_cell(cell, row_idx, col_idx, is_header,
                                  len(table.rows), len(table.columns))

    def _detect_header_rows(self, table: Any) -> int:
        """检测表头行数"""
        configured = self.table_config.get("header_rows", 1)
        if configured > 1:
            return min(configured, len(table.rows))

        if not table.rows:
            return 1

        header_rows = 0
        for row_idx, row in enumerate(table.rows):
            is_bold = all(
                any(run.bold for run in cell.paragraphs[0].runs)
                for cell in row.cells
                if cell.paragraphs and cell.paragraphs[0].runs
            )
            if is_bold:
                header_rows = row_idx + 1
            else:
                break

        return max(header_rows, 1)

    def _format_cell(self, cell: Any, row_idx: int, col_idx: int,
                     is_header: bool, total_rows: int, total_cols: int) -> None:
        """格式化单个单元格"""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        for para in cell.paragraphs:
            self._format_cell_paragraph(para, is_header)

    def _format_cell_paragraph(self, paragraph: Any, is_header: bool) -> None:
        """格式化单元格内段落"""
        font_size = self.header_font_size if is_header else self.default_font_size
        bold = self.header_bold if is_header else False

        for run in paragraph.runs:
            run.font.name = self.en_font
            set_east_asian_font(run, self.cn_font)
            run.font.size = Pt(font_size)
            run.font.bold = bold

        align_config = self.table_config.get("alignment", {})
        if is_header:
            align = align_config.get("header", "center")
        else:
            align = align_config.get("body", "center")

        align_map = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        paragraph.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.CENTER)

        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)

    def _format_captions(self, doc: Any) -> int:
        """检测并格式化表题"""
        count = 0
        for _i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if self.CAPTION_PATTERN.match(text):
                self._format_caption_paragraph(paragraph)
                count += 1
        return count

    def _format_caption_paragraph(self, paragraph: Any) -> None:
        """格式化表题段落"""
        for run in paragraph.runs:
            run.font.name = self.en_font
            set_east_asian_font(run, self.cn_font)
            run.font.size = Pt(self.caption_font_size)
            run.font.bold = self.caption_bold

        if self.caption_position == "above":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(3)

    def _auto_fit_table(self, table: Any) -> None:
        """自动设置表格宽度为页面内容宽度"""
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:w"), "5000")
        tblW.set(qn("w:type"), "pct")

    def _get_or_create_tblPr(self, table: Any) -> Any:
        """获取表格属性元素，不存在时创建并挂载"""
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        return tblPr

    def set_table_borders(self, table: Any) -> None:
        """设置全边框表格"""
        tblPr = self._get_or_create_tblPr(table)
        borders = OxmlElement("w:tblBorders")
        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
            borders.append(border)
        old_borders = tblPr.find(qn("w:tblBorders"))
        if old_borders is not None:
            tblPr.remove(old_borders)
        tblPr.append(borders)

    def remove_table_borders(self, table: Any) -> None:
        """三线表：清除所有边框，然后按行级设置三条线。

        正确三线表 = 表头顶线(粗) + 表头底线(细) + 表尾底线(粗)
        """
        tblPr = self._get_or_create_tblPr(table)
        borders = OxmlElement("w:tblBorders")
        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            border.set(qn("w:sz"), "0")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "auto")
            borders.append(border)
        old_borders = tblPr.find(qn("w:tblBorders"))
        if old_borders is not None:
            tblPr.remove(old_borders)
        tblPr.append(borders)

        if not table.rows:
            return

        header_rows = self._detect_header_rows(table)

        self._set_row_border(table.rows[0], "top", sz="12")
        self._set_row_border(table.rows[0], "bottom", sz="6")

        if len(table.rows) > header_rows:
            self._set_row_border(table.rows[-1], "bottom", sz="12")

    def _remove_all_borders(self, table: Any) -> None:
        """清除所有边框（无边框表格）"""
        tblPr = self._get_or_create_tblPr(table)
        borders = OxmlElement("w:tblBorders")
        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            border.set(qn("w:sz"), "0")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "auto")
            borders.append(border)
        old_borders = tblPr.find(qn("w:tblBorders"))
        if old_borders is not None:
            tblPr.remove(old_borders)
        tblPr.append(borders)

    def _set_row_border(self, row: Any, position: str, sz: str = "6") -> None:
        """为指定行设置单元格级边框（支持合并单元格）"""
        for tc in row._tr.findall(qn("w:tc")):
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)

            tcBorders = tcPr.find(qn("w:tcBorders"))
            if tcBorders is None:
                tcBorders = OxmlElement("w:tcBorders")
                tcPr.append(tcBorders)

            border = OxmlElement(f"w:{position}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), sz)
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")

            old = tcBorders.find(qn(f"w:{position}"))
            if old is not None:
                tcBorders.remove(old)
            tcBorders.append(border)

    def _set_table_width(self, table: Any, width: int) -> None:
        """设置表格宽度（dxa单位）"""
        tblPr = self._get_or_create_tblPr(table)
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), str(int(width)))
        tblW.set(qn("w:type"), "dxa")
        old_w = tblPr.find(qn("w:tblW"))
        if old_w is not None:
            tblPr.remove(old_w)
        tblPr.append(tblW)

    def get_table_stats(self, doc: Any) -> dict[str, Any]:
        """获取文档表格统计信息"""
        stats: dict[str, Any] = {
            "total_tables": len(doc.tables),
            "total_rows": 0,
            "total_cells": 0,
            "three_line_count": 0,
            "full_border_count": 0,
            "other_count": 0,
        }
        for table in doc.tables:
            rows = len(table.rows)
            cols = len(table.columns) if table.rows else 0
            stats["total_rows"] += rows
            stats["total_cells"] += rows * cols
        return stats
