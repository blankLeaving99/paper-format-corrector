import logging
import os
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt

from ...shared.utils.docx_utils import set_east_asian_font
from ..handlers.figure_table_handler import FigureTableHandler
from ..handlers.header_footer_handler import HeaderFooterHandler
from ..handlers.image_handler import ImageHandler
from ..handlers.table_handler import TableHandler
from ..handlers.toc_handler import TOCHandler
from ..parsers.cross_reference import CrossReferenceUpdater
from ..parsers.reference_formatter import ReferenceFormatter
from ..parsers.section_parser import (
    SectionDetector,
    SectionType,
    detect_document_language,
    map_override_to_section_type,
)

logger = logging.getLogger(__name__)


class FormatCorrector:
    """论文格式矫正器（完整版）"""

    def __init__(self, template_path, config, type_overrides: dict[str, str] | None = None):
        if template_path and Path(template_path).is_file():
            self.template = Document(template_path)
            self.template_path = template_path
        else:
            logger.warning(f"模板文件不存在 ({template_path})，使用默认空白模板")
            self.template = Document()
            self.template_path = None
        self.config = config
        self._runtime_font_overrides = None  # 语言检测的运行时字体覆盖
        self._type_overrides = type_overrides or {}  # 手动段落类型修正

        # 从模板提取样式（作为配置的回退）
        self.template_styles = {}
        self.template_margins = {}
        if self.template_path:
            try:
                from .style_extractor import StyleExtractor
                extractor = StyleExtractor(self.template_path)
                self.template_styles = extractor.extract_all_styles()
                self.template_margins = extractor.extract_page_margins()
            except Exception as e:
                logger.warning(f"模板样式提取失败: {e}")

        # 子模块
        self.section_detector = SectionDetector(config)
        self.ref_formatter = ReferenceFormatter(config)
        self.fig_table_handler = FigureTableHandler(config)
        self.table_handler = TableHandler(config)
        self.image_handler = ImageHandler(config)
        self.hf_handler = HeaderFooterHandler(config)
        self.toc_handler = TOCHandler(config)
        self.cross_ref_updater = CrossReferenceUpdater()

        # 处理报告
        self.report = self._empty_report()
        self._current_chapter = 0

    def _empty_report(self):
        return {
            "paragraphs_corrected": 0,
            "headings_fixed": 0,
            "body_fixed": 0,
            "figures_renumbered": 0,
            "tables_renumbered": 0,
            "tables_formatted": 0,
            "images_centered": 0,
            "ref_issues": [],
            "fig_table_issues": [],
            "warnings": [],
            "cross_refs_updated": 0,
        }

    def _get_font_rules(self):
        """获取字体规则，运行时覆盖优先于 config。

        Returns:
            dict: 合并后的字体规则字典
        """
        base = self.config.get("format_rules", {}).get("font", {})
        if self._runtime_font_overrides:
            merged = dict(base)
            merged.update(self._runtime_font_overrides)
            return merged
        return base

    def correct_document(self, input_path, output_path, backup=True):
        """矫正文档格式

        Args:
            input_path: 输入文档路径
            output_path: 输出文档路径
            backup: 是否在修改前备份原始文件

        Returns:
            矫正报告字典
        """
        logger.info(f"正在处理: {input_path}")

        # 修改前备份
        backup_path = None
        if backup:
            import tempfile
            output_dir = Path(output_path).parent
            output_dir.mkdir(parents=True, exist_ok=True)
            backup_fd, backup_path = tempfile.mkstemp(suffix=".backup.docx", dir=str(output_dir))
            os.close(backup_fd)
            shutil.copy2(input_path, backup_path)
            logger.info(f"已备份原始文件: {backup_path}")

        doc = Document(input_path)

        # 重置状态
        self.section_detector.reset()
        self.fig_table_handler.reset()
        self.report = self._empty_report()
        self.report["backup_path"] = backup_path

        # 检测文档语言并调整字体配置
        self._detect_and_apply_language(doc)

        # 1. 页面设置
        self._apply_page_setup(doc)

        # 2. 段落格式矫正（核心循环）
        self._correct_paragraphs(doc)

        # 3. 表格格式矫正
        table_result = self.table_handler.format_all_tables(doc)
        self.report["tables_formatted"] = table_result.tables_formatted if hasattr(table_result, 'tables_formatted') else table_result

        # 4. 图片处理（居中 + 调整大小 + DPI检查）
        image_result = self.image_handler.process_all_images(doc)
        self.report["images_centered"] = image_result.images_centered if hasattr(image_result, 'images_centered') else image_result
        if hasattr(image_result, 'warnings') and image_result.warnings:
            self.report["warnings"].extend(image_result.warnings)

        # 5. 参考文献格式化
        self._format_references(doc)

        # 6. 插入目录（如果需要）
        if self.toc_handler.enabled and not self.toc_handler.has_toc(doc):
            self.toc_handler.insert_toc(doc)

        # 7. 页眉页脚（传入章名映射）
        chapter_map = self._build_section_chapter_map(doc)
        self.hf_handler.apply(doc, chapter_map=chapter_map)

        # 8. 公式编号
        eq_renumber_map = self._renumber_formulas(doc)

        # 9. 脚注格式矫正
        self._correct_footnotes(doc)

        # 10. 交叉引用更新
        fig_map = getattr(self.fig_table_handler, 'fig_renumber_map', {})
        tab_map = getattr(self.fig_table_handler, 'tab_renumber_map', {})
        self._update_cross_references(doc, fig_map, tab_map, eq_renumber_map)

        doc.save(output_path)
        logger.info(f"已保存: {output_path}")
        return self.report

    def _detect_and_apply_language(self, doc):
        """检测文档语言并调整字体配置（不修改 self.config）"""
        lang_config = self.config.get("language", {})
        primary = lang_config.get("primary", "auto")

        if primary == "auto":
            primary = detect_document_language(doc)
            logger.info(f"  检测到文档语言: {primary}")

        # 根据语言设置运行时字体覆盖（不修改 config，避免污染后续文档）
        fonts = lang_config.get("fonts", {})
        lang_fonts = fonts.get(primary, {})

        if lang_fonts:
            self._runtime_font_overrides = {}
            if "body" in lang_fonts:
                self._runtime_font_overrides["chinese"] = lang_fonts["body"]
            if "heading" in lang_fonts:
                self._runtime_font_overrides["heading_chinese"] = lang_fonts["heading"]
            if "english_in_chinese" in lang_fonts:
                self._runtime_font_overrides["english"] = lang_fonts["english_in_chinese"]
        else:
            self._runtime_font_overrides = None

        self._detected_language = primary

    def _apply_page_setup(self, doc):
        margins = self.config.get("format_rules", {}).get("margins", {})
        # 回退到模板边距
        if not margins and self.template_margins:
            margins = self.template_margins
        if not margins:
            return
        for section in doc.sections:
            section.top_margin = Cm(margins.get("top", 2.54))
            section.bottom_margin = Cm(margins.get("bottom", 2.54))
            section.left_margin = Cm(margins.get("left", 3.17))
            section.right_margin = Cm(margins.get("right", 3.17))

    def _correct_paragraphs(self, doc):  # noqa: C901
        """逐段矫正格式（使用多模块检测管道）"""
        # 多模块检测：各模块独立检测 → 聚合 → 上下文校验
        pipeline_result = self.section_detector.detect_all(doc.paragraphs)

        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue

            section_type = pipeline_result.final_labels[i]
            extra = pipeline_result.final_extras[i]

            # 应用手动类型修正（如果有）
            if self._type_overrides:
                # 获取当前段落的类型字符串（用于匹配覆盖规则）
                current_type_str = section_type.name.lower()
                # 特殊映射：CHAPTER -> heading1, SECTION -> heading2, SUBSECTION -> heading3
                type_str_to_check = current_type_str
                if section_type == SectionType.CHAPTER:
                    type_str_to_check = "heading1"
                elif section_type == SectionType.SECTION:
                    type_str_to_check = "heading2"
                elif section_type == SectionType.SUBSECTION:
                    type_str_to_check = "heading3"
                elif section_type == SectionType.REFERENCE_ITEM:
                    type_str_to_check = "reference"
                elif section_type == SectionType.FORMULA_CONTENT:
                    type_str_to_check = "formula"

                if type_str_to_check in self._type_overrides:
                    override_val = self._type_overrides[type_str_to_check]
                    override_section_type = map_override_to_section_type(override_val)
                    if override_section_type and override_section_type != section_type:
                        section_type = override_section_type

            # 根据类型应用不同格式
            if section_type == SectionType.TITLE:
                self._apply_title_style(paragraph)
            elif section_type == SectionType.AUTHORS:
                self._apply_author_style(paragraph)
            elif section_type == SectionType.AFFILIATION:
                self._apply_affiliation_style(paragraph)
            elif section_type in (SectionType.ABSTRACT_CN, SectionType.ABSTRACT_EN):
                self._apply_abstract_style(paragraph, section_type)
            elif section_type in (SectionType.KEYWORDS_CN, SectionType.KEYWORDS_EN):
                self._apply_keywords_style(paragraph, section_type)
            elif section_type == SectionType.CHAPTER:
                self._apply_heading_style(paragraph, "heading1")
                chapter_num = extra.get("chapter_num", 0)
                self.fig_table_handler.update_chapter(chapter_num)
                self._current_chapter = chapter_num
                self.report["headings_fixed"] += 1
            elif section_type == SectionType.SECTION:
                self._apply_heading_style(paragraph, "heading2")
                self.report["headings_fixed"] += 1
            elif section_type == SectionType.SUBSECTION:
                self._apply_heading_style(paragraph, "heading3")
                self.report["headings_fixed"] += 1
            elif section_type == SectionType.REFERENCE_TITLE:
                self._apply_heading_style(paragraph, "heading1")
                self.report["headings_fixed"] += 1
            elif section_type in (SectionType.ACKNOWLEDGMENT_TITLE, SectionType.APPENDIX_TITLE):
                self._apply_heading_style(paragraph, "heading1")
                self.report["headings_fixed"] += 1
            elif section_type == SectionType.CODE:
                self._preserve_code_style(paragraph)
            elif section_type == SectionType.FORMULA_CONTENT:
                self._preserve_formula_content_style(paragraph)
            elif section_type == SectionType.FIGURE_CAPTION:
                self.fig_table_handler.process_paragraph(paragraph, section_type, extra)
            elif section_type == SectionType.TABLE_CAPTION:
                self.fig_table_handler.process_paragraph(paragraph, section_type, extra)
            elif section_type == SectionType.BODY:
                self._apply_body_style(paragraph)
                self.report["body_fixed"] += 1
            elif section_type == SectionType.REFERENCE_ITEM:
                pass  # 参考文献条目不计入矫正数

            # 统计实际处理的段落数（排除 REFERENCE_ITEM 和未识别类型）
            if section_type not in (SectionType.REFERENCE_ITEM, SectionType.BLANK, SectionType.UNKNOWN):
                self.report["paragraphs_corrected"] += 1

        self.report["fig_table_issues"] = self.fig_table_handler.get_issues()

    def _format_references(self, doc):
        ref_start = None
        ref_keywords = self.config.get("auto_detect", {}).get(
            "reference_keywords", ["参考文献"]
        )
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            # 模糊匹配：支持 "参考文献"、"参考文献[1]"、"参考文献：" 等变体
            if any(text.startswith(kw) for kw in ref_keywords):
                ref_start = i
                break
        if ref_start is not None:
            # 检测引用风格
            citation_style = self.ref_formatter.detect_citation_style(doc, ref_start)
            self.report["citation_style"] = self.ref_formatter.get_citation_style_name(citation_style)

            self.ref_formatter.format_references(doc, ref_start)
            self.report["ref_issues"] = self.ref_formatter.validate_references(doc, ref_start)
            # 引用一致性检查
            consistency_issues = self.ref_formatter.check_citation_consistency(doc, ref_start)
            for issue in consistency_issues:
                self.report["ref_issues"].append(issue["message"])

    def _build_section_chapter_map(self, doc):
        """将段落级章名映射转为 section 级映射。

        通过遍历段落的 <w:sectPr>（节结束标记）精确确定每个 section 的边界。

        Returns:
            dict[int, str]: {section_index: chapter_name}
        """
        from docx.oxml.ns import qn as _qn

        para_chapters = self.section_detector.get_chapter_map()
        if not para_chapters:
            return {}

        section_map = {}
        current_section = 0
        last_chapter = None

        for i, para in enumerate(doc.paragraphs):
            # 更新当前章名
            if i in para_chapters:
                last_chapter = para_chapters[i]

            # 检查此段落是否包含 sectPr（节结束标记）
            pPr = para._element.find(_qn("w:pPr"))
            if pPr is not None and pPr.find(_qn("w:sectPr")) is not None:
                section_map[current_section] = last_chapter
                current_section += 1

        # 最后一个 section（body 级别的 sectPr）
        section_map[current_section] = last_chapter

        return section_map

    def _renumber_formulas(self, doc):  # noqa: C901
        """重编号公式并返回编号映射

        Returns:
            dict: {old_num: new_num} 映射，用于交叉引用更新
        """
        formula_config = self.config.get("format_rules", {}).get("formulas", {})
        if not formula_config.get("numbering", True):
            return {}

        formula_pattern = re.compile(
            self.config.get("auto_detect", {}).get("formula_pattern", r"^\(?\d+[-\.]\d+\)?$")
        )

        numbering = formula_config.get("numbering_style", "chapter")  # chapter | sequential
        separator = formula_config.get("separator", "-")

        renumber_map = {}
        chapter_counter = {}  # {chapter_num: formula_count}
        global_counter = 0
        current_chapter = 0

        # 先检测当前章节（从 section_detector 的状态获取）
        if hasattr(self, '_current_chapter'):
            current_chapter = self._current_chapter

        for para in doc.paragraphs:
            text = para.text.strip()
            if not formula_pattern.match(text):
                continue

            # 提取旧编号
            old_num_match = re.search(r"\d+(?:[-\.]\d+)?", text)
            if not old_num_match:
                continue
            old_num = old_num_match.group(0)

            # 生成新编号
            if numbering == "chapter":
                if current_chapter not in chapter_counter:
                    chapter_counter[current_chapter] = 0
                chapter_counter[current_chapter] += 1
                new_num = f"{current_chapter}{separator}{chapter_counter[current_chapter]}"
            else:
                global_counter += 1
                new_num = str(global_counter)

            # 应用样式
            self._apply_formula_style(para, formula_config)

            # 如果编号需要修正，重写文本（保留 run 级格式）
            clean_old = old_num.strip("()")
            if clean_old != new_num:
                # 保留原始格式（括号等）
                if text.startswith("(") and text.endswith(")"):
                    new_text = f"({new_num})"
                else:
                    new_text = new_num

                # 优先：找到包含旧编号的 run，只替换那个 run（保留内联格式）
                replaced = False
                for run in para.runs:
                    if clean_old in run.text or old_num in run.text:
                        run.text = run.text.replace(old_num, new_text).replace(clean_old, new_text)
                        replaced = True
                        break

                # fallback：如果没找到，清空所有 run，用第一个 run 写入
                if not replaced:
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = new_text
                    else:
                        para.add_run(new_text)

                renumber_map[clean_old] = new_num

        return renumber_map

    def _correct_footnotes(self, doc):  # noqa: C901
        """矫正脚注格式"""
        from docx.oxml.ns import qn as _qn

        # 查找文档中的脚注部分
        footnotes_part = None
        try:
            footnotes_part = doc.part.notes_part
        except Exception:
            pass

        if footnotes_part is None:
            # 尝试通过 XML 查找
            footnotes_elements = doc.element.findall(_qn('w:footnotes'))
            if not footnotes_elements:
                return

            for footnotes_elem in footnotes_elements:
                for footnote in footnotes_elem.findall(_qn('w:footnote')):
                    # 跳过分隔符等特殊脚注
                    if footnote.get(_qn('w:type')) in ('separator', 'continuationSeparator'):
                        continue
                    for p_elem in footnote.findall(_qn('w:p')):
                        for r_elem in p_elem.findall(_qn('w:r')):
                            rpr = r_elem.find(_qn('w:rPr'))
                            if rpr is None:
                                rpr = OxmlElement('w:rPr')
                                r_elem.insert(0, rpr)
                            # 设置脚注字号（小五号 = 9pt）
                            sz = rpr.find(_qn('w:sz'))
                            if sz is None:
                                sz = OxmlElement('w:sz')
                                rpr.append(sz)
                            sz.set(_qn('w:val'), '18')  # 9pt = 18 half-points
        else:
            # 通过 notes_part 处理
            for footnote in footnotes_part.footnotes:
                for para in footnote.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)  # 小五号

    def _update_cross_references(self, doc, fig_map, tab_map, eq_map):
        """更新交叉引用

        Args:
            doc: Document 对象
            fig_map: 图编号映射 {old: new}
            tab_map: 表编号映射 {old: new}
            eq_map: 公式编号映射 {old: new}
        """
        if not fig_map and not tab_map and not eq_map:
            return

        updated = self.cross_ref_updater.update(doc, fig_map, tab_map, eq_map)
        if updated > 0:
            self.report["cross_refs_updated"] = updated
            logger.info(f"  交叉引用更新: {updated} 个段落")

    def _apply_formula_style(self, paragraph, config):
        style_rules = {"font_size": config.get("font_size", 12)}
        font_rules = self._get_font_rules()
        for run in paragraph.runs:
            self._set_run_font(run, font_rules, style_rules)
        if config.get("numbering_position") == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _preserve_code_style(self, paragraph):
        """保留代码段落的原始格式，不做字体/行距矫正"""
        code_config = self.config.get("format_rules", {}).get("code", {})

        if code_config.get("force_mono", False):
            mono_font = code_config.get("mono_font", "Consolas")
            mono_size = code_config.get("mono_font_size", 10)
            for run in paragraph.runs:
                run.font.name = mono_font
                run.font.size = Pt(mono_size)

    def _preserve_formula_content_style(self, paragraph):
        """保留公式内容的原始格式，不做字体矫正"""
        formula_config = self.config.get("format_rules", {}).get("formulas", {})
        if formula_config.get("center", False):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ========== 各类段落样式应用 ==========

    def _apply_title_style(self, paragraph):
        rules = self.config.get("format_rules", {})
        tp = rules.get("title_page", {})
        font_rules = rules.get("font", {})
        for run in paragraph.runs:
            run.font.name = font_rules.get("english", "Times New Roman")
            set_east_asian_font(run, font_rules.get("heading_chinese", "黑体"))
            run.font.size = Pt(tp.get("title_font_size", 22))
            run.font.bold = tp.get("title_bold", True)
        align_map = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        paragraph.alignment = align_map.get(tp.get("title_align", "center"), WD_ALIGN_PARAGRAPH.CENTER)
        paragraph.paragraph_format.space_before = Pt(24)
        paragraph.paragraph_format.space_after = Pt(18)
        paragraph.paragraph_format.line_spacing = 1.5

    def _apply_author_style(self, paragraph):
        rules = self.config.get("format_rules", {})
        tp = rules.get("title_page", {})
        font_rules = rules.get("font", {})
        for run in paragraph.runs:
            run.font.name = font_rules.get("english", "Times New Roman")
            set_east_asian_font(run, font_rules.get("chinese", "宋体"))
            run.font.size = Pt(tp.get("author_font_size", 12))
            run.font.bold = False
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)

    def _apply_affiliation_style(self, paragraph):
        rules = self.config.get("format_rules", {})
        tp = rules.get("title_page", {})
        font_rules = rules.get("font", {})
        for run in paragraph.runs:
            run.font.name = font_rules.get("english", "Times New Roman")
            set_east_asian_font(run, font_rules.get("chinese", "宋体"))
            run.font.size = Pt(tp.get("affiliation_font_size", 10.5))
            run.font.bold = False
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _apply_abstract_style(self, paragraph, section_type):
        rules = self.config.get("format_rules", {})
        abs_config = rules.get("abstract", {})
        font_rules = rules.get("font", {})
        text = paragraph.text.strip()

        abstract_label_cn = self.config.get("auto_detect", {}).get("abstract_pattern", r"^摘\s*要$")
        abstract_label_en = self.config.get("auto_detect", {}).get("abstract_en_pattern", r"^Abstract$|^ABSTRACT$")

        is_title_line = re.match(abstract_label_cn, text) or re.match(abstract_label_en, text)

        if is_title_line:
            for run in paragraph.runs:
                run.font.name = font_rules.get("english", "Times New Roman")
                set_east_asian_font(run, font_rules.get("heading_chinese", "黑体"))
                run.font.size = Pt(abs_config.get("title_font_size", 16))
                run.font.bold = abs_config.get("title_bold", True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            for run in paragraph.runs:
                run.font.name = font_rules.get("english", "Times New Roman")
                set_east_asian_font(run, font_rules.get("chinese", "宋体"))
                run.font.size = Pt(abs_config.get("content_font_size", 12))
                run.font.bold = False
            self._apply_line_spacing(paragraph, abs_config.get("content_line_spacing", 1.5))
            indent = abs_config.get("content_first_line_indent", 0)
            if indent:
                paragraph.paragraph_format.first_line_indent = Pt(indent * 12)

    def _apply_keywords_style(self, paragraph, section_type):
        rules = self.config.get("format_rules", {})
        kw_config = rules.get("keywords", {})
        font_rules = rules.get("font", {})
        text = paragraph.text.strip()

        if "：" in text or ":" in text:
            sep_pos = max(text.find("："), text.find(":"))
            label_part = text[:sep_pos + 1]
            content_part = text[sep_pos + 1:]
        else:
            label_part = ""
            content_part = text

        for run in paragraph.runs:
            run.text = ""

        if label_part and paragraph.runs:
            paragraph.runs[0].text = label_part
            paragraph.runs[0].font.bold = kw_config.get("bold_label", True)
            paragraph.runs[0].font.size = Pt(kw_config.get("font_size", 12))
            paragraph.runs[0].font.name = font_rules.get("english", "Times New Roman")
            set_east_asian_font(paragraph.runs[0], font_rules.get("chinese", "宋体"))
            if content_part:
                content_run = paragraph.add_run(content_part)
                content_run.font.bold = False
                content_run.font.size = Pt(kw_config.get("font_size", 12))
                content_run.font.name = font_rules.get("english", "Times New Roman")
                set_east_asian_font(content_run, font_rules.get("chinese", "宋体"))
        else:
            for run in paragraph.runs:
                run.font.name = font_rules.get("english", "Times New Roman")
                set_east_asian_font(run, font_rules.get("chinese", "宋体"))
                run.font.size = Pt(kw_config.get("font_size", 12))

        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)

    def _apply_heading_style(self, paragraph, style_name):
        rules = self.config.get("format_rules", {})
        heading_rules = rules.get("headings", {}).get(style_name, {})
        font_rules = rules.get("font", {})

        for run in paragraph.runs:
            self._set_run_font(run, font_rules, heading_rules)
            set_east_asian_font(run, font_rules.get("heading_chinese", "黑体"))

        align_map = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        align = heading_rules.get("align", "left")
        paragraph.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)

        sb = heading_rules.get("space_before", 0)
        sa = heading_rules.get("space_after", 0)
        paragraph.paragraph_format.space_before = Pt(sb)
        paragraph.paragraph_format.space_after = Pt(sa)

        if "line_spacing" in heading_rules:
            self._apply_line_spacing(paragraph, heading_rules["line_spacing"])

    def _apply_body_style(self, paragraph):
        rules = self.config.get("format_rules", {})
        body_rules = rules.get("body_text", {})
        font_rules = rules.get("font", {})

        self._apply_mixed_font(paragraph, font_rules, body_rules)

        self._apply_line_spacing(paragraph, body_rules.get("line_spacing", 1.5))

        align = body_rules.get("align", "justify")
        if align == "justify":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        indent_chars = body_rules.get("first_line_indent", 0)
        if indent_chars:
            paragraph.paragraph_format.first_line_indent = Pt(
                indent_chars * body_rules.get("font_size", 12)
            )

        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

    def _apply_line_spacing(self, paragraph, spacing_value):
        """应用行距（支持多种格式）"""
        if isinstance(spacing_value, dict):
            stype = spacing_value.get("type", "multiple")
            value = spacing_value.get("value", 1.5)
            if stype == "multiple":
                paragraph.paragraph_format.line_spacing = value
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            elif stype == "exact":
                paragraph.paragraph_format.line_spacing = Pt(value)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            elif stype == "atLeast":
                paragraph.paragraph_format.line_spacing = Pt(value)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        else:
            paragraph.paragraph_format.line_spacing = spacing_value

    def _apply_mixed_font(self, paragraph, font_rules, style_rules):
        cn_font = font_rules.get("chinese", "宋体")
        en_font = font_rules.get("english", "Times New Roman")
        font_size = style_rules.get("font_size", 12)

        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            has_chinese = bool(re.search(r"[一-鿿]", text))
            has_english = bool(re.search(r"[a-zA-Z]", text))

            if has_chinese and has_english:
                run.font.name = en_font
                set_east_asian_font(run, cn_font)
            elif has_chinese:
                run.font.name = cn_font
                set_east_asian_font(run, cn_font)
            else:
                run.font.name = en_font

            run.font.size = Pt(font_size)
            if style_rules.get("bold") is not None:
                run.font.bold = style_rules["bold"]

    # ========== 工具方法 ==========

    def _set_run_font(self, run, font_rules, style_rules):
        run.font.name = font_rules.get("english", "Times New Roman")
        set_east_asian_font(run, font_rules.get("chinese", "宋体"))
        if style_rules.get("bold"):
            run.font.bold = True
        if style_rules.get("font_size"):
            run.font.size = Pt(style_rules["font_size"])

    def get_report(self):
        return self.report

    @staticmethod
    def restore_from_backup(backup_path, target_path):
        """从备份恢复原始文件

        Args:
            backup_path: 备份文件路径
            target_path: 恢复目标路径

        Returns:
            True if restored successfully, False otherwise
        """
        backup = Path(backup_path).resolve()
        target = Path(target_path).resolve()

        if not backup.is_file():
            logger.error(f"备份文件不存在: {backup_path}")
            return False

        # 校验文件类型
        if backup.suffix.lower() not in (".docx", ".doc"):
            logger.error(f"备份文件类型不允许: {backup.suffix}")
            return False
        if target.suffix.lower() not in (".docx", ".doc"):
            logger.error(f"目标文件类型不允许: {target.suffix}")
            return False

        # 确保目标目录存在
        target.parent.mkdir(parents=True, exist_ok=True)

        shutil.copy2(str(backup), str(target))
        logger.info(f"已从备份恢复: {target_path}")
        return True
