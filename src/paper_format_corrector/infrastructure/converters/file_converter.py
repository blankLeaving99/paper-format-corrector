"""文件格式转换模块

支持将多种文档格式转换为 .docx，以便后续处理：
- .doc (旧版 Word) → LibreOffice / docx2docx
- .odt (OpenDocument) → LibreOffice
- .rtf (Rich Text) → LibreOffice
- .pdf → 文本提取后生成 .docx
- .txt → 包装为 .docx
- .md → 解析后生成 .docx
"""

import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from ...infrastructure.external_tools import find_libreoffice
from ..path_security import ALLOWED_INPUT_EXTENSIONS, validate_input_path


class FileConverter:
    """多格式文件转换器"""

    # 支持的输入格式
    SUPPORTED_INPUT_FORMATS = (
        ".docx", ".doc", ".odt", ".rtf", ".pdf", ".txt", ".md", ".markdown", ".tex",
    )

    def __init__(self, config=None):
        self.config = config or {}

    @staticmethod
    def is_supported(file_path: str) -> bool:
        """检查文件格式是否支持"""
        suffix = Path(file_path).suffix.lower()
        return suffix in FileConverter.SUPPORTED_INPUT_FORMATS

    @staticmethod
    def needs_conversion(file_path: str) -> bool:
        """检查文件是否需要转换（非 .docx 格式）"""
        suffix = Path(file_path).suffix.lower()
        return suffix != ".docx"

    def convert(self, input_path: str, output_dir: str = None) -> str:  # noqa: C901
        """将文件转换为 .docx 格式

        Args:
            input_path: 输入文件路径
            output_dir: 输出目录，默认为临时目录

        Returns:
            转换后的 .docx 文件路径
        """
        input_path = validate_input_path(input_path, ALLOWED_INPUT_EXTENSIONS)
        suffix = input_path.suffix.lower()

        if suffix == ".docx":
            return str(input_path)

        temp_dir = None
        if output_dir is None:
            temp_dir = tempfile.mkdtemp()
            output_dir = Path(temp_dir)
        else:
            output_dir = Path(output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)

        output_path = output_dir / f"{input_path.stem}.docx"

        try:
            if suffix == ".doc":
                return self._convert_doc(input_path, output_path)
            elif suffix == ".odt":
                return self._convert_odt(input_path, output_path)
            elif suffix == ".rtf":
                return self._convert_rtf(input_path, output_path)
            elif suffix == ".pdf":
                return self._convert_pdf(input_path, output_path)
            elif suffix in (".txt", ".md", ".markdown"):
                return self._convert_text(input_path, output_path, is_markdown=suffix in (".md", ".markdown"))
            elif suffix == ".tex":
                return self._convert_tex(input_path, output_path)
            else:
                raise ValueError(f"不支持的文件格式: {suffix}")
        except Exception:
            # 转换失败时清理临时目录
            if temp_dir:
                shutil.rmtree(temp_dir, ignore_errors=True)
            raise

    # ── .doc 转换 ──────────────────────────────────────────────

    def _convert_doc(self, input_path: Path, output_path: Path) -> str:
        """转换 .doc 为 .docx"""
        # 方法1: LibreOffice
        lo = find_libreoffice()
        if lo:
            return self._libreoffice_convert(input_path, output_path, lo)

        # 方法2: docx2docx
        try:
            from docx2docx import convert
            convert(str(input_path), str(output_path))
            return str(output_path)
        except ImportError:
            pass

        raise RuntimeError(
            "转换 .doc 文件需要 LibreOffice 或 docx2docx。\n"
            "请安装其中之一：\n"
            "  1. 安装 LibreOffice: https://www.libreoffice.org/\n"
            "  2. 安装 docx2docx: pip install docx2docx"
        )

    # ── .odt 转换 ──────────────────────────────────────────────

    def _convert_odt(self, input_path: Path, output_path: Path) -> str:
        """转换 .odt 为 .docx"""
        lo = find_libreoffice()
        if lo:
            return self._libreoffice_convert(input_path, output_path, lo)

        # 尝试用 python-docx 的 ODF 支持（有限）
        try:
            return self._odt_via_odfdo(input_path, output_path)
        except Exception:
            pass

        raise RuntimeError(
            "转换 .odt 文件需要 LibreOffice。\n"
            "请安装 LibreOffice: https://www.libreoffice.org/"
        )

    # ── .rtf 转换 ──────────────────────────────────────────────

    def _convert_rtf(self, input_path: Path, output_path: Path) -> str:
        """转换 .rtf 为 .docx"""
        lo = find_libreoffice()
        if lo:
            return self._libreoffice_convert(input_path, output_path, lo)

        # 尝试解析 RTF 并转换
        try:
            return self._rtf_to_docx(input_path, output_path)
        except Exception:
            pass

        raise RuntimeError(
            "转换 .rtf 文件需要 LibreOffice。\n"
            "请安装 LibreOffice: https://www.libreoffice.org/"
        )

    # ── .pdf 转换 ──────────────────────────────────────────────

    def _convert_pdf(self, input_path: Path, output_path: Path) -> str:
        """从 PDF 提取文本并生成 .docx"""
        text_pages = self._extract_pdf_text(input_path)
        if not text_pages:
            raise RuntimeError("无法从 PDF 中提取文本内容")

        return self._text_to_docx("\n\n".join(text_pages), output_path, source_path=input_path)

    def _extract_pdf_text(self, pdf_path: Path) -> list:
        """从 PDF 提取文本，按页分隔"""
        # 方法1: pdfplumber
        try:
            import pdfplumber
            pages = []
            with pdfplumber.open(str(pdf_path)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text.strip())
            return pages
        except ImportError:
            pass

        # 方法2: PyMuPDF (fitz)
        try:
            import fitz
            doc = fitz.open(str(pdf_path))
            pages = []
            for page in doc:
                text = page.get_text()
                if text.strip():
                    pages.append(text.strip())
            doc.close()
            return pages
        except ImportError:
            pass

        # 方法3: PyPDF2
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(str(pdf_path))
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text and text.strip():
                    pages.append(text.strip())
            return pages
        except ImportError:
            pass

        raise RuntimeError(
            "PDF 文本提取需要安装以下库之一：\n"
            "  pip install pdfplumber\n"
            "  pip install PyMuPDF\n"
            "  pip install PyPDF2"
        )

    # ── .txt / .md 转换 ────────────────────────────────────────

    def _convert_text(self, input_path: Path, output_path: Path, is_markdown: bool = False) -> str:
        """将纯文本或 Markdown 转换为 .docx"""
        # 文件大小限制（50MB）
        max_size = 50 * 1024 * 1024
        file_size = input_path.stat().st_size
        if file_size > max_size:
            raise ValueError(f"文件过大: {file_size / 1024 / 1024:.1f}MB（最大 50MB）")

        # 检测编码
        encoding = self._detect_encoding(input_path)

        with open(input_path, encoding=encoding) as f:
            content = f.read()

        if is_markdown:
            return self._markdown_to_docx(content, output_path)
        else:
            return self._text_to_docx(content, output_path)

    def _text_to_docx(self, text: str, output_path: Path, source_path: Path = None) -> str:
        """将纯文本转换为 .docx"""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt

        # 从 config 读取字体和边距（有默认值）
        font_rules = self.config.get("format_rules", {}).get("font", {})
        body_config = self.config.get("format_rules", {}).get("body_text", {})
        margins = self.config.get("format_rules", {}).get("margins", {})

        doc = Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        font = style.font
        font.name = font_rules.get("chinese", "宋体")
        font.size = Pt(body_config.get("font_size", 12))

        # 设置页面边距
        for section in doc.sections:
            section.top_margin = Cm(margins.get("top", 2.54))
            section.bottom_margin = Cm(margins.get("bottom", 2.54))
            section.left_margin = Cm(margins.get("left", 3.17))
            section.right_margin = Cm(margins.get("right", 3.17))

        # 添加来源信息（如果有）
        if source_path:
            p = doc.add_paragraph()
            run = p.add_run(f"来源文件: {source_path.name}")
            run.font.size = Pt(9)
            run.font.color.rgb = None  # 默认颜色
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 添加分隔线
            doc.add_paragraph("─" * 50)

        # 按段落分割
        paragraphs = text.split("\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                doc.add_paragraph("")
                continue

            # 简单标题检测
            if self._looks_like_heading(para_text):
                doc.add_heading(para_text, level=2)
            else:
                p = doc.add_paragraph(para_text)
                p.paragraph_format.first_line_indent = Cm(0.74)  # 两个字符缩进

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return str(output_path)

    def _markdown_to_docx(self, md_text: str, output_path: Path) -> str:  # noqa: C901
        """将 Markdown 转换为 .docx"""
        from docx import Document
        from docx.shared import Cm, Pt

        # 从 config 读取字体和边距（有默认值）
        font_rules = self.config.get("format_rules", {}).get("font", {})
        body_config = self.config.get("format_rules", {}).get("body_text", {})
        margins = self.config.get("format_rules", {}).get("margins", {})

        doc = Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        font = style.font
        font.name = font_rules.get("chinese", "宋体")
        font.size = Pt(body_config.get("font_size", 12))

        # 设置页面边距
        for section in doc.sections:
            section.top_margin = Cm(margins.get("top", 2.54))
            section.bottom_margin = Cm(margins.get("bottom", 2.54))
            section.left_margin = Cm(margins.get("left", 3.17))
            section.right_margin = Cm(margins.get("right", 3.17))

        lines = md_text.split("\n")
        in_code_block = False
        code_lines = []

        for line in lines:
            stripped = line.strip()

            # 代码块处理
            if stripped.startswith("```"):
                if in_code_block:
                    # 结束代码块
                    code_text = "\n".join(code_lines)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = "Consolas"
                    run.font.size = Pt(10)
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue

            if in_code_block:
                code_lines.append(line)
                continue

            # 空行
            if not stripped:
                doc.add_paragraph("")
                continue

            # 标题
            heading_match = re.match(r"^(#{1,6})\s+(.+)", stripped)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()
                doc.add_heading(text, level=min(level, 4))
                continue

            # 列表项
            list_match = re.match(r"^[-*+]\s+(.+)", stripped)
            if list_match:
                p = doc.add_paragraph(list_match.group(1), style="List Bullet")
                continue

            # 有序列表
            ordered_match = re.match(r"^\d+\.\s+(.+)", stripped)
            if ordered_match:
                p = doc.add_paragraph(ordered_match.group(1), style="List Number")
                continue

            # 引用
            if stripped.startswith("> "):
                p = doc.add_paragraph(stripped[2:])
                p.paragraph_format.left_indent = Cm(1)
                continue

            # 水平线
            if re.match(r"^[-*_]{3,}\s*$", stripped):
                doc.add_paragraph("─" * 50)
                continue

            # 普通段落，处理行内格式
            p = doc.add_paragraph()
            self._add_markdown_runs(p, stripped)
            p.paragraph_format.first_line_indent = Cm(0.74)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return str(output_path)

    def _add_markdown_runs(self, paragraph, text: str):
        """将 Markdown 行内格式添加到段落"""
        from docx.shared import Pt

        # 处理行内格式：**bold**, *italic*, `code`, [link](url)
        pattern = r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|\[(.+?)\]\((.+?)\))"
        last_end = 0

        for match in re.finditer(pattern, text):
            # 添加匹配前的普通文本
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])

            if match.group(2):  # **bold**
                run = paragraph.add_run(match.group(2))
                run.bold = True
            elif match.group(3):  # *italic*
                run = paragraph.add_run(match.group(3))
                run.italic = True
            elif match.group(4):  # `code`
                run = paragraph.add_run(match.group(4))
                run.font.name = "Consolas"
                run.font.size = Pt(10)
            elif match.group(5):  # [text](url)
                run = paragraph.add_run(f"{match.group(5)} ({match.group(6)})")

            last_end = match.end()

        # 添加剩余文本
        if last_end < len(text):
            paragraph.add_run(text[last_end:])

    # ── 辅助方法 ────────────────────────────────────────────────

    def _looks_like_heading(self, text: str) -> bool:
        """简单判断文本是否像标题"""
        # 中文章节标题
        if re.match(r"^第[一二三四五六七八九十百零\d]+[章部分篇]", text):
            return True
        # 数字编号标题：必须是 "1.1 xxx" 或 "1. xxx" 格式且长度适中
        if re.match(r"^\d+\.\d+\.?\s+\S", text) and len(text) < 80:
            return True
        if re.match(r"^\d+\.\s+[A-Z一-鿿ぁ-んァ-ヶ가-힣]", text) and len(text) < 80:
            return True
        # 全大写英文（至少10个字符，避免误匹配缩写）
        if text.isupper() and 10 < len(text) < 100:
            return True
        return False

    @staticmethod
    def _detect_encoding(file_path: Path) -> str:
        """检测文件编码"""
        encodings = ["utf-8-sig", "utf-8", "gbk", "gb18030", "big5", "latin-1"]

        # 读取原始字节用于检测
        raw_bytes = file_path.read_bytes()[:8192]

        for enc in encodings:
            try:
                raw_bytes.decode(enc)
                return enc
            except (UnicodeDecodeError, UnicodeError):
                continue

        return "utf-8"  # 默认回退

    def _libreoffice_convert(self, input_path: Path, output_path: Path, lo_path: str) -> str:
        """使用 LibreOffice 转换文件"""
        output_dir = output_path.parent
        output_dir.mkdir(parents=True, exist_ok=True)

        try:
            result = subprocess.run(
                [
                    lo_path,
                    "--headless",
                    "--convert-to", "docx",
                    "--outdir", str(output_dir),
                    str(input_path),
                ],
                capture_output=True,
                text=True,
                timeout=120,
            )

            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice 转换失败: {result.stderr}")

            # LibreOffice 输出文件名基于输入文件名
            lo_output = output_dir / f"{input_path.stem}.docx"
            if lo_output.exists() and lo_output != output_path:
                lo_output.rename(output_path)

            if not output_path.exists():
                raise RuntimeError("LibreOffice 转换完成但未找到输出文件")

            return str(output_path)

        except subprocess.TimeoutExpired as err:
            raise RuntimeError("LibreOffice 转换超时（120秒）") from err

    def _odt_via_odfdo(self, input_path: Path, output_path: Path) -> str:
        """使用 odfdo 将 ODT 转换为 DOCX（备用方案）"""
        try:
            from odfdo import Document as OdfDocument
        except ImportError:
            raise

        odf_doc = OdfDocument(str(input_path))
        text = odf_doc.get_formatted_text()

        return self._text_to_docx(text, output_path, source_path=input_path)

    def _rtf_to_docx(self, input_path: Path, output_path: Path) -> str:
        """将 RTF 转换为 DOCX（基本实现）"""
        with open(input_path, encoding="latin-1") as f:
            rtf_content = f.read()

        # 简单的 RTF 文本提取
        # 移除 RTF 控制字符
        text = re.sub(r"\\[a-z]+\d*\s?", "", rtf_content)
        text = re.sub(r"[{}]", "", text)
        text = re.sub(r"\\\'[0-9a-f]{2}", "", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = text.strip()

        if not text:
            raise RuntimeError("无法从 RTF 文件中提取文本")

        return self._text_to_docx(text, output_path, source_path=input_path)

    def _convert_tex(self, input_path: Path, output_path: Path) -> str:
        """将 LaTeX (.tex) 转换为 .docx"""
        # 方法1: 使用 pandoc（优先系统 PATH，再试常见路径）
        pandoc = shutil.which("pandoc")
        if not pandoc:
            for p in (r"C:\Program Files\Pandoc\pandoc.exe", "/usr/bin/pandoc", "/usr/local/bin/pandoc"):
                if Path(p).exists():
                    pandoc = p
                    break
        if pandoc:
            try:
                cmd = [
                    pandoc,
                    str(input_path),
                    "-o", str(output_path),
                    "--from=latex",
                    "--to=docx",
                    "--standalone",
                    "--wrap=none",
                ]
                result = subprocess.run(
                    cmd,
                    capture_output=True, text=True, timeout=120,
                )
                if result.returncode == 0 and output_path.exists():
                    return str(output_path)
            except (subprocess.TimeoutExpired, Exception):
                pass

        # 方法2: 手动解析 LaTeX 提取文本
        encoding = self._detect_encoding(input_path)
        with open(input_path, encoding=encoding) as f:
            tex_content = f.read()

        text = self._extract_text_from_latex(tex_content)
        if not text:
            raise RuntimeError("无法从 LaTeX 文件中提取文本内容")

        return self._text_to_docx(text, output_path, source_path=input_path)

    @staticmethod
    def _extract_text_from_latex(tex: str) -> str:
        r"""从 LaTeX 源码中提取文本内容，保留基本格式标记

        提取规则：
        - 移除注释行、preamble、\end{document}
        - 保留 \textbf{} → **bold**、\textit{} → *italic*、\emph{} → *italic*
        - 保留 \texttt{} → `code`
        - 提取 \includegraphics 的 alt 文本占位
        - 保留表格内容（tabular 环境提取为 Markdown 表格）
        - 章节标题映射为 Markdown 格式
        """
        # 移除注释行（行首 %）
        tex = re.sub(r"(?m)^%.*$", "", tex)

        # 移除 preamble（\documentclass 到 \begin{document} 之间）
        tex = re.sub(r"\\documentclass.*?\\begin\{document\}", "", tex, flags=re.DOTALL)

        # 移除 \end{document} 及之后
        tex = re.sub(r"\\end\{document\}.*", "", tex, flags=re.DOTALL)

        # ── 提取表格内容（tabular 环境 → Markdown 表格）──
        def _tabular_to_markdown(match: re.Match) -> str:
            tabular = match.group(1)
            # 提取列对齐信息
            col_match = re.search(r"\{([lcr|]+)\}", tabular)
            col_spec = col_match.group(1) if col_match else ""
            # 按 \\ 分行
            rows = re.split(r"\\\\", tabular)
            md_rows = []
            for row in rows:
                # 移除 \hline、\midrule 等
                row = re.sub(r"\\(?:hline|midrule|toprule|bottomrule|cline\{[^}]*\})", "", row)
                row = row.strip()
                if not row:
                    continue
                cells = [c.strip() for c in row.split("&")]
                md_rows.append("| " + " | ".join(cells) + " |")
            if md_rows and len(md_rows) >= 1:
                # 插入分隔行
                sep_cells = ["---"] * (len(md_rows[0].split("|")) - 2)
                sep = "| " + " | ".join(sep_cells) + " |"
                md_rows.insert(1, sep)
            return "\n\n" + "\n".join(md_rows) + "\n\n" if md_rows else ""

        tex = re.sub(
            r"\\begin\{tabular\}.*?\n(.*?)\\end\{tabular\}",
            _tabular_to_markdown,
            tex,
            flags=re.DOTALL,
        )

        # ── 提取章节标题 ──
        tex = re.sub(r"\\(?:part|chapter)\*?\{(.+?)\}", r"\n\n# \1\n", tex)
        tex = re.sub(r"\\section\*?\{(.+?)\}", r"\n\n## \1\n", tex)
        tex = re.sub(r"\\subsection\*?\{(.+?)\}", r"\n\n### \1\n", tex)
        tex = re.sub(r"\\subsubsection\*?\{(.+?)\}", r"\n\n#### \1\n", tex)
        tex = re.sub(r"\\paragraph\*?\{(.+?)\}", r"\n\n##### \1\n", tex)

        # ── 提取摘要 ──
        tex = re.sub(r"\\begin\{abstract\}", "\n\n**摘要**\n\n", tex)
        tex = re.sub(r"\\end\{abstract\}", "\n\n", tex)

        # ── 保留行内格式（转为 Markdown） ──
        tex = re.sub(r"\\textbf\{([^}]*)\}", r"**\1**", tex)
        tex = re.sub(r"\\(?:textit|emph)\{([^}]*)\}", r"*\1*", tex)
        tex = re.sub(r"\\texttt\{([^}]*)\}", r"`\1`", tex)
        # underline/sout 转换放在清理之后处理，避免 _ 和 ~ 被清除

        # ── 图片占位 ──
        tex = re.sub(
            r"\\includegraphics(?:\[[^\]]*\])?\{([^}]*)\}",
            r"[图片: \1]",
            tex,
        )

        # ── 移除引用/标签命令 ──
        tex = re.sub(r"\\(?:cite|ref|label|pageref|autoref|eqref)\{[^}]*\}", "", tex)

        # ── 提取列表项 ──
        tex = re.sub(r"\\item\s+", "- ", tex)

        # ── 提取脚注 ──
        tex = re.sub(r"\\footnote\{([^}]*)\}", r" (脚注: \1)", tex)

        # ── 移除环境标记（保留内容） ──
        # 对于 figure/equation 等，移除 begin/end 标记但保留内容
        env_names = "figure|table|equation|align|gather|itemize|enumerate|description|lstlisting|verbatim|quote|quotation|theorem|lemma|proof"
        tex = re.sub(rf"\\begin\{{(?:{env_names})\}}", "", tex)
        tex = re.sub(rf"\\end\{{(?:{env_names})\}}", "", tex)

        # ── 移除剩余的 LaTeX 命令（保留参数内容） ──
        # 匹配 \command[opts]{args} 形式，保留最内层花括号内容
        def _strip_command(match: re.Match) -> str:
            # 返回最后一个 { } 内的内容
            groups = re.findall(r"\{([^{}]*)\}", match.group(0))
            return groups[-1] if groups else ""

        tex = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?(?:\{[^}]*\})+", _strip_command, tex)

        # ── 移除单独的 \command（无参数） ──
        # 保留常见换行/分页命令的语义
        tex = re.sub(r"\\(?:newline|linebreak|pagebreak)\b", "\n", tex)
        tex = re.sub(r"\\(?:newpage)\b", "\n\n", tex)
        tex = re.sub(r"\\(?:par|bigskip|medskip|smallskip)\b", "\n\n", tex)
        tex = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?\s*", "", tex)

        # ── 清理特殊字符 ──
        tex = re.sub(r"[{}$\\]", "", tex)
        # 保留 & （表格分隔符已被处理）
        tex = re.sub(r"[&^]", "", tex)

        # ── 行内格式后置处理（避免 _ 和 ~ 被清理误删） ──
        # 注意：此时 \underline 和 \sout 已被通用命令清理移除，
        # 因此这些转换仅用于处理未被清理的残留情况
        # \underline{...} → <u>...</u>（HTML 格式在 docx 中更可靠）
        tex = re.sub(r"\\underline\{([^}]*)\}", r"<u>\1</u>", tex)
        tex = re.sub(r"\\sout\{([^}]*)\}", r"<s>\1</s>", tex)

        # ── 清理空行和空格 ──
        tex = re.sub(r"\n{3,}", "\n\n", tex)
        tex = re.sub(r"[ \t]+\n", "\n", tex)
        tex = re.sub(r"\n[ \t]+", "\n", tex)
        tex = tex.strip()

        return tex


def get_supported_formats_display() -> str:
    """返回支持的格式说明字符串"""
    return ".docx, .doc, .odt, .rtf, .pdf, .txt, .md"
