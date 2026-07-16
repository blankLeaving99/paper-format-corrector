"""离线文档分析器

分析 .docx 文档的结构和格式，无需 LLM API。
可检测段落类型、提取格式信息、检查格式一致性。

功能：
- 段落类型检测（标题、正文、摘要、参考文献等）
- 字体/字号/对齐方式提取
- 格式一致性检查
- 与标准规范的差异分析
"""

import re
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt


class ParagraphType(Enum):
    """段落类型枚举"""
    TITLE = "title"
    AUTHOR = "author"
    AFFILIATION = "affiliation"
    ABSTRACT_CN = "abstract_cn"
    ABSTRACT_EN = "abstract_en"
    KEYWORDS_CN = "keywords_cn"
    KEYWORDS_EN = "keywords_en"
    HEADING1 = "heading1"
    HEADING2 = "heading2"
    HEADING3 = "heading3"
    BODY = "body"
    FIGURE_CAPTION = "figure_caption"
    TABLE_CAPTION = "table_caption"
    FORMULA = "formula"
    CODE = "code"
    REFERENCE = "reference"
    ACKNOWLEDGMENT = "acknowledgment"
    APPENDIX = "appendix"
    TOC = "toc"
    PAGE_BREAK = "page_break"
    UNKNOWN = "unknown"


@dataclass
class ParagraphInfo:
    """段落信息"""
    index: int
    text: str
    paragraph_type: ParagraphType
    style_name: str = ""
    font_name: str = ""
    font_size: float = 0.0
    is_bold: bool = False
    is_italic: bool = False
    alignment: str = ""
    line_spacing: float = 1.0
    first_line_indent: float = 0.0
    space_before: float = 0.0
    space_after: float = 0.0


@dataclass
class DocumentAnalysis:
    """文档分析结果"""
    total_paragraphs: int = 0
    paragraphs: list[ParagraphInfo] = field(default_factory=list)
    structure: dict[str, int] = field(default_factory=dict)
    fonts_used: dict[str, int] = field(default_factory=dict)
    font_sizes_used: dict[float, int] = field(default_factory=dict)
    issues: list[dict[str, Any]] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)


class DocumentAnalyzer:
    """文档分析器（离线版）

    分析 .docx 文档的结构和格式，检测段落类型，
    提取格式信息，检查格式一致性。
    """

    # 默认检测模式
    DEFAULT_PATTERNS = {
        "title": r"^论文题目[:：]?|^题目[:：]?",
        "chapter": r"^第[一二三四五六七八九十百零\d]+[章部分篇]",
        "section": r"^\d+\.\d+",
        "subsection": r"^\d+\.\d+\.\d+",
        "abstract_cn": r"^摘\s*要$",
        "abstract_en": r"^Abstract$|^ABSTRACT$",
        "keywords_cn": r"^关键词[:：]",
        "keywords_en": r"^Key\s*[Ww]ords[:：]?",
        "reference": r"^参考文献$|^References$|^REFERENCES$",
        "acknowledgment": r"^致\s*谢$|^Acknowledge?ments?$",
        "appendix": r"^附\s*录[A-Z]?",
        "figure_caption": r"^图\s*\d",
        "table_caption": r"^表\s*\d",
        "formula": r"^\(?\d+[-\.]\d+\)?$",
    }

    def __init__(self, patterns: dict[str, str] | None = None):
        """初始化分析器

        Args:
            patterns: 自定义检测模式字典，None 使用默认模式
        """
        self.patterns = patterns or self.DEFAULT_PATTERNS
        self._compiled = {
            name: re.compile(pattern)
            for name, pattern in self.patterns.items()
        }

    def analyze(self, doc_path: str | Path) -> DocumentAnalysis:
        """分析文档

        Args:
            doc_path: .docx 文件路径

        Returns:
            DocumentAnalysis 分析结果
        """
        doc = Document(str(doc_path))
        result = DocumentAnalysis()

        # 分析页面设置
        self._analyze_page_setup(doc, result)

        # 分析每个段落
        for i, para in enumerate(doc.paragraphs):
            info = self._analyze_paragraph(i, para, doc)
            result.paragraphs.append(info)

            # 统计结构
            ptype = info.paragraph_type.value
            result.structure[ptype] = result.structure.get(ptype, 0) + 1

            # 统计字体
            if info.font_name:
                result.fonts_used[info.font_name] = result.fonts_used.get(info.font_name, 0) + 1

            # 统计字号
            if info.font_size > 0:
                result.font_sizes_used[info.font_size] = result.font_sizes_used.get(info.font_size, 0) + 1

        result.total_paragraphs = len(doc.paragraphs)

        # 检查格式问题
        self._check_issues(result)

        return result

    def _analyze_page_setup(self, doc: Document, result: DocumentAnalysis) -> None:
        """分析页面设置"""
        if doc.sections:
            section = doc.sections[0]
            from docx.shared import Cm
            result.metadata["margins"] = {
                "top": section.top_margin / Cm(1) if section.top_margin else 0,
                "bottom": section.bottom_margin / Cm(1) if section.bottom_margin else 0,
                "left": section.left_margin / Cm(1) if section.left_margin else 0,
                "right": section.right_margin / Cm(1) if section.right_margin else 0,
            }
            result.metadata["page_width"] = section.page_width
            result.metadata["page_height"] = section.page_height

    def _analyze_paragraph(self, index: int, para, doc: Document) -> ParagraphInfo:
        """分析单个段落"""
        text = para.text.strip()
        ptype = self._detect_paragraph_type(text, para)

        # 提取格式信息
        font_name = ""
        font_size = 0.0
        is_bold = False
        is_italic = False

        for run in para.runs:
            if run.font.name:
                font_name = run.font.name
            if run.font.size:
                font_size = run.font.size.pt
            if run.font.bold is not None:
                is_bold = run.font.bold
            if run.font.italic is not None:
                is_italic = run.font.italic
            # 只取第一个 run 的信息
            break

        # 对齐方式
        alignment = ""
        if para.alignment is not None:
            alignment_map = {
                0: "left", 1: "center", 2: "right", 3: "justify"
            }
            alignment = alignment_map.get(para.alignment, "")

        # 段落格式
        pf = para.paragraph_format
        line_spacing = 1.0
        first_line_indent = 0.0
        space_before = 0.0
        space_after = 0.0

        if pf.line_spacing is not None:
            line_spacing = pf.line_spacing
        if pf.first_line_indent is not None:
            first_line_indent = pf.first_line_indent
        if pf.space_before is not None:
            space_before = pf.space_before
        if pf.space_after is not None:
            space_after = pf.space_after

        return ParagraphInfo(
            index=index,
            text=text[:100],  # 截断过长文本
            paragraph_type=ptype,
            style_name=para.style.name if para.style else "",
            font_name=font_name,
            font_size=font_size,
            is_bold=is_bold,
            is_italic=is_italic,
            alignment=alignment,
            line_spacing=line_spacing,
            first_line_indent=first_line_indent,
            space_before=space_before,
            space_after=space_after,
        )

    def _detect_paragraph_type(self, text: str, para) -> ParagraphType:
        """检测段落类型"""
        if not text:
            return ParagraphType.UNKNOWN

        # 按优先级检查各类模式
        checks = [
            ("reference", ParagraphType.REFERENCE),
            ("acknowledgment", ParagraphType.ACKNOWLEDGMENT),
            ("appendix", ParagraphType.APPENDIX),
            ("abstract_cn", ParagraphType.ABSTRACT_CN),
            ("abstract_en", ParagraphType.ABSTRACT_EN),
            ("keywords_cn", ParagraphType.KEYWORDS_CN),
            ("keywords_en", ParagraphType.KEYWORDS_EN),
            ("title", ParagraphType.TITLE),
            ("chapter", ParagraphType.HEADING1),
            ("section", ParagraphType.HEADING2),
            ("subsection", ParagraphType.HEADING3),
            ("figure_caption", ParagraphType.FIGURE_CAPTION),
            ("table_caption", ParagraphType.TABLE_CAPTION),
            ("formula", ParagraphType.FORMULA),
        ]

        for pattern_name, ptype in checks:
            if pattern_name in self._compiled:
                if self._compiled[pattern_name].search(text):
                    return ptype

        # 检查是否是标题样式
        if para.style and para.style.name:
            style_name = para.style.name.lower()
            if "heading 1" in style_name or "标题 1" in style_name:
                return ParagraphType.HEADING1
            elif "heading 2" in style_name or "标题 2" in style_name:
                return ParagraphType.HEADING2
            elif "heading 3" in style_name or "标题 3" in style_name:
                return ParagraphType.HEADING3
            elif "toc" in style_name:
                return ParagraphType.TOC

        # 默认为正文
        return ParagraphType.BODY

    def _check_issues(self, result: DocumentAnalysis) -> None:
        """检查格式问题"""
        issues = []

        # 检查字体一致性
        if len(result.fonts_used) > 3:
            issues.append({
                "type": "font_consistency",
                "severity": "warning",
                "message": f"使用了 {len(result.fonts_used)} 种字体，建议统一字体",
                "details": dict(list(result.fonts_used.items())[:5]),
            })

        # 检查字号一致性
        if len(result.font_sizes_used) > 5:
            issues.append({
                "type": "size_consistency",
                "severity": "warning",
                "message": f"使用了 {len(result.font_sizes_used)} 种字号，建议统一字号",
                "details": result.font_sizes_used,
            })

        # 检查是否有标题但无正文
        heading_count = sum(
            result.structure.get(t, 0)
            for t in ["heading1", "heading2", "heading3"]
        )
        body_count = result.structure.get("body", 0)
        if heading_count > 0 and body_count == 0:
            issues.append({
                "type": "structure",
                "severity": "error",
                "message": "检测到标题但无正文内容",
            })

        # 检查是否有摘要
        if "abstract_cn" not in result.structure and "abstract_en" not in result.structure:
            issues.append({
                "type": "structure",
                "severity": "warning",
                "message": "未检测到摘要",
            })

        # 检查页边距
        margins = result.metadata.get("margins", {})
        if margins:
            expected = {"top": 2.54, "bottom": 2.54, "left": 3.17, "right": 3.17}
            for side, expected_val in expected.items():
                actual_val = margins.get(side, 0)
                if abs(actual_val - expected_val) > 0.1:
                    issues.append({
                        "type": "margin",
                        "severity": "info",
                        "message": f"{side}边距: {actual_val:.2f}cm (预期 {expected_val}cm)",
                    })

        result.issues = issues

    def compare_with_config(
        self, analysis: DocumentAnalysis, config: dict
    ) -> list[dict[str, Any]]:
        """将分析结果与配置进行比较

        Args:
            analysis: 文档分析结果
            config: 标准配置字典

        Returns:
            差异列表
        """
        diffs = []
        format_rules = config.get("format_rules", {})

        # 检查字体
        expected_font = format_rules.get("font", {})
        if expected_font:
            cn_font = expected_font.get("chinese", "")
            en_font = expected_font.get("english", "")
            for font_name, count in analysis.fonts_used.items():
                if cn_font and cn_font not in font_name and font_name not in cn_font:
                    diffs.append({
                        "type": "font",
                        "severity": "warning",
                        "message": f"字体 '{font_name}' 不符合预期 '{cn_font}'",
                        "count": count,
                    })

        # 检查字号
        body_size = format_rules.get("body_text", {}).get("font_size")
        if body_size:
            for size, count in analysis.font_sizes_used.items():
                if abs(size - body_size) > 0.5:
                    diffs.append({
                        "type": "font_size",
                        "severity": "info",
                        "message": f"字号 {size}pt 出现 {count} 次 (预期正文 {body_size}pt)",
                        "count": count,
                    })

        # 检查页边距
        expected_margins = format_rules.get("margins", {})
        if expected_margins:
            actual_margins = analysis.metadata.get("margins", {})
            for side in ["top", "bottom", "left", "right"]:
                expected_val = expected_margins.get(side)
                actual_val = actual_margins.get(side)
                if expected_val and actual_val:
                    if abs(actual_val - expected_val) > 0.1:
                        diffs.append({
                            "type": "margin",
                            "severity": "warning",
                            "message": f"{side}边距: {actual_val:.2f}cm (预期 {expected_val}cm)",
                        })

        return diffs


def analyze_document(doc_path: str | Path) -> DocumentAnalysis:
    """便捷函数：分析文档

    Args:
        doc_path: .docx 文件路径

    Returns:
        DocumentAnalysis 分析结果
    """
    analyzer = DocumentAnalyzer()
    return analyzer.analyze(doc_path)
