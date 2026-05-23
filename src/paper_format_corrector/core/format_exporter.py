"""多格式导出模块

支持将 .docx 导出为:
- PDF (通过 docx2pdf 或 LibreOffice)
- HTML (通过 mammoth)
- TXT (纯文本)
- Markdown
"""

import re
from pathlib import Path

from ..infra.path_security import validate_input_path, validate_output_path, ALLOWED_INPUT_EXTENSIONS, ALLOWED_OUTPUT_EXTENSIONS


class FormatExporter:
    """论文格式导出器"""

    SUPPORTED_FORMATS = ("pdf", "html", "txt", "md", "markdown")

    def __init__(self, config=None):
        self._docx2pdf = None
        self._mammoth = None
        self._config = config or {}
        # 从 config 读取标题检测模式
        detect = self._config.get("auto_detect", {})
        self._chapter_re = re.compile(
            detect.get("chapter_pattern", r"^第[一二三四五六七八九十百零\d]+[章部分篇]")
        )
        self._section_re = re.compile(detect.get("section_pattern", r"^\d+\.\d+"))
        self._subsection_re = re.compile(detect.get("subsection_pattern", r"^\d+\.\d+\.\d+"))
        self._fig_re = re.compile(detect.get("figure_caption_pattern", r"^图\s*\d"))
        self._tab_re = re.compile(detect.get("table_caption_pattern", r"^表\s*\d"))

    def export(self, docx_path, output_path, fmt):
        """导出为指定格式"""
        fmt = fmt.lower().strip(".")
        if fmt == "markdown":
            fmt = "md"

        if fmt not in self.SUPPORTED_FORMATS:
            raise ValueError(
                f"不支持的格式: {fmt}，支持: {', '.join(self.SUPPORTED_FORMATS)}"
            )

        docx_path = validate_input_path(docx_path, ALLOWED_INPUT_EXTENSIONS)
        output_path = validate_output_path(output_path, ALLOWED_OUTPUT_EXTENSIONS)

        if not docx_path.exists():
            raise FileNotFoundError(f"文件不存在: {docx_path}")

        output_path.parent.mkdir(parents=True, exist_ok=True)

        if fmt == "pdf":
            return self._export_pdf(docx_path, output_path)
        elif fmt == "html":
            return self._export_html(docx_path, output_path)
        elif fmt == "txt":
            return self._export_txt(docx_path, output_path)
        elif fmt == "md":
            return self._export_markdown(docx_path, output_path)

    def _export_pdf(self, docx_path, output_path):
        """导出为 PDF"""
        try:
            from docx2pdf import convert
            convert(str(docx_path), str(output_path))
            return str(output_path)
        except ImportError:
            pass

        # 回退：使用 LibreOffice
        import subprocess

        lo_path = self._find_libreoffice()
        if lo_path:
            out_dir = output_path.parent
            subprocess.run(
                [
                    lo_path,
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", str(out_dir),
                    str(docx_path),
                ],
                check=True,
                capture_output=True,
            )
            # LibreOffice 输出文件名基于输入文件名
            lo_output = out_dir / f"{docx_path.stem}.pdf"
            if lo_output != output_path and lo_output.exists():
                lo_output.rename(output_path)
            return str(output_path)

        raise RuntimeError(
            "PDF 导出需要安装 docx2pdf (pip install docx2pdf) 或 LibreOffice"
        )

    def _export_html(self, docx_path, output_path):
        """导出为 HTML"""
        try:
            import mammoth
            with open(docx_path, "rb") as f:
                result = mammoth.convert_to_html(f)
                html_content = result.value

            full_html = self._wrap_html(html_content, docx_path.stem)
            output_path.write_text(full_html, encoding="utf-8")
            return str(output_path)
        except ImportError:
            pass

        # 回退：使用 python-docx 手动转 HTML
        html_content = self._docx_to_html_manual(docx_path)
        output_path.write_text(html_content, encoding="utf-8")
        return str(output_path)

    def _export_txt(self, docx_path, output_path):
        """导出为纯文本"""
        from docx import Document

        doc = Document(str(docx_path))
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
            else:
                lines.append("")

        output_path.write_text("\n".join(lines), encoding="utf-8")
        return str(output_path)

    def _export_markdown(self, docx_path, output_path):
        """导出为 Markdown"""
        from docx import Document
        import re

        doc = Document(str(docx_path))
        lines = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                lines.append("")
                continue

            style_name = para.style.name.lower() if para.style else ""

            # 标题映射
            if "heading 1" in style_name or "heading1" in style_name:
                lines.append(f"## {text}")
            elif "heading 2" in style_name or "heading2" in style_name:
                lines.append(f"### {text}")
            elif "heading 3" in style_name or "heading3" in style_name:
                lines.append(f"#### {text}")
            elif self._is_chapter(text):
                lines.append(f"## {text}")
            elif self._is_section(text):
                lines.append(f"### {text}")
            elif self._is_subsection(text):
                lines.append(f"#### {text}")
            elif self._is_figure_caption(text):
                lines.append(f"*{text}*")
            elif self._is_table_caption(text):
                lines.append(f"*{text}*")
            else:
                # 正文：处理加粗和斜体
                md_text = self._runs_to_markdown(para)
                lines.append(md_text)

        output_path.write_text("\n".join(lines), encoding="utf-8")
        return str(output_path)

    def _runs_to_markdown(self, paragraph):
        """将 paragraph 的 runs 转换为 Markdown"""
        parts = []
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            if run.bold and run.italic:
                parts.append(f"***{text}***")
            elif run.bold:
                parts.append(f"**{text}**")
            elif run.italic:
                parts.append(f"*{text}*")
            else:
                parts.append(text)
        return "".join(parts)

    def _docx_to_html_manual(self, docx_path):
        """手动将 docx 转换为 HTML（不依赖 mammoth）"""
        from docx import Document

        doc = Document(str(docx_path))
        body_parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                body_parts.append("<br>")
                continue

            style_name = para.style.name.lower() if para.style else ""

            if "heading 1" in style_name or self._is_chapter(text):
                body_parts.append(f"<h2>{self._escape_html(text)}</h2>")
            elif "heading 2" in style_name or self._is_section(text):
                body_parts.append(f"<h3>{self._escape_html(text)}</h3>")
            elif "heading 3" in style_name or self._is_subsection(text):
                body_parts.append(f"<h4>{self._escape_html(text)}</h4>")
            else:
                html_runs = []
                for run in para.runs:
                    t = self._escape_html(run.text)
                    if run.bold:
                        t = f"<strong>{t}</strong>"
                    if run.italic:
                        t = f"<em>{t}</em>"
                    html_runs.append(t)
                body_parts.append(f"<p>{''.join(html_runs)}</p>")

        return self._wrap_html("\n".join(body_parts), docx_path.stem)

    def _wrap_html(self, body_html, title="论文"):
        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{self._escape_html(title)}</title>
    <style>
        body {{ font-family: "SimSun", "Times New Roman", serif; max-width: 800px; margin: 40px auto; padding: 0 20px; line-height: 1.8; }}
        h1, h2 {{ text-align: center; font-family: "SimHei", sans-serif; }}
        h3 {{ font-family: "SimHei", sans-serif; }}
        p {{ text-indent: 2em; margin: 0.5em 0; }}
        table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
        td, th {{ border: 1px solid #999; padding: 6px 12px; }}
        .caption {{ text-align: center; font-size: 0.9em; margin: 0.5em 0; }}
    </style>
</head>
<body>
{body_html}
</body>
</html>"""

    def _escape_html(self, text):
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
        )

    def _is_chapter(self, text):
        return bool(self._chapter_re.match(text))

    def _is_section(self, text):
        return bool(self._section_re.match(text))

    def _is_subsection(self, text):
        return bool(self._subsection_re.match(text))

    def _is_figure_caption(self, text):
        return bool(self._fig_re.match(text))

    def _is_table_caption(self, text):
        return bool(self._tab_re.match(text))

    def _find_libreoffice(self):
        """查找 LibreOffice 安装路径"""
        import shutil

        # 只检查绝对路径，避免 PATH 污染风险
        candidates = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            "/usr/bin/libreoffice",
            "/usr/bin/soffice",
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        ]
        for candidate in candidates:
            if Path(candidate).exists():
                return candidate

        # PATH 搜索仅作最后回退
        for name in ("soffice", "libreoffice"):
            found = shutil.which(name)
            if found:
                return found
        return None
