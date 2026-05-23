import re
import logging

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

from ..parsers.section_detector import SectionDetector, SectionType, detect_document_language
from ..parsers.reference_formatter import ReferenceFormatter
from ..handlers.figure_table_handler import FigureTableHandler
from ..handlers.table_handler import TableHandler
from ..handlers.image_handler import ImageHandler
from ..handlers.header_footer_handler import HeaderFooterHandler
from ..handlers.toc_handler import TOCHandler

logger = logging.getLogger(__name__)


class FormatCorrector:
    """论文格式矫正器（完整版）"""

    def __init__(self, template_path, config):
        self.template = Document(template_path)
        self.config = config
        self.style_mapping = {}

        # 子模块
        self.section_detector = SectionDetector(config)
        self.ref_formatter = ReferenceFormatter(config)
        self.fig_table_handler = FigureTableHandler(config)
        self.table_handler = TableHandler(config)
        self.image_handler = ImageHandler(config)
        self.hf_handler = HeaderFooterHandler(config)
        self.toc_handler = TOCHandler(config)

        # 处理报告
        self.report = self._empty_report()

    def _empty_report(self):
        return {
            "paragraphs_corrected": 0,
            "headings_fixed": 0,
            "body_fixed": 0,
            "figures_renumbered": 0,
            "tables_renumbered": 0,
            "tables_formatted": 0,
            "images_centered": 0,
            "ref_issues": [],
            "fig_table_issues": [],
            "warnings": [],
        }

    def load_template_styles(self):
        for style in self.template.styles:
            if style.type == 1:
                self.style_mapping[style.name.lower()] = style

    def correct_document(self, input_path, output_path):
        logger.info(f"正在处理: {input_path}")
        doc = Document(input_path)

        # 重置状态
        self.section_detector.reset()
        self.fig_table_handler.reset()
        self.report = self._empty_report()

        # 检测文档语言并调整字体配置
        self._detect_and_apply_language(doc)

        # 1. 页面设置
        self._apply_page_setup(doc)

        # 2. 段落格式矫正（核心循环）
        self._correct_paragraphs(doc)

        # 3. 表格格式矫正
        self.report["tables_formatted"] = self.table_handler.format_all_tables(doc)

        # 4. 图片处理（居中 + 调整大小 + DPI检查）
        self.report["images_centered"] = self.image_handler.process_all_images(doc)

        # 5. 参考文献格式化
        self._format_references(doc)

        # 6. 插入目录（如果需要）
        if self.toc_handler.enabled and not self.toc_handler.has_toc(doc):
            self.toc_handler.insert_toc(doc)

        # 7. 页眉页脚
        self.hf_handler.apply(doc)

        # 8. 公式编号
        self._correct_formulas(doc)

        doc.save(output_path)
        logger.info(f"已保存: {output_path}")
        return self.report

    def _detect_and_apply_language(self, doc):
        """检测文档语言并调整字体配置"""
        lang_config = self.config.get("language", {})
        primary = lang_config.get("primary", "auto")

        if primary == "auto":
            primary = detect_document_language(doc)
            logger.info(f"  检测到文档语言: {primary}")

        # 根据语言调整字体配置
        fonts = lang_config.get("fonts", {})
        lang_fonts = fonts.get(primary, {})

        if lang_fonts:
            font_rules = self.config.setdefault("format_rules", {}).setdefault("font", {})
            if "body" in lang_fonts:
                font_rules["chinese"] = lang_fonts["body"]
            if "heading" in lang_fonts:
                font_rules["heading_chinese"] = lang_fonts["heading"]
            if "english_in_chinese" in lang_fonts:
                font_rules["english"] = lang_fonts["english_in_chinese"]

        self._detected_language = primary

    def _apply_page_setup(self, doc):
        margins = self.config.get("format_rules", {}).get("margins", {})
        if not margins:
            return
        for section in doc.sections:
            section.top_margin = Cm(margins.get("top", 2.54))
            section.bottom_margin = Cm(margins.get("bottom", 2.54))
            section.left_margin = Cm(margins.get("left", 3.17))
            section.right_margin = Cm(margins.get("right", 3.17))

    def _correct_paragraphs(self, doc):
        """逐段矫正格式"""
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue

            # 检测段落类型
            section_type, extra = self.section_detector.detect(paragraph)

            # 根据类型应用不同格式
            if section_type == SectionType.TITLE:
                self._apply_title_style(paragraph)
            elif section_type == SectionType.AUTHORS:
                self._apply_author_style(paragraph)
            elif section_type == SectionType.AFFILIATION:
                self._apply_affiliation_style(paragraph)
            elif section_type in (SectionType.ABSTRACT_CN, SectionType.ABSTRACT_EN):
                self._apply_abstract_style(paragraph, section_type)
            elif section_type in (SectionType.KEYWORDS_CN, SectionType.KEYWORDS_EN):
                self._apply_keywords_style(paragraph, section_type)
            elif section_type == SectionType.CHAPTER:
                self._apply_heading_style(paragraph, "heading1")
                chapter_num = extra.get("chapter_num", 0)
                self.fig_table_handler.update_chapter(chapter_num)
                self.report["headings_fixed"] += 1
            elif section_type == SectionType.SECTION:
                self._apply_heading_style(paragraph, "heading2")
                self.report["headings_fixed"] += 1
            elif section_type == SectionType.SUBSECTION:
                self._apply_heading_style(paragraph, "heading3")
                self.report["headings_fixed"] += 1
            elif section_type == SectionType.REFERENCE_TITLE:
                self._apply_heading_style(paragraph, "heading1")
                self.report["headings_fixed"] += 1
            elif section_type in (SectionType.ACKNOWLEDGMENT_TITLE, SectionType.APPENDIX_TITLE):
                self._apply_heading_style(paragraph, "heading1")
                self.report["headings_fixed"] += 1
            elif section_type == SectionType.FIGURE_CAPTION:
                self.fig_table_handler.process_paragraph(paragraph, section_type, extra)
            elif section_type == SectionType.TABLE_CAPTION:
                self.fig_table_handler.process_paragraph(paragraph, section_type, extra)
            elif section_type == SectionType.BODY:
                self._apply_body_style(paragraph)
                self.report["body_fixed"] += 1
            elif section_type == SectionType.REFERENCE_ITEM:
                pass

            self.report["paragraphs_corrected"] += 1

        self.report["fig_table_issues"] = self.fig_table_handler.get_issues()

    def _format_references(self, doc):
        ref_start = None
        ref_keywords = self.config.get("auto_detect", {}).get(
            "reference_keywords", ["参考文献"]
        )
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text in ref_keywords or any(text.startswith(kw) for kw in ref_keywords):
                ref_start = i
                break
        if ref_start is not None:
            self.ref_formatter.format_references(doc, ref_start)
            self.report["ref_issues"] = self.ref_formatter.validate_references(doc, ref_start)

    def _correct_formulas(self, doc):
        formula_config = self.config.get("format_rules", {}).get("formulas", {})
        if not formula_config.get("numbering", True):
            return
        formula_pattern = re.compile(
            self.config.get("auto_detect", {}).get("formula_pattern", r"^\(?\d+[-\.]\d+\)?$")
        )
        for para in doc.paragraphs:
            text = para.text.strip()
            if formula_pattern.match(text):
                self._apply_formula_style(para, formula_config)

    def _apply_formula_style(self, paragraph, config):
        style_rules = {"font_size": config.get("font_size", 12)}
        font_rules = self.config.get("format_rules", {}).get("font", {})
        for run in paragraph.runs:
            self._set_run_font(run, font_rules, style_rules)
        if config.get("numbering_position") == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # ========== 各类段落样式应用 ==========

    def _apply_title_style(self, paragraph):
        rules = self.config.get("format_rules", {})
        tp = rules.get("title_page", {})
        font_rules = rules.get("font", {})
        for run in paragraph.runs:
            run.font.name = font_rules.get("english", "Times New Roman")
            self._set_east_asian_font(run, font_rules.get("heading_chinese", "黑体"))
            run.font.size = Pt(tp.get("title_font_size", 22))
            run.font.bold = tp.get("title_bold", True)
        align_map = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        paragraph.alignment = align_map.get(tp.get("title_align", "center"), WD_ALIGN_PARAGRAPH.CENTER)
        paragraph.paragraph_format.space_before = Pt(24)
        paragraph.paragraph_format.space_after = Pt(18)
        paragraph.paragraph_format.line_spacing = 1.5

    def _apply_author_style(self, paragraph):
        rules = self.config.get("format_rules", {})
        tp = rules.get("title_page", {})
        font_rules = rules.get("font", {})
        for run in paragraph.runs:
            run.font.name = font_rules.get("english", "Times New Roman")
            self._set_east_asian_font(run, font_rules.get("chinese", "宋体"))
            run.font.size = Pt(tp.get("author_font_size", 12))
            run.font.bold = False
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)

    def _apply_affiliation_style(self, paragraph):
        rules = self.config.get("format_rules", {})
        tp = rules.get("title_page", {})
        font_rules = rules.get("font", {})
        for run in paragraph.runs:
            run.font.name = font_rules.get("english", "Times New Roman")
            self._set_east_asian_font(run, font_rules.get("chinese", "宋体"))
            run.font.size = Pt(tp.get("affiliation_font_size", 10.5))
            run.font.bold = False
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _apply_abstract_style(self, paragraph, section_type):
        rules = self.config.get("format_rules", {})
        abs_config = rules.get("abstract", {})
        font_rules = rules.get("font", {})
        text = paragraph.text.strip()

        abstract_label_cn = self.config.get("auto_detect", {}).get("abstract_pattern", r"^摘\s*要$")
        abstract_label_en = self.config.get("auto_detect", {}).get("abstract_en_pattern", r"^Abstract$|^ABSTRACT$")

        is_title_line = re.match(abstract_label_cn, text) or re.match(abstract_label_en, text)

        if is_title_line:
            for run in paragraph.runs:
                run.font.name = font_rules.get("english", "Times New Roman")
                self._set_east_asian_font(run, font_rules.get("heading_chinese", "黑体"))
                run.font.size = Pt(abs_config.get("title_font_size", 16))
                run.font.bold = abs_config.get("title_bold", True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            for run in paragraph.runs:
                run.font.name = font_rules.get("english", "Times New Roman")
                self._set_east_asian_font(run, font_rules.get("chinese", "宋体"))
                run.font.size = Pt(abs_config.get("content_font_size", 12))
                run.font.bold = False
            self._apply_line_spacing(paragraph, abs_config.get("content_line_spacing", 1.5))
            indent = abs_config.get("content_first_line_indent", 0)
            if indent:
                paragraph.paragraph_format.first_line_indent = Pt(indent * 12)

    def _apply_keywords_style(self, paragraph, section_type):
        rules = self.config.get("format_rules", {})
        kw_config = rules.get("keywords", {})
        font_rules = rules.get("font", {})
        text = paragraph.text.strip()

        if "：" in text or ":" in text:
            sep_pos = max(text.find("："), text.find(":"))
            label_part = text[:sep_pos + 1]
            content_part = text[sep_pos + 1:]
        else:
            label_part = ""
            content_part = text

        for run in paragraph.runs:
            run.text = ""

        if label_part and paragraph.runs:
            paragraph.runs[0].text = label_part
            paragraph.runs[0].font.bold = kw_config.get("bold_label", True)
            paragraph.runs[0].font.size = Pt(kw_config.get("font_size", 12))
            paragraph.runs[0].font.name = font_rules.get("english", "Times New Roman")
            self._set_east_asian_font(paragraph.runs[0], font_rules.get("chinese", "宋体"))
            if content_part:
                content_run = paragraph.add_run(content_part)
                content_run.font.bold = False
                content_run.font.size = Pt(kw_config.get("font_size", 12))
                content_run.font.name = font_rules.get("english", "Times New Roman")
                self._set_east_asian_font(content_run, font_rules.get("chinese", "宋体"))
        else:
            for run in paragraph.runs:
                run.font.name = font_rules.get("english", "Times New Roman")
                self._set_east_asian_font(run, font_rules.get("chinese", "宋体"))
                run.font.size = Pt(kw_config.get("font_size", 12))

        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)

    def _apply_heading_style(self, paragraph, style_name):
        rules = self.config.get("format_rules", {})
        heading_rules = rules.get("headings", {}).get(style_name, {})
        font_rules = rules.get("font", {})

        for run in paragraph.runs:
            self._set_run_font(run, font_rules, heading_rules)
            self._set_east_asian_font(run, font_rules.get("heading_chinese", "黑体"))

        align_map = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        align = heading_rules.get("align", "left")
        paragraph.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)

        sb = heading_rules.get("space_before", 0)
        sa = heading_rules.get("space_after", 0)
        paragraph.paragraph_format.space_before = Pt(sb)
        paragraph.paragraph_format.space_after = Pt(sa)

        if "line_spacing" in heading_rules:
            self._apply_line_spacing(paragraph, heading_rules["line_spacing"])

    def _apply_body_style(self, paragraph):
        rules = self.config.get("format_rules", {})
        body_rules = rules.get("body_text", {})
        font_rules = rules.get("font", {})

        self._apply_mixed_font(paragraph, font_rules, body_rules)

        self._apply_line_spacing(paragraph, body_rules.get("line_spacing", 1.5))

        align = body_rules.get("align", "justify")
        if align == "justify":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        indent_chars = body_rules.get("first_line_indent", 0)
        if indent_chars:
            paragraph.paragraph_format.first_line_indent = Pt(
                indent_chars * body_rules.get("font_size", 12)
            )

        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

    def _apply_line_spacing(self, paragraph, spacing_value):
        """应用行距（支持多种格式）"""
        if isinstance(spacing_value, dict):
            stype = spacing_value.get("type", "multiple")
            value = spacing_value.get("value", 1.5)
            if stype == "multiple":
                paragraph.paragraph_format.line_spacing = value
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            elif stype == "exact":
                paragraph.paragraph_format.line_spacing = Pt(value)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            elif stype == "atLeast":
                paragraph.paragraph_format.line_spacing = Pt(value)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        else:
            paragraph.paragraph_format.line_spacing = spacing_value

    def _apply_mixed_font(self, paragraph, font_rules, style_rules):
        cn_font = font_rules.get("chinese", "宋体")
        en_font = font_rules.get("english", "Times New Roman")
        font_size = style_rules.get("font_size", 12)

        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            has_chinese = bool(re.search(r"[一-鿿]", text))
            has_english = bool(re.search(r"[a-zA-Z]", text))

            if has_chinese and has_english:
                run.font.name = en_font
                self._set_east_asian_font(run, cn_font)
            elif has_chinese:
                run.font.name = cn_font
                self._set_east_asian_font(run, cn_font)
            else:
                run.font.name = en_font

            run.font.size = Pt(font_size)
            if style_rules.get("bold") is not None:
                run.font.bold = style_rules["bold"]

    # ========== 工具方法 ==========

    def _set_run_font(self, run, font_rules, style_rules):
        run.font.name = font_rules.get("english", "Times New Roman")
        self._set_east_asian_font(run, font_rules.get("chinese", "宋体"))
        if style_rules.get("bold"):
            run.font.bold = True
        if style_rules.get("font_size"):
            run.font.size = Pt(style_rules["font_size"])

    def _set_east_asian_font(self, run, font_name):
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = run._element.makeelement(qn("w:rFonts"), {})
            rpr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), font_name)

    def get_report(self):
        return self.report
