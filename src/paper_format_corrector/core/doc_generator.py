"""文档生成器 v2

从结构化JSON数据从零创建Word文档，支持：
- 标题层级（Heading 1-4）
- 正文段落（含中英文字体混排）
- 表格
- 列表（有序/无序）
- 图片占位符
- 分页符
- 目录
- Markdown预览（不生成docx）
- 增量更新章节内容

与 FormatCorrector 不同，本模块专注于「从无到有」创建文档，
而非修改已有文档的格式。
"""

from __future__ import annotations

import copy
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from ..utils.docx_utils import set_east_asian_font


class DocGenerator:
    """结构化文档生成器"""

    # 默认格式配置
    DEFAULT_CONFIG = {
        "font": {
            "chinese": "宋体",
            "english": "Times New Roman",
            "heading_chinese": "黑体",
        },
        "headings": {
            "heading1": {"font_size": 22, "bold": True, "align": "center"},
            "heading2": {"font_size": 16, "bold": True, "align": "left"},
            "heading3": {"font_size": 14, "bold": True, "align": "left"},
            "heading4": {"font_size": 12, "bold": True, "align": "left"},
        },
        "body_text": {
            "font_size": 12,
            "line_spacing": 1.5,
            "first_line_indent": 2,
            "align": "justify",
        },
        "margins": {
            "top": 2.54,
            "bottom": 2.54,
            "left": 3.17,
            "right": 3.17,
        },
    }

    def __init__(self, config: dict | None = None):
        self.config = self._merge_config(self.DEFAULT_CONFIG, config or {})

    def generate(self, structure: dict[str, Any], output_path: str) -> str:
        """从结构化数据生成文档

        Args:
            structure: 文档结构JSON，格式：
                {
                    "title": "文档标题",
                    "sections": [
                        {"type": "heading1", "title": "第一章 ..."},
                        {"type": "body", "content": "正文内容..."},
                        {"type": "table", "header": [...], "rows": [[...], ...]},
                        {"type": "list", "ordered": false, "items": ["...", ...]},
                        {"type": "page_break"},
                        {"type": "image_placeholder", "caption": "图1: ..."},
                    ]
                }
            output_path: 输出文件路径

        Returns:
            输出文件路径
        """
        doc = Document()
        self._apply_page_setup(doc)

        sections = structure.get("sections", [])
        for section in sections:
            self._add_section(doc, section)

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return output_path

    def generate_preview(self, structure: dict[str, Any]) -> str:  # noqa: C901
        """生成Markdown格式的预览（不生成docx）

        Args:
            structure: 文档结构JSON

        Returns:
            Markdown格式的文档内容
        """
        lines: list[str] = []
        title = structure.get("title", "")
        if title:
            lines.append(f"# {title}")
            lines.append("")

        sections = structure.get("sections", [])
        for section in sections:
            section_type = section.get("type", "body")

            if section_type == "heading1":
                lines.append(f"## {section.get('title', '')}")
                lines.append("")
            elif section_type == "heading2":
                lines.append(f"### {section.get('title', '')}")
                lines.append("")
            elif section_type == "heading3":
                lines.append(f"#### {section.get('title', '')}")
                lines.append("")
            elif section_type == "heading4":
                lines.append(f"##### {section.get('title', '')}")
                lines.append("")
            elif section_type == "body":
                content = section.get("content", "")
                lines.append(content)
                lines.append("")
            elif section_type == "table":
                header = section.get("header", [])
                rows = section.get("rows", [])
                if header:
                    lines.append("| " + " | ".join(str(h) for h in header) + " |")
                    lines.append("| " + " | ".join("---" for _ in header) + " |")
                for row in rows:
                    lines.append("| " + " | ".join(str(c) for c in row) + " |")
                lines.append("")
            elif section_type == "list":
                items = section.get("items", [])
                ordered = section.get("ordered", False)
                for i, item in enumerate(items):
                    prefix = f"{i + 1}." if ordered else "-"
                    lines.append(f"{prefix} {item}")
                lines.append("")
            elif section_type == "page_break":
                lines.append("---")
                lines.append("")
            elif section_type == "image_placeholder":
                caption = section.get("caption", "图片")
                lines.append(f"*[{caption}]*")
                lines.append("")

        return "\n".join(lines)

    def update_section(
        self,
        structure: dict[str, Any],
        section_index: int,
        new_content: list[dict[str, Any]],
    ) -> dict[str, Any]:
        """更新指定章节的内容

        Args:
            structure: 当前文档结构
            section_index: 要更新的章节在sections数组中的索引
            new_content: 新的章节内容列表

        Returns:
            更新后的文档结构（深拷贝）
        """
        updated = copy.deepcopy(structure)
        sections = updated.get("sections", [])

        if section_index < 0 or section_index >= len(sections):
            raise ValueError(f"章节索引 {section_index} 超出范围（共 {len(sections)} 个章节）")

        # 找到当前章节的范围（从当前heading到下一个同级或更高级heading）
        current_section = sections[section_index]
        current_level = self._get_heading_level(current_section)

        # 找到结束位置
        end_index = len(sections)
        for i in range(section_index + 1, len(sections)):
            next_section = sections[i]
            next_level = self._get_heading_level(next_section)
            if next_level is not None and next_level <= current_level:
                end_index = i
                break

        # 替换章节内容
        updated["sections"] = (
            sections[:section_index] + new_content + sections[end_index:]
        )

        return updated

    def _get_heading_level(self, section: dict[str, Any]) -> int | None:
        """获取标题层级（非标题返回None）"""
        section_type = section.get("type", "")
        if section_type.startswith("heading"):
            try:
                return int(section_type[-1])
            except (ValueError, IndexError):
                pass
        return None

    def _apply_page_setup(self, doc: Document) -> None:
        """设置页面边距"""
        margins = self.config.get("margins", {})
        section = doc.sections[0]
        section.top_margin = Cm(margins.get("top", 2.54))
        section.bottom_margin = Cm(margins.get("bottom", 2.54))
        section.left_margin = Cm(margins.get("left", 3.17))
        section.right_margin = Cm(margins.get("right", 3.17))

    def _add_section(self, doc: Document, section: dict[str, Any]) -> None:
        """根据类型添加一个文档元素"""
        section_type = section.get("type", "body")
        handler = {
            "heading1": lambda: self._add_heading(doc, section, 1),
            "heading2": lambda: self._add_heading(doc, section, 2),
            "heading3": lambda: self._add_heading(doc, section, 3),
            "heading4": lambda: self._add_heading(doc, section, 4),
            "body": lambda: self._add_body(doc, section),
            "table": lambda: self._add_table(doc, section),
            "list": lambda: self._add_list(doc, section),
            "page_break": lambda: self._add_page_break(doc),
            "image_placeholder": lambda: self._add_image_placeholder(doc, section),
            "toc": lambda: self._add_toc(doc),
        }.get(section_type)

        if handler:
            handler()

    def _add_heading(self, doc: Document, section: dict, level: int) -> None:
        """添加标题"""
        title = section.get("title", "")
        if not title:
            return

        heading_key = f"heading{level}"
        style_cfg = self.config.get("headings", {}).get(heading_key, {})
        font_cfg = self.config.get("font", {})

        p = doc.add_paragraph()
        run = p.add_run(title)

        # 字体设置
        font_size = style_cfg.get("font_size", {1: 22, 2: 16, 3: 14, 4: 12}.get(level, 12))
        run.font.size = Pt(font_size)
        run.font.bold = style_cfg.get("bold", True)
        run.font.name = font_cfg.get("english", "Times New Roman")
        set_east_asian_font(run, font_cfg.get("heading_chinese", "黑体"))

        # 对齐
        align = style_cfg.get("align", "center" if level == 1 else "left")
        p.alignment = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }.get(align, WD_ALIGN_PARAGRAPH.LEFT)

        # 段前段后间距
        space_before = style_cfg.get("space_before", 12)
        space_after = style_cfg.get("space_after", 6)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)

        # 使用 python-docx 的 Heading 样式（如果可用）
        try:
            p.style = doc.styles[f"Heading {level}"]
            # 重新应用字体（样式可能会覆盖）
            for r in p.runs:
                r.font.size = Pt(font_size)
                r.font.bold = style_cfg.get("bold", True)
                r.font.name = font_cfg.get("english", "Times New Roman")
                set_east_asian_font(r, font_cfg.get("heading_chinese", "黑体"))
        except (KeyError, ValueError):
            pass

    def _add_body(self, doc: Document, section: dict[str, Any]) -> None:
        """添加正文段落"""
        content = section.get("content", "")
        if not content:
            return

        body_cfg = self.config.get("body_text", {})
        font_cfg = self.config.get("font", {})

        # 支持多段落（用 \n\n 分隔）
        paragraphs = content.split("\n\n") if "\n\n" in content else [content]

        for text in paragraphs:
            text = text.strip()
            if not text:
                continue

            p = doc.add_paragraph()
            run = p.add_run(text)

            # 字体
            run.font.size = Pt(body_cfg.get("font_size", 12))
            run.font.name = font_cfg.get("english", "Times New Roman")
            set_east_asian_font(run, font_cfg.get("chinese", "宋体"))

            # 对齐
            align = body_cfg.get("align", "justify")
            p.alignment = {
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)

            # 首行缩进
            indent = body_cfg.get("first_line_indent", 2)
            if indent > 0:
                p.paragraph_format.first_line_indent = Pt(font_cfg.get("body_font_size", 12) * indent)

            # 行距
            line_spacing = body_cfg.get("line_spacing", 1.5)
            if isinstance(line_spacing, (int, float)):
                p.paragraph_format.line_spacing = line_spacing

    def _add_table(self, doc: Document, section: dict[str, Any]) -> None:  # noqa: C901
        """添加表格"""
        header = section.get("header", [])
        rows = section.get("rows", [])

        if not header and not rows:
            return

        # 确定列数
        num_cols = len(header) if header else (len(rows[0]) if rows else 0)
        if num_cols == 0:
            return

        # 创建表格
        table = doc.add_table(rows=1 + len(rows), cols=num_cols)
        table.style = "Table Grid"

        # 写入表头
        if header:
            for i, cell_text in enumerate(header):
                cell = table.rows[0].cells[i]
                cell.text = str(cell_text)
                # 表头加粗
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.size = Pt(10.5)
                        run.font.name = "Times New Roman"
                        set_east_asian_font(run, "黑体")

        # 写入数据行
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.rows[row_idx + 1].cells[col_idx]
                    cell.text = str(cell_text)
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(10.5)
                            run.font.name = "Times New Roman"
                            set_east_asian_font(run, "宋体")

        # 表格后空行
        doc.add_paragraph()

    def _add_list(self, doc: Document, section: dict[str, Any]) -> None:
        """添加列表"""
        items = section.get("items", [])
        ordered = section.get("ordered", False)

        for i, item in enumerate(items):
            if ordered:
                prefix = f"{i + 1}. "
            else:
                prefix = "• "

            p = doc.add_paragraph()
            run = p.add_run(f"{prefix}{item}")

            body_cfg = self.config.get("body_text", {})
            font_cfg = self.config.get("font", {})

            run.font.size = Pt(body_cfg.get("font_size", 12))
            run.font.name = font_cfg.get("english", "Times New Roman")
            set_east_asian_font(run, font_cfg.get("chinese", "宋体"))

            p.paragraph_format.left_indent = Pt(24)
            p.paragraph_format.first_line_indent = Pt(-12)

    def _add_page_break(self, doc: Document) -> None:
        """添加分页符"""
        doc.add_page_break()

    def _add_image_placeholder(self, doc: Document, section: dict[str, Any]) -> None:
        """添加图片占位符"""
        caption = section.get("caption", "图片")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[{caption}]")
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.italic = True

    def _add_toc(self, doc: Document) -> None:
        """添加目录占位符"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("（请在 Word 中右键点击此处，选择 更新域 以生成目录）")
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(128, 128, 128)

    @staticmethod
    def _merge_config(base: dict, override: dict) -> dict:
        """深度合并配置"""
        result = copy.deepcopy(base)
        for key, value in override.items():
            if key.startswith("_"):
                continue
            if key in result and isinstance(result[key], dict) and isinstance(value, dict):
                result[key] = DocGenerator._merge_config(result[key], value)
            else:
                result[key] = copy.deepcopy(value)
        return result


def generate_doc_from_structure(
    structure: dict[str, Any],
    output_path: str,
    config: dict | None = None,
) -> str:
    """便捷函数：从结构化数据生成文档"""
    generator = DocGenerator(config)
    return generator.generate(structure, output_path)
