"""封面页生成器

根据元数据自动生成标准论文封面页。

支持的封面元素：
- 论文题目（中文/英文）
- 作者姓名
- 指导教师
- 学院/专业
- 学号
- 日期
- 学校名称
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class CoverPageGenerator:
    """封面页生成器"""

    def __init__(self, config=None):
        self.config = config or {}

    def generate(self, metadata, output_path, template="standard"):
        """生成封面页"""
        doc = Document()

        # 页面设置
        section = doc.sections[0]
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

        if template == "standard":
            self._generate_standard(doc, metadata)
        elif template == "graduate":
            self._generate_graduate(doc, metadata)
        else:
            self._generate_standard(doc, metadata)

        doc.save(output_path)
        return output_path

    def _generate_standard(self, doc, meta):
        """标准封面模板"""
        # 学校名称
        if meta.get("university"):
            self._add_paragraph(doc, meta["university"], font_size=22, bold=True, spacing_after=12)

        # 论文类型
        if meta.get("paper_type"):
            self._add_paragraph(doc, meta["paper_type"], font_size=18, bold=True, spacing_after=24)

        # 空行
        doc.add_paragraph()

        # 题目
        title = meta.get("title", "论文题目")
        self._add_labeled_line(doc, "题    目", title, font_size=16)

        # 英文题目
        if meta.get("title_en"):
            self._add_labeled_line(doc, "Title", meta["title_en"], font_size=14)

        doc.add_paragraph()

        # 信息表格
        info_items = [
            ("学    院", meta.get("college", "")),
            ("专    业", meta.get("major", "")),
            ("学    号", meta.get("student_id", "")),
            ("姓    名", meta.get("author", "")),
            ("指导教师", meta.get("advisor", "")),
        ]

        for label, value in info_items:
            if value:
                self._add_labeled_line(doc, label, value, font_size=14)

        # 空行
        for _ in range(3):
            doc.add_paragraph()

        # 日期
        if meta.get("date"):
            self._add_paragraph(doc, meta["date"], font_size=14, align="center")

    def _generate_graduate(self, doc, meta):
        """研究生封面模板"""
        # 空行
        for _ in range(2):
            doc.add_paragraph()

        # 学校名称
        if meta.get("university"):
            self._add_paragraph(doc, meta["university"], font_size=26, bold=True, spacing_after=6)

        # 学位论文类型
        degree = meta.get("degree", "硕士学位论文")
        self._add_paragraph(doc, degree, font_size=22, bold=True, spacing_after=24)

        # 空行
        doc.add_paragraph()

        # 题目（大字）
        title = meta.get("title", "论文题目")
        self._add_paragraph(doc, title, font_size=22, bold=True, spacing_after=12)

        # 英文题目
        if meta.get("title_en"):
            self._add_paragraph(doc, meta["title_en"], font_size=16, bold=False, spacing_after=24)

        # 空行
        for _ in range(2):
            doc.add_paragraph()

        # 信息
        info_items = [
            ("学科专业", meta.get("major", "")),
            ("研 究 生", meta.get("author", "")),
            ("指导教师", meta.get("advisor", "")),
        ]

        for label, value in info_items:
            if value:
                self._add_labeled_line(doc, label, value, font_size=14)

        # 空行
        for _ in range(4):
            doc.add_paragraph()

        # 日期
        if meta.get("date"):
            self._add_paragraph(doc, meta["date"], font_size=14, align="center")

    def _add_paragraph(self, doc, text, font_size=12, bold=False, align="center", spacing_after=6):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.name = "宋体"
        self._set_east_asian_font(run, "黑体" if bold else "宋体")
        p.alignment = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
        }.get(align, WD_ALIGN_PARAGRAPH.CENTER)
        p.paragraph_format.space_after = Pt(spacing_after)
        return p

    def _add_labeled_line(self, doc, label, value, font_size=12):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 标签
        label_run = p.add_run(f"{label}：")
        label_run.font.size = Pt(font_size)
        label_run.font.bold = True
        label_run.font.name = "宋体"
        self._set_east_asian_font(label_run, "宋体")

        # 下划线 + 值
        value_run = p.add_run(f"  {value}  ")
        value_run.font.size = Pt(font_size)
        value_run.font.bold = False
        value_run.font.name = "宋体"
        self._set_east_asian_font(value_run, "宋体")

        # 添加下划线
        rpr = value_run._element.get_or_add_rPr()
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rpr.append(u)

        p.paragraph_format.space_after = Pt(12)
        return p

    def _set_east_asian_font(self, run, font_name):
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = run._element.makeelement(qn("w:rFonts"), {})
            rpr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), font_name)


def generate_cover_from_config(config, output_path):
    """从配置生成封面"""
    meta = config.get("cover_page", {})
    generator = CoverPageGenerator(config)
    template = meta.get("template", "standard")
    return generator.generate(meta, output_path, template)
