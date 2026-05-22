"""毕业论文格式矫正集成测试

模拟真实毕业论文场景：
- 封面页（题目/作者/单位）
- 中英文摘要+关键词
- 多章节正文（含表格、图片占位）
- 参考文献
- 致谢

用毕业论文格式需求文档矫正。
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path


def create_thesis_document(path="input/thesis_sample.docx"):
    """生成一份格式不规范的毕业论文样例"""
    doc = Document()

    # 页面边距故意不对
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ===== 封面 =====
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("基于深度学习的中文文本情感分析研究")
    run.font.size = Pt(18)
    run.font.name = "Arial"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 应该居中

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("张  三")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    run = p.add_run("计算机科学与技术学院")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("2024年6月")
    run.font.size = Pt(12)

    doc.add_page_break()

    # ===== 中文摘要 =====
    p = doc.add_paragraph()
    run = p.add_run("摘  要")
    run.font.size = Pt(14)
    run.font.bold = False  # 应该加粗

    p = doc.add_paragraph()
    run = p.add_run(
        "随着互联网的快速发展，网络上的文本信息呈爆炸式增长。情感分析作为自然语言处理的重要分支，"
        "在舆情监控、产品评价分析等领域有着广泛的应用。本文提出了一种基于BERT预训练模型的中文文本"
        "情感分析方法，通过引入注意力机制和多层特征融合策略，有效提升了情感分类的准确率。实验结果"
        "表明，本文提出的方法在多个公开数据集上取得了最优性能。"
    )
    run.font.size = Pt(10.5)  # 应该小四
    run.font.name = "微软雅黑"  # 应该宋体

    p = doc.add_paragraph()
    run = p.add_run("关键词：深度学习；情感分析；BERT；注意力机制；自然语言处理")
    run.font.size = Pt(10.5)

    doc.add_page_break()

    # ===== 英文摘要 =====
    p = doc.add_paragraph()
    run = p.add_run("ABSTRACT")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    run = p.add_run(
        "With the rapid development of the Internet, text information on the web has "
        "experienced explosive growth. Sentiment analysis, as an important branch of "
        "natural language processing, has been widely applied in public opinion monitoring "
        "and product review analysis. This paper proposes a Chinese text sentiment analysis "
        "method based on the BERT pre-training model, which effectively improves the "
        "accuracy of sentiment classification by introducing attention mechanisms and "
        "multi-layer feature fusion strategies."
    )
    run.font.size = Pt(10.5)

    p = doc.add_paragraph()
    run = p.add_run("Key Words: Deep Learning; Sentiment Analysis; BERT; Attention Mechanism")
    run.font.size = Pt(10.5)

    doc.add_page_break()

    # ===== 第一章 =====
    p = doc.add_paragraph()
    run = p.add_run("第一章 绪论")
    run.font.size = Pt(14)
    run.font.bold = False
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    run = p.add_run("1.1 研究背景及意义")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(
        "近年来，随着社交媒体和电子商务的蓬勃发展，用户生成内容呈指数级增长。"
        "这些海量文本数据中蕴含着丰富的情感信息，如何自动识别和分析这些情感信息，"
        "成为了自然语言处理领域的重要研究课题。情感分析技术能够帮助企业了解用户对产品"
        "的满意度，帮助政府机构监测网络舆情，具有重要的理论价值和实际意义。"
    )
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("1.2 国内外研究现状")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(
        "情感分析的研究可以追溯到早期基于规则和词典的方法。随着深度学习技术的发展，"
        "基于神经网络的情感分析方法逐渐成为主流。"
    )
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("1.2.1 基于词典的方法")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(
        "基于情感词典的方法通过构建情感词典，利用词典中的情感极性信息对文本进行情感分类。"
        "这种方法简单直观，但依赖于词典的质量和覆盖范围。"
    )
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("1.2.2 基于深度学习的方法")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(
        "深度学习方法能够自动学习文本的特征表示，避免了人工特征工程的繁琐过程。"
        "常用的深度学习模型包括CNN、RNN、LSTM和Transformer等。"
    )
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("1.3 本文主要工作")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(
        "本文的主要贡献包括：（1）提出了一种基于BERT的中文情感分析模型；"
        "（2）引入了多头注意力机制增强特征提取能力；（3）在多个数据集上验证了方法的有效性。"
    )
    run.font.size = Pt(12)

    # ===== 第二章 =====
    p = doc.add_paragraph()
    run = p.add_run("第二章 相关理论与技术")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    run = p.add_run("2.1 预训练语言模型")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(
        "预训练语言模型是近年来自然语言处理领域的重大突破。BERT（Bidirectional Encoder "
        "Representations from Transformers）通过在大规模语料上进行预训练，学习到了丰富的"
        "语言知识，可以有效提升下游任务的性能。"
    )
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("2.2 注意力机制")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(
        "注意力机制能够使模型关注输入中最相关的部分，在自然语言处理中得到了广泛应用。"
        "自注意力机制是Transformer架构的核心组件。"
    )
    run.font.size = Pt(12)

    # 表标题（编号不对）
    p = doc.add_paragraph()
    run = p.add_run("表2 BERT模型参数对比")
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 应该居中

    # 简单表格
    table = doc.add_table(rows=4, cols=4)
    headers = ["模型", "层数", "隐藏维度", "参数量"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["BERT-Base", "12", "768", "110M"],
        ["BERT-Large", "24", "1024", "340M"],
        ["RoBERTa", "12", "768", "125M"],
    ]
    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            table.rows[r + 1].cells[c].text = val

    # 图标题
    p = doc.add_paragraph()
    run = p.add_run("图2 BERT模型架构")
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run("2.3 本章小结")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(
        "本章介绍了预训练语言模型和注意力机制的基本原理，为后续章节的方法设计奠定了理论基础。"
    )
    run.font.size = Pt(12)

    # ===== 第三章 =====
    p = doc.add_paragraph()
    run = p.add_run("第三章 基于BERT的情感分析模型")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    run = p.add_run("3.1 模型架构")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(
        "本文提出的模型基于BERT-Base架构，在其基础上添加了多头注意力层和全连接层。"
        "模型的输入为中文文本的token序列，输出为情感类别概率分布。"
    )
    run.font.size = Pt(12)

    # 表标题
    p = doc.add_paragraph()
    run = p.add_run("表1 模型超参数设置")
    run.font.size = Pt(10)

    table2 = doc.add_table(rows=5, cols=2)
    params = [
        ["学习率", "2e-5"],
        ["批次大小", "32"],
        ["最大序列长度", "512"],
        ["训练轮数", "10"],
    ]
    table2.rows[0].cells[0].text = "参数"
    table2.rows[0].cells[1].text = "值"
    for r, (k, v) in enumerate(params):
        table2.rows[r + 1].cells[0].text = k
        table2.rows[r + 1].cells[1].text = v

    p = doc.add_paragraph()
    run = p.add_run("3.2 训练策略")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(
        "模型采用Adam优化器，学习率设置为2e-5，使用线性学习率衰减策略。"
        "训练过程中采用梯度累积和混合精度训练技术以提升训练效率。"
    )
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("3.3 本章小结")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(
        "本章详细描述了基于BERT的情感分析模型的架构设计和训练策略。"
    )
    run.font.size = Pt(12)

    # ===== 参考文献 =====
    p = doc.add_paragraph()
    run = p.add_run("参考文献")
    run.font.size = Pt(14)
    run.font.bold = True

    refs = [
        "[1] Devlin J, Chang M W, Lee K, et al. BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding[J]. NAACL, 2019: 4171-4186.",
        "[2] 王小明, 李红. 基于深度学习的中文情感分析综述[J]. 计算机学报, 2021, 44(6): 1125-1148.",
        "[3] Liu Y, Ott M, Goyal N, et al. RoBERTa: A Robustly Optimized BERT Pretraining Approach[J]. arXiv preprint arXiv:1907.11692, 2019.",
        "[4] 张三, 李四. 基于注意力机制的文本情感分类方法[J]. 软件学报, 2022, 33(8): 2891-2905.",
        "[5] Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[C]//NeurIPS. 2017: 5998-6008.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(10)

    # ===== 致谢 =====
    p = doc.add_paragraph()
    run = p.add_run("致  谢")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    run = p.add_run(
        "感谢我的导师在论文写作过程中给予的悉心指导和帮助。感谢实验室的同学们"
        "在研究过程中提供的宝贵意见和建议。感谢家人的理解和支持。"
    )
    run.font.size = Pt(12)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"毕业论文样例已生成: {path}")
    return path


def create_thesis_requirement(path="input/thesis_requirement.txt"):
    """生成毕业论文格式要求文档"""
    content = """XX大学毕业论文（设计）格式要求

一、论文题目
论文题目用黑体，二号字，居中，加粗。

二、作者信息
作者姓名用宋体，小四号，居中。
作者单位用宋体，五号，居中。

三、摘要
中文摘要标题用黑体，三号字，居中，加粗。
中文摘要正文用宋体，小四号，1.5倍行距。
英文摘要标题用Times New Roman，三号字，居中，加粗。
英文摘要正文用Times New Roman，小四号，1.5倍行距。

四、关键词
关键词标题加粗，关键词内容用宋体，小四号。

五、一级标题
一级标题用黑体，二号字，居中，加粗。

六、二级标题
二级标题用黑体，三号字，左对齐，加粗。

七、三级标题
三级标题用黑体，四号字，左对齐，加粗。

八、正文
正文用宋体，小四号（12pt），1.5倍行距，首行缩进2字符，两端对齐。

九、图标题
图标题用宋体，五号，居中。

十、表标题
表标题用宋体，五号，居中。

十一、参考文献
参考文献标题用黑体，三号字，居中，加粗。
参考文献正文用宋体，五号，1.25倍行距，悬挂缩进。

十二、致谢
致谢标题同一级标题格式。
致谢正文同正文格式。

十三、页边距
上边距2.54cm，下边距2.54cm，左边距3.17cm，右边距3.17cm。

十四、页眉页脚
页眉显示论文题目，宋体，五号，居中，有下划线。
页脚居中显示页码。
"""
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    Path(path).write_text(content, encoding="utf-8")
    print(f"毕业论文格式要求已生成: {path}")
    return path


def main():
    print("=" * 60)
    print("  毕业论文格式矫正 - 集成测试")
    print("=" * 60)

    # 确保目录存在
    for d in ("input", "output", "template"):
        Path(d).mkdir(exist_ok=True)

    # 1. 生成样例论文和需求文档
    print("\n[1/4] 生成样例毕业论文...")
    thesis_path = create_thesis_document()

    print("\n[2/4] 生成格式要求文档...")
    req_path = create_thesis_requirement()

    # 2. 确保模板存在
    template_path = Path("template/template.docx")
    if not template_path.exists():
        print("\n[3/4] 创建默认模板...")
        from test_corrector import create_default_template
        create_default_template(str(template_path))

    # 3. 运行矫正
    print("\n[4/4] 运行格式矫正器（带需求文档）...")
    from main import PaperFormatCorrector

    corrector = PaperFormatCorrector()
    corrector.apply_requirement(req_path)
    report = corrector.process_single(thesis_path, "output/thesis_formatted.docx")

    if report:
        print(f"\n{'=' * 60}")
        print("矫正完成！处理统计：")
        print(f"  矫正段落: {report['paragraphs_corrected']}")
        print(f"  标题矫正: {report['headings_fixed']}")
        print(f"  正文矫正: {report['body_fixed']}")
        print(f"  表格格式化: {report.get('tables_formatted', 0)}")
        print(f"  图片居中: {report.get('images_centered', 0)}")
        if report.get("fig_table_issues"):
            print(f"  图表修正: {len(report['fig_table_issues'])} 项")
            for issue in report["fig_table_issues"]:
                print(f"    - {issue}")
        print(f"{'=' * 60}")

    print("\n输出文件: output/thesis_formatted.docx")


if __name__ == "__main__":
    main()
